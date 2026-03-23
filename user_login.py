import os
os.environ["STREAMLIT_WATCHDOG"] = "false"
import json
import random
import string
import re
import asyncio
import io
import urllib.parse
import base64
from io import BytesIO
from collections import Counter
from datetime import datetime
import time

# Third-party library imports
import streamlit as st
import streamlit.components.v1 as components
from base64 import b64encode
import requests
import fitz
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import altair as alt
from PIL import Image
from pdf2image import convert_from_path
from dotenv import load_dotenv
from nltk.stem import WordNetLemmatizer
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from xhtml2pdf import pisa
from pydantic import BaseModel
from streamlit_pdf_viewer import pdf_viewer

# Heavy libraries - loaded with caching
import torch

# Langchain & Embeddings

from langchain_text_splitters import CharacterTextSplitter 
from langchain_community.vectorstores import FAISS 
from langchain_community.embeddings import HuggingFaceEmbeddings 
from langchain_groq import ChatGroq  # optional if you're using it













# Local project imports
from llm_manager import call_llm, load_groq_api_keys, get_healthy_keys, increment_key_usage, mark_key_failure
from db_manager import (
    db_manager,
    insert_candidate,
    get_top_domains_by_score,
    get_database_stats,
    detect_domain_from_title_and_description,
    get_domain_similarity
)
from user_login import (
    create_user_table,
    add_user,
    complete_registration,
    verify_user,
    get_logins_today,
    get_total_registered_users,
    log_user_action,
    username_exists,
    email_exists,
    is_valid_email,
    save_user_api_key,
    get_user_api_key,
    get_all_user_logs,
    generate_otp,
    send_email_otp,
    get_user_by_email,
    update_password_by_email,
    is_strong_password,
    domain_has_mx_record,
    create_login_token,
    verify_login_token,
    send_login_confirmation_email,
    get_email_by_username,
)

# ============================================================
# 💾 Persistent Storage Configuration for Streamlit Cloud
# ============================================================
# SQLite storage removed — data persists in Supabase PostgreSQL

# ── Cached DB helpers — prevent re-querying Supabase on every rerun ──────────
# These are the functions called in the script body (hero stats, admin panel,
# sidebar). Without caching they fire on EVERY widget interaction / tab click.

@st.cache_data(ttl=60)   # refresh hero counters every 60 s
def _cached_hero_stats():
    return (
        get_total_registered_users(),
        get_logins_today(),
        get_database_stats(),
    )

@st.cache_data(ttl=30)   # admin panel metrics — slightly fresher
def _cached_admin_metrics():
    return (
        get_total_registered_users(),
        get_logins_today(),
        get_all_user_logs(),
    )

@st.cache_data(ttl=300)  # API key rarely changes — 5-min cache per user
def _cached_user_api_key(username: str):
    return get_user_api_key(username)
# ─────────────────────────────────────────────────────────────────────────────

def html_to_pdf_bytes(html_string):
    styled_html = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: 400mm 297mm;  /* Original custom large page size */
                margin-top: 10mm;
                margin-bottom: 10mm;
                margin-left: 10mm;
                margin-right: 10mm;
            }}
            body {{
                font-size: 14pt;
                font-family: "Segoe UI", "Helvetica", sans-serif;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2, h3 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 10px;
            }}
            .box {{
                padding: 8px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;  /* More elegant than full border */
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.5em;
            }}
            li {{
                margin-bottom: 5px;
            }}
        </style>
    </head>
    <body>
        {html_string}
    </body>
    </html>
    """

    pdf_io = BytesIO()
    pisa.CreatePDF(styled_html, dest=pdf_io)
    pdf_io.seek(0)
    return pdf_io

def generate_cover_letter_from_resume_builder():
    name = st.session_state.get("name", "")
    job_title = st.session_state.get("job_title", "")
    summary = st.session_state.get("summary", "")
    skills = st.session_state.get("skills", "")
    location = st.session_state.get("location", "")
    today_date = datetime.today().strftime("%B %d, %Y")

    # ✅ Input boxes for contact info
    company = st.text_input("🏢 Target Company", placeholder="e.g., Google")
    linkedin = st.text_input("🔗 LinkedIn URL", placeholder="e.g., https://linkedin.com/in/username")
    email = st.text_input("📧 Email", placeholder="e.g., you@example.com")
    mobile = st.text_input("📞 Mobile Number", placeholder="e.g., +91 9876543210")

    # ✅ Button to prevent relooping
    if st.button("✉️ Generate Cover Letter"):
        # ✅ Validate input before generating
        if not all([name, job_title, summary, skills, company, linkedin, email, mobile]):
            st.warning("⚠️ Please fill in all fields including LinkedIn, email, and mobile.")
            return

        prompt = f"""
You are a world-class executive cover letter writer with 20+ years of experience helping candidates land roles at top-tier companies.

Write a compelling, personalized cover letter using the candidate information below.

COVER LETTER STRUCTURE:
1. **Header** — Date, Hiring Manager, Company Name
2. **Opening Paragraph** — Hook with specific value proposition; mention the role and why this company specifically
3. **Core Value Paragraph** — Top 2 achievements from their background (quantified if possible); connect directly to company's likely needs
4. **Skills-Fit Paragraph** — Bridge candidate skills to the role requirements; show cultural awareness
5. **Closing Paragraph** — Confident call to action; express enthusiasm; professional closing

TONE: Professional, confident, specific — NOT generic. Avoid clichés like "I am passionate about..." or "I believe I would be a great fit."
INCLUDE the contact block at the very top: Name, LinkedIn, Email, Phone.
ENSURE company name appears only once (in the header or salutation).
LENGTH: 3 short-to-medium paragraphs. Maximum 350 words.

### CANDIDATE DETAILS:
- Full Name: {name}
- Job Title Applying For: {job_title}
- Professional Summary: {summary}
- Key Skills: {skills}
- Location: {location}
- Date: {today_date}

### COMPANY DETAILS:
- Target Company: {company}
- Candidate LinkedIn: {linkedin}
- Candidate Email: {email}
- Candidate Phone: {mobile}

### INSTRUCTIONS:
- Return PLAIN TEXT ONLY — no HTML, no markdown, no asterisks
- Do NOT mention company name more than once
- Make it feel tailored — not templated
- End with: "Sincerely," followed by the candidate's full name
"""

        # ✅ Call LLM
        with st.spinner("✉️ Generating cover letter..."):
            try:
                cover_letter = call_llm(prompt, session=st.session_state).strip()
            except Exception as e:
                st.error(f"❌ Failed to generate cover letter: {e}")
                return

        # ✅ Store plain text
        st.session_state["cover_letter"] = cover_letter

        # ✅ Build HTML wrapper for preview (safe)
        cover_letter_html = f"""
        <div style="font-family: Georgia, serif; font-size: 13pt; line-height: 1.6; 
                    color: #000; background: #fff; padding: 25px; 
                    border-radius: 8px; box-shadow: 0px 2px 6px rgba(0,0,0,0.1); 
                    max-width: 800px; margin: auto;">
            <div style="text-align:center; margin-bottom:15px;">
                <div style="font-size:18pt; font-weight:bold; color:#003366;">{name}</div>
                <div style="font-size:14pt; color:#555;">{job_title}</div>
                <div style="font-size:10pt; margin-top:5px;">
                    <a href="{linkedin}" style="color:#003366;">{linkedin}</a><br/>
                    📧 {email} | 📞 {mobile}
                </div>
            </div>
            <hr/>
            <pre style="white-space: pre-wrap; font-family: Georgia, serif; font-size: 12pt; color:#000;">
{cover_letter}
            </pre>
        </div>
        """

        st.session_state["cover_letter_html"] = cover_letter_html

        # ✅ Show nicely in Streamlit
        st.markdown(cover_letter_html, unsafe_allow_html=True)

# ------------------- Initialize -------------------
# ✅ Initialize database in persistent storage
create_user_table()

# ------------------- Login Token Handler -------------------
# Runs on every page load. If the URL carries ?login_token=<token>,
# validate it and complete the login — this is the "Yes, it's me" click.
_login_token = st.query_params.get("login_token", None)
if _login_token and not st.session_state.get("authenticated", False):
    _username, _groq_key = verify_login_token(_login_token)
    if _username:
        st.session_state.authenticated = True
        st.session_state.username = _username
        st.session_state.user_groq_key = _groq_key or ""
        st.session_state.login_stage = "credentials"
        st.session_state.pending_login_username = None
        log_user_action(_username, "login")
        # Remove the token from the URL so a refresh doesn't re-submit it
        st.query_params.clear()
        st.rerun()
    else:
        # Token invalid, expired, or already used — show a one-time error
        if "token_error_shown" not in st.session_state:
            st.session_state.token_error_shown = True
            st.query_params.clear()
            notify("login", "error", "❌ This login link is invalid or has expired. Please sign in again.")
            st.rerun()

# ------------------- Tab-Specific Notification System -------------------
if "login_notification" not in st.session_state:
    st.session_state.login_notification = {"type": None, "text": None, "expires": 0.0}
if "register_notification" not in st.session_state:
    st.session_state.register_notification = {"type": None, "text": None, "expires": 0.0}

def notify(tab, msg_type, text, duration=3.0):
    """Show auto-disappearing message for specific tab (login/register)."""
    notification_key = f"{tab}_notification"
    st.session_state[notification_key] = {
        "type": msg_type,
        "text": text,
        "expires": time.time() + duration,
    }

def render_notification(tab):
    """Render notification in a fixed-height slot — button position never shifts."""
    notification_key = f"{tab}_notification"
    notif = st.session_state[notification_key]

    # Map type to inline style colours (avoids Streamlit's full-height alert boxes)
    _styles = {
        "success": ("rgba(52,211,153,0.13)", "rgba(52,211,153,0.28)", "#6ee7b7"),
        "error":   ("rgba(251,113,133,0.13)", "rgba(251,113,133,0.28)", "#fca5a5"),
        "warning": ("rgba(251,191,36,0.13)",  "rgba(251,191,36,0.28)",  "#fde68a"),
        "info":    ("rgba(56,189,248,0.13)",  "rgba(56,189,248,0.28)",  "#7dd3fc"),
    }

    # Always emit a min-height wrapper so nothing below shifts on empty state
    if notif["type"] and time.time() < notif["expires"]:
        bg, border, color = _styles.get(notif["type"], _styles["info"])
        st.markdown(
            f"""<div style='min-height:48px; display:flex; align-items:center;'>
                <div style='width:100%; padding:8px 14px; border-radius:8px;
                            background:{bg}; border:1px solid {border};
                            color:{color}; font-size:0.85rem; font-weight:500;
                            font-family:-apple-system,sans-serif; line-height:1.4;
                            white-space:normal; word-wrap:break-word; overflow:visible;'>
                    {notif["text"]}
                </div>
            </div>""",
            unsafe_allow_html=True
        )
    else:
        # Reserved space — invisible, same height
        st.markdown("<div style='height:48px;'></div>", unsafe_allow_html=True)


def display_timer(remaining_seconds, expired=False, key_suffix=""):
    """
    Display a server-synced timer with glassmorphism styling.
    Server-side validation ensures OTP expiry is accurately enforced.

    Args:
        remaining_seconds: Time remaining in seconds (server-calculated)
        expired: Whether the timer has expired
        key_suffix: Unique suffix for the timer component
    """
    minutes = remaining_seconds // 60
    seconds = remaining_seconds % 60

    if expired or remaining_seconds <= 0:
        st.markdown("""
        <div class='timer-display timer-expired' style="
            background: linear-gradient(135deg, rgba(255, 99, 71, 0.18) 0%, rgba(255, 99, 71, 0.08) 100%);
            backdrop-filter: blur(15px);
            -webkit-backdrop-filter: blur(15px);
            border: 2px solid rgba(255, 99, 71, 0.4);
            border-radius: 14px;
            padding: 16px 24px;
            margin: 20px 0;
            text-align: center;
            box-shadow: 0 4px 20px rgba(255, 99, 71, 0.15), inset 0 1px 0 rgba(255, 255, 255, 0.1);
        ">
            <span class='timer-text' style="
                color: #FF6347;
                font-size: 1.15em;
                font-weight: bold;
                font-family: 'Orbitron', sans-serif;
                text-shadow: 0 0 18px rgba(255, 99, 71, 0.5);
            ">⏱️ OTP Expired</span>
        </div>
        """, unsafe_allow_html=True)
    else:
        # Client-side countdown for UX, but server validates on action
        st.components.v1.html(f"""
        <div class='timer-display' id='timer-{key_suffix}' style="
            background: linear-gradient(135deg, rgba(255, 215, 0, 0.18) 0%, rgba(255, 165, 0, 0.08) 100%);
            backdrop-filter: blur(15px);
            -webkit-backdrop-filter: blur(15px);
            border: 2px solid rgba(255, 215, 0, 0.4);
            border-radius: 14px;
            padding: 16px 24px;
            margin: 20px 0;
            text-align: center;
            box-shadow: 0 4px 20px rgba(255, 215, 0, 0.15), inset 0 1px 0 rgba(255, 255, 255, 0.1);
        ">
            <span class='timer-text' style="
                color: #FFD700;
                font-size: 1.15em;
                font-weight: bold;
                font-family: 'Orbitron', sans-serif;
                text-shadow: 0 0 18px rgba(255, 215, 0, 0.5);
            ">⏱️ Time Remaining: <span id='countdown-{key_suffix}'>{minutes:02d}:{seconds:02d}</span></span>
        </div>
        <script>
        (function() {{
            let remaining = {remaining_seconds};
            const countdownEl = document.getElementById('countdown-{key_suffix}');
            const timerEl = document.getElementById('timer-{key_suffix}');

            const interval = setInterval(() => {{
                remaining--;
                if (remaining <= 0) {{
                    clearInterval(interval);
                    if (timerEl) {{
                        timerEl.style.background = 'linear-gradient(135deg, rgba(255, 99, 71, 0.18) 0%, rgba(255, 99, 71, 0.08) 100%)';
                        timerEl.style.border = '2px solid rgba(255, 99, 71, 0.4)';
                        timerEl.innerHTML = "<span style='color: #FF6347; font-size: 1.15em; font-weight: bold; font-family: Orbitron, sans-serif; text-shadow: 0 0 18px rgba(255, 99, 71, 0.5);'>⏱️ OTP Expired</span>";
                    }}
                }} else {{
                    const mins = Math.floor(remaining / 60);
                    const secs = remaining % 60;
                    if (countdownEl) {{
                        countdownEl.textContent = `${{mins.toString().padStart(2, '0')}}:${{secs.toString().padStart(2, '0')}}`;
                    }}
                }}
            }}, 1000);
        }})();
        </script>
        """, height=80)

# ------------------- Initialize Session State -------------------
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "username" not in st.session_state:
    st.session_state.username = None
if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

# Login confirmation flow states
if "login_stage" not in st.session_state:
    # "credentials"  → normal sign-in form
    # "awaiting_email" → credentials verified, confirmation email sent, waiting
    st.session_state.login_stage = "credentials"
if "pending_login_username" not in st.session_state:
    st.session_state.pending_login_username = None

# Forgot password session states
if "reset_stage" not in st.session_state:
    st.session_state.reset_stage = "none"
if "reset_email" not in st.session_state:
    st.session_state.reset_email = ""
if "reset_otp" not in st.session_state:
    st.session_state.reset_otp = ""
if "reset_otp_time" not in st.session_state:
    st.session_state.reset_otp_time = 0

# Validation message state for register form (populated by on_change callbacks)
# _email_msg, _user_msg, _pass_msg are initialised inside the register form block

# ------------------- CSS Styling -------------------
st.markdown("""
<style>
/* ═══════════════════════════════════════════════════════════════
   HIRELYZER — Premium Apple-Style Dark Theme
   Font Stack: SF Pro Display → Segoe UI → Roboto → sans-serif
   Design Language: Glassmorphism · Soft gradients · Refined motion
   ═══════════════════════════════════════════════════════════════ */

@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=DM+Sans:wght@300;400;500;600;700&family=Orbitron:wght@600&display=swap');

:root {
    --bg-primary:       #080c12;
    --bg-secondary:     #0e1420;
    --bg-tertiary:      #141c2b;
    --surface-01:       rgba(255,255,255,0.04);
    --surface-02:       rgba(255,255,255,0.07);
    --surface-hover:    rgba(255,255,255,0.10);
    --border-subtle:    rgba(255,255,255,0.07);
    --border-accent:    rgba(99,179,237,0.30);
    --accent-blue:      #4fa3e3;
    --accent-cyan:      #38bdf8;
    --accent-violet:    #818cf8;
    --accent-emerald:   #34d399;
    --accent-amber:     #fbbf24;
    --accent-rose:      #fb7185;
    --text-primary:     #f0f4f8;
    --text-secondary:   #94a3b8;
    --text-muted:       #4a5568;
    --radius-sm:        8px;
    --radius-md:        14px;
    --radius-lg:        20px;
    --radius-xl:        28px;
    --shadow-glow-blue: 0 0 30px rgba(79,163,227,0.15);
    --shadow-card:      0 8px 40px rgba(0,0,0,0.45), 0 1px 0 rgba(255,255,255,0.06) inset;
    --font-sans:        -apple-system, BlinkMacSystemFont, "SF Pro Display", "DM Sans", "Segoe UI", Roboto, sans-serif;
    --transition-fast:  0.18s cubic-bezier(0.4,0,0.2,1);
    --transition-base:  0.28s cubic-bezier(0.4,0,0.2,1);
    --transition-slow:  0.45s cubic-bezier(0.4,0,0.2,1);
}

/* ── Base Reset ── */
html, body, [class*="css"], .stApp {
    font-family: var(--font-sans) !important;
    background-color: var(--bg-primary) !important;
    color: var(--text-primary) !important;
    -webkit-font-smoothing: antialiased;
    -moz-osx-font-smoothing: grayscale;
    scroll-behavior: smooth;
}

/* ── Force unified background — eliminates the horizontal seam glitch ── */
.stApp > header,
.stApp [data-testid="stAppViewContainer"],
.stApp [data-testid="stAppViewBlockContainer"],
.stApp [data-testid="block-container"],
.main,
.main > div,
section[data-testid="stMain"],
section[data-testid="stMain"] > div {
    background-color: var(--bg-primary) !important;
    background: var(--bg-primary) !important;
}

/* ── Streamlit top toolbar / header bar — the actual seam source ── */
header[data-testid="stHeader"],
header[data-testid="stHeader"] > div,
header[data-testid="stHeader"] > div > div,
.stApp header,
div[data-testid="stToolbar"],
div[data-testid="stStatusWidget"] {
    background-color: var(--bg-primary) !important;
    background: var(--bg-primary) !important;
    border-bottom: none !important;
    box-shadow: none !important;
}

/* ── Remove the decorative top colour bar Streamlit injects ── */
div[data-testid="stDecoration"],
#stDecoration {
    background: var(--bg-primary) !important;
    background-image: none !important;
    display: none !important;
}

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: var(--bg-secondary); }
::-webkit-scrollbar-thumb { background: rgba(79,163,227,0.35); border-radius: 99px; }
::-webkit-scrollbar-thumb:hover { background: rgba(79,163,227,0.6); }

/* ── Main container ── */
.main .block-container {
    padding: 1.5rem 2rem 3rem !important;
    max-width: 1280px;
}

/* ══════════════════════════════════════
   FADE ANIMATIONS
   ══════════════════════════════════════ */
@keyframes fadein  { from { opacity: 0; transform: translateY(6px); } to { opacity: 1; transform: translateY(0); } }
@keyframes fadeout { from { opacity: 1; } to { opacity: 0; } }
@keyframes fadeSlideUp { from { opacity: 0; transform: translateY(16px); } to { opacity: 1; transform: translateY(0); } }
@keyframes pulseGlow {
    0%, 100% { box-shadow: var(--shadow-card); }
    50%       { box-shadow: var(--shadow-card), var(--shadow-glow-blue); }
}
@keyframes shimmerSlide {
    0%   { transform: translateX(-100%) skewX(-12deg); }
    100% { transform: translateX(220%) skewX(-12deg); }
}
@keyframes glassShimmer {
    0%   { transform: translateX(-100%) skewX(-15deg); }
    100% { transform: translateX(200%) skewX(-15deg); }
}
@keyframes slideIn {
    0%   { transform: translateX(-50px); opacity: 0; }
    100% { transform: translateX(0); opacity: 1; }
}
@keyframes floatUp {
    0%, 100% { transform: translateY(0); }
    50%       { transform: translateY(-6px); }
}

/* ── Animated cards removed — replaced by premium hero section ── */

/* ══════════════════════════════════════
   STREAMLIT ALERT TOASTS
   ══════════════════════════════════════ */
div.stAlert {
    border-radius: var(--radius-md) !important;
    padding: 12px 18px !important;
    animation: fadein 0.3s ease, fadeout 0.3s 2.7s ease;
    backdrop-filter: blur(20px);
    border: 1px solid var(--border-subtle);
    font-size: 0.875rem;
    font-family: var(--font-sans) !important;
}

/* ══════════════════════════════════════
   GLASSMORPHISM CARD — base class
   ══════════════════════════════════════ */
.glass-card {
    background: var(--surface-01);
    backdrop-filter: blur(24px) saturate(180%);
    -webkit-backdrop-filter: blur(24px) saturate(180%);
    border: 1px solid var(--border-subtle);
    border-radius: var(--radius-lg);
    box-shadow: var(--shadow-card);
    transition: transform var(--transition-base), box-shadow var(--transition-base);
    position: relative;
    overflow: hidden;
}
.glass-card::before {
    content: '';
    position: absolute;
    inset: 0;
    background: linear-gradient(135deg, rgba(255,255,255,0.05) 0%, transparent 60%);
    pointer-events: none;
    border-radius: inherit;
}
.glass-card:hover {
    transform: translateY(-4px);
    box-shadow: var(--shadow-card), 0 0 50px rgba(79,163,227,0.10);
    border-color: var(--border-accent);
}

/* ══════════════════════════════════════
   LOGIN / AUTH CARD
   ══════════════════════════════════════ */
.login-card {
    background: linear-gradient(160deg,
        rgba(14,20,32,0.95) 0%,
        rgba(8,12,18,0.98) 100%);
    backdrop-filter: blur(32px) saturate(160%);
    -webkit-backdrop-filter: blur(32px) saturate(160%);
    border: 1px solid rgba(99,179,237,0.18);
    border-radius: var(--radius-xl);
    padding: 28px 32px 36px;
    box-shadow: var(--shadow-card), 0 0 60px rgba(79,163,227,0.07);
    transition: all var(--transition-slow);
    position: relative;
    overflow: hidden;
    animation: fadeSlideUp 0.7s cubic-bezier(0.22,1,0.36,1) forwards;
}
.login-card::after {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 60%; height: 100%;
    background: linear-gradient(90deg, transparent, rgba(79,163,227,0.06), transparent);
    animation: shimmerSlide 3.5s ease-in-out infinite;
}
.login-card:hover {
    border-color: rgba(99,179,237,0.32);
    box-shadow: var(--shadow-card), 0 0 80px rgba(79,163,227,0.12);
}

/* ══════════════════════════════════════
   TEXT INPUTS & TEXTAREAS
   ══════════════════════════════════════ */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stSelectbox > div > div > div {
    background: rgba(255,255,255,0.04) !important;
    color: var(--text-primary) !important;
    border: 1px solid rgba(255,255,255,0.10) !important;
    border-radius: var(--radius-sm) !important;
    font-family: var(--font-sans) !important;
    font-size: 0.9rem !important;
    padding: 10px 14px !important;
    transition: border-color var(--transition-fast), box-shadow var(--transition-fast) !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: rgba(79,163,227,0.5) !important;
    box-shadow: 0 0 0 3px rgba(79,163,227,0.12) !important;
    outline: none !important;
}
.stTextInput > div > div > input:hover,
.stTextArea > div > div > textarea:hover {
    border-color: rgba(79,163,227,0.28) !important;
}
.stTextInput > label,
.stTextArea > label,
.stSelectbox > label,
.stSlider > label,
.stFileUploader > label {
    color: var(--text-secondary) !important;
    font-size: 0.8rem !important;
    font-weight: 500 !important;
    letter-spacing: 0.03em !important;
    text-transform: uppercase !important;
    font-family: var(--font-sans) !important;
}

/* ══════════════════════════════════════
   BUTTONS — Primary style
   ══════════════════════════════════════ */
.stButton > button {
    background: linear-gradient(135deg,
        rgba(56,189,248,0.18) 0%,
        rgba(79,163,227,0.12) 100%) !important;
    color: var(--accent-cyan) !important;
    border: 1px solid rgba(56,189,248,0.3) !important;
    border-radius: var(--radius-sm) !important;
    font-family: var(--font-sans) !important;
    font-weight: 600 !important;
    font-size: 0.875rem !important;
    letter-spacing: 0.02em !important;
    padding: 10px 22px !important;
    backdrop-filter: blur(12px) !important;
    box-shadow: 0 2px 12px rgba(56,189,248,0.08), inset 0 1px 0 rgba(255,255,255,0.08) !important;
    transition: all var(--transition-fast) !important;
    position: relative;
    overflow: hidden;
}
.stButton > button:hover {
    background: linear-gradient(135deg,
        rgba(56,189,248,0.28) 0%,
        rgba(79,163,227,0.22) 100%) !important;
    border-color: rgba(56,189,248,0.55) !important;
    box-shadow: 0 4px 20px rgba(56,189,248,0.18), inset 0 1px 0 rgba(255,255,255,0.12) !important;
    transform: translateY(-2px) !important;
    color: #e0f6ff !important;
}
.stButton > button:active {
    transform: translateY(0px) !important;
    box-shadow: 0 1px 6px rgba(56,189,248,0.10) !important;
}

/* ══════════════════════════════════════
   DOWNLOAD BUTTONS
   ══════════════════════════════════════ */
.stDownloadButton > button {
    background: linear-gradient(135deg,
        rgba(52,211,153,0.16) 0%,
        rgba(52,211,153,0.08) 100%) !important;
    color: var(--accent-emerald) !important;
    border: 1px solid rgba(52,211,153,0.28) !important;
    border-radius: var(--radius-sm) !important;
    font-family: var(--font-sans) !important;
    font-weight: 600 !important;
    font-size: 0.875rem !important;
    transition: all var(--transition-fast) !important;
}
.stDownloadButton > button:hover {
    background: linear-gradient(135deg,
        rgba(52,211,153,0.26) 0%,
        rgba(52,211,153,0.16) 100%) !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 4px 18px rgba(52,211,153,0.15) !important;
}

/* ══════════════════════════════════════
   METRICS
   ══════════════════════════════════════ */
div[data-testid="metric-container"] {
    background: var(--surface-01) !important;
    border: 1px solid var(--border-subtle) !important;
    border-radius: var(--radius-md) !important;
    padding: 18px 20px !important;
    backdrop-filter: blur(16px) !important;
    transition: all var(--transition-base) !important;
    animation: fadeSlideUp 0.5s ease forwards;
}
div[data-testid="metric-container"]:hover {
    border-color: var(--border-accent) !important;
    background: var(--surface-02) !important;
    transform: translateY(-3px) !important;
    box-shadow: var(--shadow-glow-blue) !important;
}
div[data-testid="metric-container"] label {
    color: var(--text-secondary) !important;
    font-size: 0.75rem !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.06em !important;
    font-family: var(--font-sans) !important;
}
div[data-testid="metric-container"] [data-testid="stMetricValue"] {
    color: var(--text-primary) !important;
    font-size: 1.75rem !important;
    font-weight: 700 !important;
    letter-spacing: -0.02em !important;
    line-height: 1.2 !important;
    font-family: var(--font-sans) !important;
}

/* ══════════════════════════════════════
   TABS
   ══════════════════════════════════════ */
.stTabs [data-baseweb="tab-list"] {
    gap: 4px !important;
    background: rgba(255,255,255,0.03) !important;
    padding: 5px !important;
    border-radius: var(--radius-md) !important;
    border: 1px solid var(--border-subtle) !important;
    backdrop-filter: blur(16px) !important;
}
.stTabs [data-baseweb="tab"] {
    border-radius: var(--radius-sm) !important;
    color: var(--text-secondary) !important;
    font-family: var(--font-sans) !important;
    font-weight: 500 !important;
    font-size: 0.875rem !important;
    padding: 9px 18px !important;
    transition: all var(--transition-fast) !important;
    border: none !important;
    background: transparent !important;
}
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg,
        rgba(56,189,248,0.18) 0%,
        rgba(79,163,227,0.12) 100%) !important;
    color: var(--accent-cyan) !important;
    font-weight: 600 !important;
    border: 1px solid rgba(56,189,248,0.25) !important;
    box-shadow: 0 2px 10px rgba(56,189,248,0.12) !important;
}
.stTabs [data-baseweb="tab-highlight"] { display: none !important; }
.stTabs [data-baseweb="tab-panel"] {
    background: transparent !important;
    padding: 20px 0 !important;
}

/* ══════════════════════════════════════
   DATAFRAME / TABLE
   ══════════════════════════════════════ */
.dataframe, .stDataFrame {
    border-radius: var(--radius-md) !important;
    overflow: hidden !important;
    border: 1px solid var(--border-subtle) !important;
}
.stDataFrame [data-testid="stDataFrameResizable"] {
    background: var(--bg-secondary) !important;
}

/* ══════════════════════════════════════
   SIDEBAR
   ══════════════════════════════════════ */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg,
        rgba(10,14,22,0.98) 0%,
        rgba(8,12,18,1) 100%) !important;
    border-right: 1px solid var(--border-subtle) !important;
}
section[data-testid="stSidebar"] .block-container {
    padding: 1.5rem 1rem !important;
}
section[data-testid="stSidebar"] label {
    color: var(--text-secondary) !important;
    font-size: 0.8rem !important;
    text-transform: uppercase !important;
    letter-spacing: 0.05em !important;
    font-weight: 600 !important;
}
section[data-testid="stSidebar"] .stTextInput > div > div > input,
section[data-testid="stSidebar"] .stTextArea > div > div > textarea {
    background: rgba(255,255,255,0.03) !important;
    border-color: rgba(255,255,255,0.08) !important;
}

/* ══════════════════════════════════════
   EXPANDER
   ══════════════════════════════════════ */
.streamlit-expanderHeader {
    background: var(--surface-01) !important;
    border: 1px solid var(--border-subtle) !important;
    border-radius: var(--radius-md) !important;
    color: var(--text-primary) !important;
    font-family: var(--font-sans) !important;
    font-weight: 500 !important;
    transition: all var(--transition-fast) !important;
}
.streamlit-expanderHeader:hover {
    background: var(--surface-hover) !important;
    border-color: var(--border-accent) !important;
}
.streamlit-expanderContent {
    background: rgba(255,255,255,0.02) !important;
    border: 1px solid var(--border-subtle) !important;
    border-top: none !important;
    border-radius: 0 0 var(--radius-md) var(--radius-md) !important;
}

/* ══════════════════════════════════════
   SLIDERS
   ══════════════════════════════════════ */
.stSlider [data-baseweb="slider"] [role="slider"] {
    background: var(--accent-cyan) !important;
    border: 2px solid var(--bg-primary) !important;
    box-shadow: 0 0 0 3px rgba(56,189,248,0.3) !important;
}
.stSlider [data-baseweb="slider"] [data-testid="stTickBar"] > div {
    background: rgba(56,189,248,0.6) !important;
}

/* ══════════════════════════════════════
   FILE UPLOADER
   ══════════════════════════════════════ */
.stFileUploader > div {
    background: var(--surface-01) !important;
    border: 1.5px dashed rgba(79,163,227,0.3) !important;
    border-radius: var(--radius-lg) !important;
    transition: all var(--transition-base) !important;
}
.stFileUploader > div:hover {
    border-color: rgba(79,163,227,0.6) !important;
    background: rgba(79,163,227,0.04) !important;
    box-shadow: 0 0 40px rgba(79,163,227,0.07) !important;
}

/* ══════════════════════════════════════
   DIVIDER
   ══════════════════════════════════════ */
hr {
    border: none !important;
    border-top: 1px solid var(--border-subtle) !important;
    margin: 28px 0 !important;
}

/* ══════════════════════════════════════
   HEADINGS
   ══════════════════════════════════════ */
h1, h2, h3, h4, h5, h6,
.stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {
    font-family: var(--font-sans) !important;
    color: var(--text-primary) !important;
    letter-spacing: -0.02em !important;
    font-weight: 700 !important;
}
h1, .stMarkdown h1 { font-size: 2rem !important; }
h2, .stMarkdown h2 { font-size: 1.4rem !important; }
h3, .stMarkdown h3 {
    font-size: 1.1rem !important;
    color: var(--text-secondary) !important;
    font-weight: 600 !important;
}

/* ══════════════════════════════════════
   FEATURE CARDS (sidebar pre-login)
   ══════════════════════════════════════ */
.feature-card {
    background: var(--surface-01);
    backdrop-filter: blur(20px);
    -webkit-backdrop-filter: blur(20px);
    border: 1px solid var(--border-subtle);
    border-radius: var(--radius-md);
    padding: 18px 16px;
    margin-bottom: 12px;
    transition: all var(--transition-base);
    position: relative;
    overflow: hidden;
    animation: fadeSlideUp 0.5s ease forwards;
}
.feature-card::before {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 50%; height: 100%;
    background: linear-gradient(90deg, transparent, rgba(79,163,227,0.05), transparent);
    transition: left 0.6s ease;
}
.feature-card:hover { transform: translateY(-4px); border-color: var(--border-accent); }
.feature-card:hover::before { left: 150%; }
.feature-card h3 {
    color: var(--accent-cyan) !important;
    font-size: 0.9rem !important;
    font-weight: 600 !important;
    margin: 8px 0 4px !important;
}
.feature-card p {
    color: var(--text-secondary) !important;
    font-size: 0.78rem !important;
    line-height: 1.5 !important;
    margin: 0 !important;
}

/* ══════════════════════════════════════
   SLIDE MESSAGES (inline notifications)
   ══════════════════════════════════════ */
.slide-message {
    position: relative;
    overflow: hidden;
    margin: 12px 0;
    padding: 12px 18px;
    border-radius: var(--radius-md);
    font-weight: 500;
    font-size: 0.875rem;
    display: flex;
    align-items: center;
    justify-content: flex-start;
    gap: 10px;
    animation: fadein 0.4s cubic-bezier(0.34,1.56,0.64,1) forwards;
    backdrop-filter: blur(16px);
    -webkit-backdrop-filter: blur(16px);
    box-shadow: 0 4px 16px rgba(0,0,0,0.15), inset 0 1px 0 rgba(255,255,255,0.08);
    width: 100%;
    box-sizing: border-box;
    font-family: var(--font-sans) !important;
    transition: all var(--transition-fast);
    min-height: 46px;
}
.slide-message:hover { transform: translateY(-2px); }
.slide-message-text { flex: 1; position: relative; z-index: 2; word-wrap: break-word; }
.success-msg {
    background: linear-gradient(135deg, rgba(52,211,153,0.15) 0%, rgba(52,211,153,0.05) 100%);
    border: 1px solid rgba(52,211,153,0.30);
    color: #6ee7b7;
}
.error-msg {
    background: linear-gradient(135deg, rgba(251,113,133,0.15) 0%, rgba(251,113,133,0.05) 100%);
    border: 1px solid rgba(251,113,133,0.30);
    color: #fca5a5;
}
.info-msg {
    background: linear-gradient(135deg, rgba(56,189,248,0.15) 0%, rgba(56,189,248,0.05) 100%);
    border: 1px solid rgba(56,189,248,0.30);
    color: #7dd3fc;
}
.warn-msg {
    background: linear-gradient(135deg, rgba(251,191,36,0.15) 0%, rgba(251,191,36,0.05) 100%);
    border: 1px solid rgba(251,191,36,0.30);
    color: #fde68a;
}

/* ══════════════════════════════════════
   COUNTER GRID (landing page stats)
   ══════════════════════════════════════ */
.counter-grid {
    display: grid;
    grid-template-columns: repeat(2, 1fr);
    gap: 16px;
    padding: 24px 0;
    max-width: 520px;
    margin: 0 auto;
}
.counter-box {
    background: var(--surface-01);
    backdrop-filter: blur(20px);
    -webkit-backdrop-filter: blur(20px);
    border: 1px solid var(--border-subtle);
    border-radius: var(--radius-md);
    padding: 22px 18px;
    display: flex;
    flex-direction: column;
    justify-content: center;
    align-items: center;
    position: relative;
    overflow: hidden;
    transition: all var(--transition-base);
    animation: floatUp 4s ease-in-out infinite;
}
.counter-box::before {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 50%; height: 100%;
    background: linear-gradient(90deg, transparent, rgba(56,189,248,0.06), transparent);
    animation: shimmerSlide 3s infinite;
}
.counter-box:hover {
    transform: translateY(-6px) scale(1.02);
    border-color: var(--border-accent);
    box-shadow: 0 12px 40px rgba(56,189,248,0.10);
}
.counter-box:nth-child(1) { animation-delay: 0s; }
.counter-box:nth-child(2) { animation-delay: 0.6s; }
.counter-box:nth-child(3) { animation-delay: 1.2s; }
.counter-box:nth-child(4) { animation-delay: 1.8s; }
.counter-number {
    font-size: 2rem;
    font-weight: 700;
    color: var(--accent-cyan);
    letter-spacing: -0.03em;
    line-height: 1;
    position: relative;
    z-index: 2;
    font-family: var(--font-sans);
}
.counter-label {
    margin-top: 6px;
    font-size: 0.78rem;
    color: var(--text-secondary);
    font-weight: 500;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    position: relative;
    z-index: 2;
    font-family: var(--font-sans);
}

/* ══════════════════════════════════════
   ATS SECTION CARDS (analysis results)
   ══════════════════════════════════════ */
.ats-section-header {
    background: linear-gradient(135deg,
        rgba(56,189,248,0.14) 0%,
        rgba(79,163,227,0.08) 100%);
    border: 1px solid rgba(56,189,248,0.22);
    border-radius: var(--radius-sm) var(--radius-sm) 0 0;
    padding: 12px 18px;
    font-family: var(--font-sans) !important;
    font-weight: 700;
    font-size: 0.875rem;
    color: var(--accent-cyan);
    letter-spacing: 0.02em;
    text-transform: uppercase;
}
.ats-section-body {
    background: rgba(255,255,255,0.025);
    border: 1px solid rgba(255,255,255,0.06);
    border-top: none;
    border-radius: 0 0 var(--radius-sm) var(--radius-sm);
    padding: 16px 18px;
    color: var(--text-secondary);
    font-family: var(--font-sans) !important;
    font-size: 0.875rem;
    line-height: 1.65;
    margin-bottom: 14px;
}
.score-badge {
    display: inline-flex;
    align-items: center;
    background: linear-gradient(135deg, rgba(56,189,248,0.18) 0%, rgba(56,189,248,0.08) 100%);
    border: 1px solid rgba(56,189,248,0.30);
    border-radius: 99px;
    padding: 4px 14px;
    font-size: 0.78rem;
    font-weight: 700;
    color: var(--accent-cyan);
    letter-spacing: 0.04em;
    font-family: var(--font-sans);
    margin-bottom: 10px;
}

/* ══════════════════════════════════════
   WELCOME BANNER (post-login)
   ══════════════════════════════════════ */
.welcome-banner {
    background: linear-gradient(135deg,
        rgba(14,20,32,0.9) 0%,
        rgba(10,16,26,0.95) 100%);
    border: 1px solid rgba(56,189,248,0.15);
    border-radius: var(--radius-lg);
    padding: 24px 32px;
    margin-bottom: 28px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    animation: fadeSlideUp 0.6s ease forwards;
    position: relative;
    overflow: hidden;
}
.welcome-banner::before {
    content: '';
    position: absolute;
    top: -50%; left: -20%;
    width: 60%; height: 200%;
    background: radial-gradient(ellipse, rgba(56,189,248,0.05) 0%, transparent 70%);
    pointer-events: none;
}
.welcome-title {
    font-family: var(--font-sans) !important;
    font-size: 1.45rem !important;
    font-weight: 700 !important;
    letter-spacing: -0.025em !important;
    color: var(--text-primary) !important;
    line-height: 1.3 !important;
}
.welcome-subtitle {
    font-size: 0.85rem;
    color: var(--text-secondary);
    margin-top: 4px;
    font-family: var(--font-sans);
}
.welcome-username {
    color: var(--accent-cyan);
    font-weight: 700;
}

/* ══════════════════════════════════════
   ADMIN DASHBOARD
   ══════════════════════════════════════ */
.admin-header {
    background: linear-gradient(135deg,
        rgba(129,140,248,0.12) 0%,
        rgba(99,102,241,0.06) 100%);
    border: 1px solid rgba(129,140,248,0.20);
    border-radius: var(--radius-md);
    padding: 16px 24px;
    margin-bottom: 24px;
}
.admin-header h2 {
    color: var(--accent-violet) !important;
    margin: 0 !important;
    font-size: 1.2rem !important;
}

/* ══════════════════════════════════════
   SECTION DIVIDER with label
   ══════════════════════════════════════ */
.section-label {
    font-family: var(--font-sans) !important;
    font-size: 0.72rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.12em !important;
    text-transform: uppercase !important;
    color: var(--text-muted) !important;
    margin-bottom: 12px !important;
    padding-bottom: 8px !important;
    border-bottom: 1px solid var(--border-subtle) !important;
}

/* ══════════════════════════════════════
   SPINNER
   ══════════════════════════════════════ */
.stSpinner > div {
    border-top-color: var(--accent-cyan) !important;
}

/* ══════════════════════════════════════
   TIMER DISPLAY
   ══════════════════════════════════════ */
.timer-display {
    background: linear-gradient(135deg,
        rgba(251,191,36,0.12) 0%,
        rgba(251,191,36,0.05) 100%);
    backdrop-filter: blur(16px);
    -webkit-backdrop-filter: blur(16px);
    border: 1px solid rgba(251,191,36,0.28);
    border-radius: var(--radius-md);
    padding: 16px 24px;
    margin: 18px 0;
    text-align: center;
    box-shadow: 0 4px 20px rgba(251,191,36,0.08), inset 0 1px 0 rgba(255,255,255,0.06);
    transition: all var(--transition-base);
    position: relative;
    overflow: hidden;
}
.timer-display::before {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 50%; height: 100%;
    background: linear-gradient(90deg, transparent, rgba(251,191,36,0.08), transparent);
    animation: shimmerSlide 3s infinite;
}
.timer-display:hover { transform: translateY(-2px); box-shadow: 0 8px 28px rgba(251,191,36,0.12); }
.timer-text {
    color: var(--accent-amber);
    font-size: 1rem;
    font-weight: 700;
    font-family: var(--font-sans);
    letter-spacing: 0.04em;
    position: relative;
    z-index: 2;
}
.timer-expired { 
    background: linear-gradient(135deg, rgba(251,113,133,0.12) 0%, rgba(251,113,133,0.05) 100%);
    border-color: rgba(251,113,133,0.28);
}
.timer-expired .timer-text { color: var(--accent-rose); }

/* ══════════════════════════════════════
   BANNER / MARQUEE (tab1 dashboard)
   ══════════════════════════════════════ */
.banner-container {
    width: 100%;
    height: 60px;
    background: linear-gradient(90deg,
        rgba(8,12,18,1) 0%,
        rgba(14,20,32,0.9) 50%,
        rgba(8,12,18,1) 100%);
    border: 1px solid var(--border-subtle);
    overflow: hidden;
    display: flex;
    align-items: center;
    position: relative;
    margin-bottom: 24px;
    border-radius: var(--radius-md);
    backdrop-filter: blur(20px);
}
.pulse-bar {
    position: absolute;
    display: flex;
    align-items: center;
    font-size: 0.9rem;
    font-weight: 600;
    font-family: var(--font-sans);
    color: var(--accent-cyan);
    white-space: nowrap;
    letter-spacing: 0.04em;
    animation: glideIn 14s linear infinite;
}
.pulse-bar .bar {
    width: 3px;
    height: 18px;
    margin-right: 12px;
    background: var(--accent-cyan);
    border-radius: 2px;
    box-shadow: 0 0 8px var(--accent-cyan);
    animation: pulse 1s ease-in-out infinite;
}
@keyframes glideIn {
    0%   { left: -40%; opacity: 0; }
    8%   { opacity: 1; }
    92%  { opacity: 1; }
    100% { left: 110%; opacity: 0; }
}
@keyframes pulse {
    0%, 100% { height: 14px; background: var(--accent-cyan); }
    50%       { height: 22px; background: var(--accent-violet); }
}

/* ══════════════════════════════════════
   HEADER BOX (dashboard title area)
   ══════════════════════════════════════ */
.header {
    font-size: 1.5rem;
    font-weight: 700;
    text-align: center;
    letter-spacing: -0.02em;
    padding: 20px 28px;
    color: var(--text-primary);
    position: relative;
    overflow: hidden;
    border-radius: var(--radius-md);
    background: linear-gradient(135deg,
        rgba(14,20,32,0.8) 0%,
        rgba(10,16,26,0.9) 100%);
    border: 1px solid rgba(56,189,248,0.18);
    box-shadow: var(--shadow-card);
    font-family: var(--font-sans);
}
.header span { color: var(--accent-cyan); }
.header::before {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 50%; height: 100%;
    background: linear-gradient(90deg, transparent, rgba(56,189,248,0.05), transparent);
    transition: left 0.7s ease;
}
.header:hover::before { left: 150%; }

/* ══════════════════════════════════════
   ANALYSIS RESULT CARDS
   ══════════════════════════════════════ */
.result-card {
    background: var(--surface-01);
    border: 1px solid var(--border-subtle);
    border-radius: var(--radius-md);
    padding: 20px;
    margin-bottom: 14px;
    transition: all var(--transition-base);
    animation: fadeSlideUp 0.5s ease forwards;
    position: relative;
    overflow: hidden;
}
.result-card::before {
    content: '';
    position: absolute;
    left: 0; top: 0; bottom: 0;
    width: 3px;
    background: linear-gradient(180deg, var(--accent-cyan) 0%, var(--accent-violet) 100%);
    border-radius: 0 0 0 var(--radius-md);
}
.result-card:hover {
    border-color: var(--border-accent);
    background: var(--surface-02);
    transform: translateX(3px);
}

/* ══════════════════════════════════════
   DARK STREAMLIT OVERRIDES
   ══════════════════════════════════════ */
.stMarkdown p, .stText {
    color: var(--text-secondary) !important;
    font-family: var(--font-sans) !important;
    font-size: 0.9rem !important;
    line-height: 1.65 !important;
}
.stMarkdown strong { color: var(--text-primary) !important; }
.stMarkdown code {
    background: rgba(255,255,255,0.06) !important;
    color: var(--accent-cyan) !important;
    border-radius: 4px !important;
    padding: 2px 6px !important;
    font-size: 0.83rem !important;
}
.stInfo, .stSuccess, .stWarning, .stError {
    border-radius: var(--radius-md) !important;
    font-family: var(--font-sans) !important;
    font-size: 0.875rem !important;
}

/* Caption */
.stCaption {
    color: var(--text-muted) !important;
    font-size: 0.78rem !important;
    font-family: var(--font-sans) !important;
}

/* ══════════════════════════════════════
   FILE UPLOADER INNER LABEL
   ══════════════════════════════════════ */
.stFileUploader [data-testid="stFileUploaderDropzone"] {
    background: rgba(56,189,248,0.02) !important;
}
.stFileUploader [data-testid="stFileUploaderDropzone"] span {
    color: var(--text-secondary) !important;
    font-family: var(--font-sans) !important;
}
</style>
""", unsafe_allow_html=True)
# 🔹 VIDEO BACKGROUND & GLOW TEXT

# ------------------- BEFORE LOGIN -------------------
if not st.session_state.authenticated:
    

    # -------- Sidebar --------
    with st.sidebar:
        st.markdown("""
        <div style="
            padding: 16px 4px 20px;
            border-bottom: 1px solid rgba(255,255,255,0.07);
            margin-bottom: 16px;
        ">
            <div style="
                font-family: var(--font-sans, -apple-system, sans-serif);
                font-size: 1.1rem;
                font-weight: 700;
                letter-spacing: -0.02em;
                color: #f0f4f8;
                line-height: 1.2;
            ">HIRELYZER</div>
            <div style="
                font-size: 0.72rem;
                font-weight: 600;
                letter-spacing: 0.1em;
                text-transform: uppercase;
                color: #38bdf8;
                margin-top: 3px;
            ">AI Resume Intelligence</div>
        </div>
        <p style="
            color: #64748b;
            font-size: 0.8rem;
            line-height: 1.55;
            font-family: var(--font-sans, -apple-system, sans-serif);
            margin-bottom: 16px;
        ">Transform your career with AI-powered resume analysis, job matching, and smart insights.</p>
        """, unsafe_allow_html=True)

        features = [
            ("https://img.icons8.com/fluency/48/resume.png", "Resume Analyzer", "Get feedback, scores, and tips powered by AI along with the biased words detection and rewriting the resume in an inclusive way."),
            ("https://img.icons8.com/fluency/48/resume-website.png", "Resume Builder", "Build modern, eye-catching resumes easily."),
            ("https://img.icons8.com/fluency/48/job.png", "Job Search", "Find tailored job matches."),
            ("https://img.icons8.com/fluency/48/classroom.png", "Course Suggestions", "Get upskilling recommendations based on your goals."),
            ("https://img.icons8.com/fluency/48/combo-chart.png", "Interactive Dashboard", "Visualize trends, scores, and analytics."),
        ]

        for icon, title, desc in features:
            st.markdown(f"""
            <div class="feature-card">
                <img src="{icon}" width="40"/>
                <h3>{title}</h3>
                <p>{desc}</p>
            </div>
            """, unsafe_allow_html=True)

    # -------- Premium Hero Section --------
    # Fetch live stats for subtle ribbon (cached — no Supabase hit on every rerun)
    total_users, active_logins, stats = _cached_hero_stats()
    resumes_uploaded = stats.get("total_candidates", 0)
    active_domains = stats.get("unique_domains", 0)

    # ── Hero HTML (no script — Streamlit strips <script> from st.markdown) ──
    st.markdown(f"""
    <style>
    .hero-section {{
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 36px 24px 14px;
        position: relative;
        overflow: hidden;
    }}
    .hero-section::before {{
        content: '';
        position: absolute;
        top: -60px; left: 50%;
        transform: translateX(-50%);
        width: 520px; height: 320px;
        background: radial-gradient(ellipse, rgba(56,189,248,0.08) 0%, transparent 70%);
        pointer-events: none;
        z-index: 0;
    }}
    .hero-brand {{
        position: relative;
        z-index: 2;
        text-align: center;
        margin-bottom: 18px;
        animation: fadeSlideUp 0.7s cubic-bezier(0.22,1,0.36,1) both;
    }}
    .hero-wordmark {{
        font-size: 3rem;
        font-weight: 800;
        letter-spacing: -0.04em;
        line-height: 1;
        color: #f0f4f8;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", "DM Sans", sans-serif;
    }}
    .hero-wordmark span {{
        background: linear-gradient(135deg, #38bdf8 0%, #818cf8 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }}
    .hero-tagline {{
        margin-top: 8px;
        font-size: 0.8rem;
        color: #334155;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
        letter-spacing: 0.04em;
        text-transform: uppercase;
        font-weight: 500;
    }}
    .hero-pills-container {{
        display: inline-flex;
        gap: 6px;
        flex-wrap: wrap;
        justify-content: center;
        align-items: center;
        background: rgba(255,255,255,0.03);
        border: 1px solid rgba(255,255,255,0.08);
        border-radius: 99px;
        padding: 6px 10px;
        margin-bottom: 32px;
        position: relative;
        z-index: 2;
        backdrop-filter: blur(16px);
        -webkit-backdrop-filter: blur(16px);
        animation: fadeSlideUp 1s cubic-bezier(0.22,1,0.36,1) 0.25s both;
    }}
    .hero-pill {{
        display: inline-flex;
        align-items: center;
        gap: 5px;
        padding: 4px 11px;
        border-radius: 99px;
        background: transparent;
        font-size: 0.75rem;
        font-weight: 500;
        color: #64748b;
        font-family: -apple-system, BlinkMacSystemFont, sans-serif;
        letter-spacing: 0.01em;
        transition: color 0.2s ease;
    }}
    .hero-pill:hover {{
        color: #38bdf8;
    }}
    .hero-pill-dot {{
        width: 5px; height: 5px;
        border-radius: 50%;
        background: #38bdf8;
        opacity: 0.55;
        flex-shrink: 0;
    }}
    .hero-pill-sep {{
        width: 1px; height: 12px;
        background: rgba(255,255,255,0.10);
        flex-shrink: 0;
        align-self: center;
    }}
    .hero-stat-ribbon {{
        display: inline-flex;
        gap: 0;
        justify-content: center;
        align-items: stretch;
        position: relative;
        z-index: 2;
        margin-bottom: 8px;
        background: rgba(255,255,255,0.025);
        border: 1px solid rgba(255,255,255,0.07);
        border-radius: 16px;
        padding: 0;
        overflow: hidden;
        animation: fadeSlideUp 1.1s cubic-bezier(0.22,1,0.36,1) 0.35s both;
    }}
    .hero-stat-item {{
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 16px 28px;
        gap: 2px;
        transition: background 0.2s ease;
    }}
    .hero-stat-item:hover {{
        background: rgba(56,189,248,0.05);
    }}
    .hero-stat-item:not(:last-child) {{
        border-right: 1px solid rgba(255,255,255,0.07);
    }}
    .hero-stat-num {{
        font-size: 1.4rem;
        font-weight: 700;
        color: #f0f4f8;
        letter-spacing: -0.03em;
        line-height: 1;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
    }}
    .hero-stat-icon {{
        display: flex;
        align-items: center;
        justify-content: center;
        width: 28px; height: 28px;
        border-radius: 8px;
        background: rgba(56,189,248,0.08);
        border: 1px solid rgba(56,189,248,0.15);
        margin-bottom: 4px;
        flex-shrink: 0;
    }}
    .hero-stat-lbl {{
        font-size: 0.65rem;
        color: #4a5568;
        text-transform: uppercase;
        letter-spacing: 0.09em;
        font-weight: 600;
        font-family: -apple-system, BlinkMacSystemFont, sans-serif;
    }}
    </style>

    <div class="hero-section">
        <div class="hero-brand">
            <div class="hero-wordmark">HIRE<span>LYZER</span></div>
            <div class="hero-tagline">AI-Powered Resume Intelligence Platform</div>
        </div>
        <div class="hero-pills-container">
            <div class="hero-pill"><span class="hero-pill-dot"></span>Bias Detection</div>
            <div class="hero-pill-sep"></div>
            <div class="hero-pill"><span class="hero-pill-dot"></span>ATS Scoring</div>
            <div class="hero-pill-sep"></div>
            <div class="hero-pill"><span class="hero-pill-dot"></span>Resume Builder</div>
            <div class="hero-pill-sep"></div>
            <div class="hero-pill"><span class="hero-pill-dot"></span>Job Matching</div>
            <div class="hero-pill-sep"></div>
            <div class="hero-pill"><span class="hero-pill-dot"></span>AI Coach</div>
        </div>
        <div class="hero-stat-ribbon">
            <div class="hero-stat-item">
                <div class="hero-stat-icon">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <circle cx="12" cy="7" r="4" stroke="#38bdf8" stroke-width="1.8" fill="rgba(56,189,248,0.12)"/>
                        <path d="M4 20c0-4 3.582-7 8-7s8 3 8 7" stroke="#38bdf8" stroke-width="1.8" stroke-linecap="round"/>
                    </svg>
                </div>
                <div class="hero-stat-num">{total_users}</div>
                <div class="hero-stat-lbl">Users</div>
            </div>
            <div class="hero-stat-item">
                <div class="hero-stat-icon">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <rect x="4" y="2" width="16" height="20" rx="2" stroke="#38bdf8" stroke-width="1.8" fill="rgba(56,189,248,0.08)"/>
                        <path d="M8 7h8M8 11h8M8 15h5" stroke="#38bdf8" stroke-width="1.5" stroke-linecap="round"/>
                    </svg>
                </div>
                <div class="hero-stat-num">{resumes_uploaded}</div>
                <div class="hero-stat-lbl">Resumes Analysed</div>
            </div>
            <div class="hero-stat-item">
                <div class="hero-stat-icon">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <circle cx="12" cy="12" r="9" stroke="#38bdf8" stroke-width="1.8" fill="rgba(56,189,248,0.08)"/>
                        <path d="M12 3v18M3 12h18" stroke="#38bdf8" stroke-width="1.5" stroke-linecap="round"/>
                    </svg>
                </div>
                <div class="hero-stat-num">{active_domains}</div>
                <div class="hero-stat-lbl">Domains</div>
            </div>
            <div class="hero-stat-item">
                <div class="hero-stat-icon">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <path d="M13 2L4.5 13.5H12L11 22L19.5 10.5H12L13 2Z" stroke="#38bdf8" stroke-width="1.8" stroke-linejoin="round" fill="rgba(56,189,248,0.12)"/>
                    </svg>
                </div>
                <div class="hero-stat-num">{active_logins}</div>
                <div class="hero-stat-lbl">Active Today</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Typewriter animation via components.html (scripts work here) ──
    st.components.v1.html("""
    <style>
    * { box-sizing: border-box; }
    body { margin: 0; padding: 0; background: transparent; overflow: hidden; }
    .tw-wrap {
        text-align: center;
        padding: 4px 0 8px;
        width: 100%;
    }
    .tw-text {
        font-size: 1rem;
        font-weight: 600;
        color: #38bdf8;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
        letter-spacing: -0.01em;
        border-right: 2px solid #38bdf8;
        white-space: nowrap;
        display: inline-block;
        max-width: 100%;
        animation: twBlink 1s ease-in-out infinite;
        background: transparent;
        vertical-align: middle;
    }
    @keyframes twBlink {
        0%, 100% { border-color: #38bdf8; }
        50%       { border-color: transparent; }
    }
    </style>
    <div class="tw-wrap"><span class="tw-text" id="tw">&nbsp;</span></div>
    <script>
    (function() {
        var phrases = [
            "Analyse resumes with zero bias.",
            "Score smarter. Hire better.",
            "AI that reads between the lines.",
            "Ethical hiring starts here.",
            "10x faster resume screening."
        ];
        var idx = 0, charIdx = 0, deleting = false;
        var el = document.getElementById('tw');
        function tick() {
            var phrase = phrases[idx];
            if (!deleting) {
                el.textContent = phrase.slice(0, ++charIdx);
                if (charIdx === phrase.length) { deleting = true; setTimeout(tick, 1800); return; }
            } else {
                el.textContent = phrase.slice(0, --charIdx);
                if (charIdx === 0) { deleting = false; idx = (idx + 1) % phrases.length; }
            }
            setTimeout(tick, deleting ? 38 : 62);
        }
        setTimeout(tick, 600);
    })();
    </script>
    """, height=44, scrolling=False)

if not st.session_state.get("authenticated", False):

    # -------- Login/Register Layout --------
    left, center, right = st.columns([1, 2, 1])

    with center:
        st.markdown(
            """<div class='login-card'>
            <div style='text-align:center; margin-bottom:12px;'>
                <div style='display:inline-flex; align-items:center; justify-content:center; width:40px; height:40px; border-radius:10px; background:rgba(79,140,255,0.12); border:1px solid rgba(79,140,255,0.22); margin-bottom:12px;'>
                    <svg width="18" height="18" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <rect x="5" y="11" width="14" height="10" rx="2" stroke="#4f8cff" stroke-width="1.5" fill="rgba(79,140,255,0.12)"/>
                        <path d="M8 11V7a4 4 0 0 1 8 0v4" stroke="#4f8cff" stroke-width="1.5" stroke-linecap="round"/>
                        <circle cx="12" cy="16" r="1.2" fill="#4f8cff"/>
                    </svg>
                </div>
            </div>
            <h2 id='auth-heading' style='text-align:center; font-family:-apple-system,BlinkMacSystemFont,"SF Pro Display","Segoe UI",Roboto,sans-serif; font-size:1.3rem; font-weight:700; letter-spacing:-0.025em; color:#e6edf3; margin:0 0 18px 0;'>Sign in to <span style='color:#4f8cff;'>HIRELYZER</span></h2>
            <script>
            (function() {
                function updateHeading() {
                    var tabs = window.parent.document.querySelectorAll('[data-baseweb="tab"]');
                    var heading = window.parent.document.getElementById('auth-heading');
                    if (!heading || tabs.length < 2) return;
                    var activeTab = window.parent.document.querySelector('[data-baseweb="tab"][aria-selected="true"]');
                    if (activeTab) {
                        var label = activeTab.textContent.trim().toLowerCase();
                        if (label === 'register') {
                            heading.innerHTML = 'Register to <span style="color:#4f8cff;">HIRELYZER</span>';
                        } else {
                            heading.innerHTML = 'Sign in to <span style="color:#4f8cff;">HIRELYZER</span>';
                        }
                    }
                    tabs.forEach(function(tab) {
                        tab.addEventListener('click', function() {
                            setTimeout(updateHeading, 80);
                        });
                    });
                }
                setTimeout(updateHeading, 400);
                var observer = new MutationObserver(function() { updateHeading(); });
                setTimeout(function() {
                    var tabBar = window.parent.document.querySelector('[data-baseweb="tab-list"]');
                    if (tabBar) observer.observe(tabBar, { attributes: true, subtree: true, attributeFilter: ['aria-selected'] });
                }, 600);
            })();
            </script>""",
            unsafe_allow_html=True,
        )

        login_tab, register_tab = st.tabs(["Login", "Register"])

        # ---------------- LOGIN TAB ----------------
        with login_tab:
            # Show login or forgot password flow based on reset_stage
            if st.session_state.reset_stage == "none":

                # ── Stage A: Credentials entry ──────────────────────────────
                if st.session_state.login_stage == "credentials":
                    st.markdown("""<h3 style='color:#9aa4af; text-align:center; font-family:-apple-system,BlinkMacSystemFont,"SF Pro Display","Segoe UI",Roboto,sans-serif; font-size:0.82rem; font-weight:500; letter-spacing:0.06em; text-transform:uppercase; margin-bottom:24px;'>Welcome Back</h3>""", unsafe_allow_html=True)

                    user = st.text_input("Username or Email", key="login_user")
                    pwd = st.text_input("Password", type="password", key="login_pass")

                    # Render notification area (reserves space)
                    render_notification("login")

                    if st.button("Sign In", key="login_btn", use_container_width=True):
                        success, _ = verify_user(user.strip(), pwd.strip())
                        if success:
                            # Credentials OK — look up email and send confirmation link
                            _uname = st.session_state.username   # set by verify_user()
                            _email = get_email_by_username(_uname)
                            if _email:
                                _token = create_login_token(_uname)
                                sent = send_login_confirmation_email(_email, _uname, _token)
                                if sent:
                                    # Move to waiting stage; clear the partial auth set by verify_user
                                    st.session_state.authenticated = False
                                    st.session_state.pending_login_username = _uname
                                    st.session_state.login_stage = "awaiting_email"
                                    st.rerun()
                                else:
                                    st.session_state.authenticated = False
                                    notify("login", "error", "❌ Could not send confirmation email. Please try again.")
                                    st.rerun()
                            else:
                                # Account has no email on record — cannot send link
                                st.session_state.authenticated = False
                                notify("login", "error", "❌ No email address is associated with this account. Please contact support.")
                                st.rerun()
                        else:
                            st.session_state.authenticated = False
                            notify("login", "error", "❌ Invalid credentials. Please try again.")
                            st.rerun()

                    st.markdown("<br>", unsafe_allow_html=True)

                    # Forgot Password Link
                    if st.button("Forgot Password?", key="forgot_pw_link"):
                        st.session_state.reset_stage = "request_email"
                        st.rerun()

                # ── Stage B: Waiting for email confirmation ─────────────────
                elif st.session_state.login_stage == "awaiting_email":
                    _pending = st.session_state.pending_login_username or "you"
                    _masked_email = ""
                    _raw_email = get_email_by_username(_pending) if _pending != "you" else None
                    if _raw_email:
                        _parts = _raw_email.split("@")
                        _masked_email = _parts[0][:2] + "***@" + _parts[1]

                    st.markdown(f"""
                    <div style="
                        text-align:center;
                        padding: 28px 16px 20px;
                        font-family:-apple-system,BlinkMacSystemFont,'SF Pro Display','Segoe UI',Roboto,sans-serif;
                    ">
                        <div style="font-size:2.6rem; margin-bottom:12px;">📧</div>
                        <div style="font-size:1.05rem; font-weight:600; color:#e6edf3; margin-bottom:8px;">
                            Check your email
                        </div>
                        <div style="font-size:0.85rem; color:#8b9ab0; line-height:1.6;">
                            We sent a confirmation link to<br>
                            <strong style="color:#c9d1d9;">{_masked_email}</strong><br><br>
                            Click <em>"Yes, it's me"</em> in the email to complete your sign-in.<br>
                            The link expires in <strong style="color:#fbbf24;">10 minutes</strong>.
                        </div>
                    </div>
                    """, unsafe_allow_html=True)

                    render_notification("login")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("🔄 Resend link", key="resend_confirm_btn", use_container_width=True):
                            _uname = st.session_state.pending_login_username
                            _email = get_email_by_username(_uname)
                            if _email:
                                _token = create_login_token(_uname)
                                sent = send_login_confirmation_email(_email, _uname, _token)
                                if sent:
                                    notify("login", "success", "📨 New confirmation link sent!")
                                else:
                                    notify("login", "error", "❌ Failed to resend. Please try again.")
                            st.rerun()
                    with col2:
                        if st.button("↩️ Back", key="back_from_confirm_btn", use_container_width=True):
                            st.session_state.login_stage = "credentials"
                            st.session_state.pending_login_username = None
                            st.session_state.authenticated = False
                            st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 1: Request Email
            # ============================================================
            elif st.session_state.reset_stage == "request_email":
                st.markdown("""<h3 style='color:#9aa4af; text-align:center; font-family:-apple-system,BlinkMacSystemFont,"SF Pro Display","Segoe UI",Roboto,sans-serif; font-size:0.82rem; font-weight:500; letter-spacing:0.06em; text-transform:uppercase; margin-bottom:16px;'>Reset Password</h3>""", unsafe_allow_html=True)
                st.markdown("<p style='color:#c9d1d9; text-align:center;'>Enter your registered email to receive an OTP</p>", unsafe_allow_html=True)

                email_input = st.text_input("Email Address", key="reset_email_input")

                # Render notification area (reserves space)
                render_notification("login")

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("📤 Send OTP", key="send_otp_btn", use_container_width=True):
                        if email_input.strip():
                            if get_user_by_email(email_input.strip()):
                                # Generate and send OTP
                                otp = generate_otp()
                                success = send_email_otp(email_input.strip(), otp)

                                if success:
                                    st.session_state.reset_email = email_input.strip()
                                    st.session_state.reset_otp = otp
                                    st.session_state.reset_otp_time = time.time()
                                    st.session_state.reset_stage = "verify_otp"

                                    notify("login", "success", "✅ OTP sent successfully to your email!")
                                    time.sleep(0.5)
                                    st.rerun()
                                else:
                                    notify("login", "error", "❌ Failed to send OTP. Please try again.")
                                    st.rerun()
                            else:
                                notify("login", "error", "❌ Email not found. Please register first.")
                                st.rerun()
                        else:
                            notify("login", "warning", "⚠️ Please enter your email address.")
                            st.rerun()

                with col2:
                    if st.button("↩️ Back to Login", key="back_to_login_1", use_container_width=True):
                        st.session_state.reset_stage = "none"
                        st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 2: Verify OTP
            # ============================================================
            elif st.session_state.reset_stage == "verify_otp":
                st.markdown("""<h3 style='color:#e6edf3; text-align:center; font-family:-apple-system,sans-serif; font-size:1.05rem; font-weight:600;'>
                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" style="vertical-align:-3px; margin-right:6px;" xmlns="http://www.w3.org/2000/svg"><path d="M4 4h16v16H4z" rx="2" stroke="#38bdf8" stroke-width="1.5" fill="none"/><path d="M4 9h16" stroke="#38bdf8" stroke-width="1.5"/><path d="M8 4v5" stroke="#38bdf8" stroke-width="1.5" stroke-linecap="round"/><path d="M16 4v5" stroke="#38bdf8" stroke-width="1.5" stroke-linecap="round"/></svg>
                    Verify OTP</h3>""", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#c9d1d9; text-align:center;'>Enter the 6-digit OTP sent to <strong>{st.session_state.reset_email}</strong></p>", unsafe_allow_html=True)

                # Calculate elapsed and remaining time (server-side)
                elapsed_time = time.time() - st.session_state.reset_otp_time
                remaining_time = max(0, int(180 - elapsed_time))

                # Display timer
                display_timer(remaining_time, expired=(remaining_time == 0), key_suffix="forgot_pw")

                # Check if OTP expired (3 minutes)
                if remaining_time == 0:
                    # OTP Expired - Show resend option
                    render_notification("login")
                    notify("login", "error", "⏱️ OTP expired. Please request a new one.")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("🔄 Resend OTP", key="resend_otp_btn", use_container_width=True):
                            # Generate new OTP
                            otp = generate_otp()
                            success = send_email_otp(st.session_state.reset_email, otp)

                            if success:
                                st.session_state.reset_otp = otp
                                st.session_state.reset_otp_time = time.time()
                                notify("login", "info", "📨 New OTP sent!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("login", "error", "❌ Failed to send OTP. Please try again.")
                                st.rerun()

                    with col2:
                        if st.button("↩️ Back to Login", key="back_to_login_expired", use_container_width=True):
                            st.session_state.reset_stage = "none"
                            st.rerun()
                else:
                    # OTP still valid - Show verification form
                    otp_input = st.text_input("🔢 Enter 6-Digit OTP", key="otp_input", max_chars=6)

                    # Render notification area (reserves space)
                    render_notification("login")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("✅ Verify OTP", key="verify_otp_btn", use_container_width=True):
                            # Re-check expiry on server side before verifying
                            current_elapsed = time.time() - st.session_state.reset_otp_time
                            if current_elapsed >= 180:
                                notify("login", "error", "⏱️ OTP has expired. Please request a new one.")
                                st.rerun()
                            elif otp_input.strip() == st.session_state.reset_otp:
                                st.session_state.reset_stage = "reset_password"
                                notify("login", "success", "✅ OTP verified successfully!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("login", "error", "❌ Invalid OTP. Please try again.")
                                st.rerun()

                    with col2:
                        if st.button("↩️ Back to Login", key="back_to_login_2", use_container_width=True):
                            st.session_state.reset_stage = "none"
                            st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 3: Reset Password
            # ============================================================
            elif st.session_state.reset_stage == "reset_password":
                st.markdown("""<h3 style='color:#e6edf3; text-align:center; font-family:-apple-system,sans-serif; font-size:1.05rem; font-weight:600;'>
                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" style="vertical-align:-3px; margin-right:6px;" xmlns="http://www.w3.org/2000/svg"><path d="M12 3a4 4 0 0 1 4 4v1H8V7a4 4 0 0 1 4-4z" stroke="#38bdf8" stroke-width="1.5" fill="none"/><rect x="5" y="11" width="14" height="10" rx="2" stroke="#38bdf8" stroke-width="1.5" fill="rgba(56,189,248,0.08)"/><circle cx="12" cy="16" r="1.5" fill="#38bdf8"/></svg>
                    Set New Password</h3>""", unsafe_allow_html=True)
                st.markdown("<p style='color:#c9d1d9; text-align:center;'>Enter your new password</p>", unsafe_allow_html=True)

                new_password = st.text_input("New Password", type="password", key="new_password_input")
                confirm_password = st.text_input("Confirm Password", type="password", key="confirm_password_input")

                st.caption("Password must be at least 8 characters, include uppercase, lowercase, number, and special character.")

                # Render notification area (reserves space)
                render_notification("login")

                if st.button("✅ Reset Password", key="reset_password_btn", use_container_width=True):
                    if new_password.strip() and confirm_password.strip():
                        if new_password == confirm_password:
                            success = update_password_by_email(st.session_state.reset_email, new_password)

                            if success:
                                notify("login", "success", "✅ Password reset successful! Please log in again.")

                                # Log the password reset action
                                log_user_action(st.session_state.reset_email, "password_reset")

                                # Reset all forgot password session states
                                st.session_state.reset_stage = "none"
                                st.session_state.reset_email = ""
                                st.session_state.reset_otp = ""
                                st.session_state.reset_otp_time = 0

                                time.sleep(1)
                                st.rerun()
                            else:
                                notify("login", "error", "❌ Failed to reset password. Please try again.")
                                st.rerun()
                        else:
                            notify("login", "error", "❌ Passwords do not match.")
                            st.rerun()
                    else:
                        notify("login", "warning", "⚠️ Please fill in both password fields.")
                        st.rerun()

                if st.button("↩️ Back to Login", key="back_to_login_3"):
                    st.session_state.reset_stage = "none"
                    st.rerun()

        # ---------------- REGISTER TAB ----------------
        with register_tab:
            # Check if OTP was sent and pending verification
            if 'pending_registration' in st.session_state:
                st.markdown("""<h3 style='color:#e6edf3; text-align:center; font-family:-apple-system,sans-serif; font-size:1.05rem; font-weight:600;'>
                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" style="vertical-align:-3px; margin-right:6px;" xmlns="http://www.w3.org/2000/svg"><path d="M20 4H4a2 2 0 0 0-2 2v12a2 2 0 0 0 2 2h16a2 2 0 0 0 2-2V6a2 2 0 0 0-2-2z" stroke="#38bdf8" stroke-width="1.5" fill="none"/><path d="M2 8l10 7 10-7" stroke="#38bdf8" stroke-width="1.5"/></svg>
                    Verify Your Email</h3>""", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#c9d1d9; text-align:center;'>Enter the 6-digit OTP sent to <strong>{st.session_state.pending_registration['email']}</strong></p>", unsafe_allow_html=True)

                # Calculate remaining time
                from datetime import datetime
                elapsed = (datetime.now(st.session_state.pending_registration['timestamp'].tzinfo) - st.session_state.pending_registration['timestamp']).total_seconds()
                remaining = max(0, 180 - int(elapsed))

                # Display timer
                display_timer(remaining, expired=(remaining == 0), key_suffix="register")

                if remaining == 0:
                    # OTP Expired
                    render_notification("register")
                    notify("register", "error", "⏱️ OTP expired. Please request a new one.")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("🔄 Resend OTP", key="reg_resend_expired_btn", use_container_width=True):
                            pending = st.session_state.pending_registration
                            success, message = add_user(pending['username'], pending['password'], pending['email'])
                            if success:
                                notify("register", "success", "✅ New OTP sent!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("register", "error", f"❌ {message}")
                                st.rerun()
                    with col2:
                        if st.button("↩️ Start Over", key="reg_start_over_btn", use_container_width=True):
                            del st.session_state.pending_registration
                            st.rerun()
                else:
                    # OTP still valid
                    otp_input = st.text_input("🔢 Enter 6-Digit OTP", key="reg_otp_input", max_chars=6)

                    # Render notification area (reserves space)
                    render_notification("register")

                    col1, col2, col3 = st.columns(3)
                    with col1:
                        if st.button("✅ Verify", key="verify_reg_otp_btn", use_container_width=True):
                            # Cache username BEFORE calling complete_registration
                            cached_username = st.session_state.pending_registration['username']

                            # Re-check expiry before verification
                            current_elapsed = (datetime.now(st.session_state.pending_registration['timestamp'].tzinfo) - st.session_state.pending_registration['timestamp']).total_seconds()
                            if current_elapsed >= 180:
                                notify("register", "error", "⏱️ OTP has expired. Please request a new one.")
                                st.rerun()
                            else:
                                success, message = complete_registration(otp_input.strip())
                                if success:
                                    notify("register", "success", message)
                                    log_user_action(cached_username, "register")
                                    time.sleep(0.5)
                                    st.rerun()
                                else:
                                    notify("register", "error", message)
                                    st.rerun()

                    with col2:
                        if st.button("🔄 Resend", key="resend_reg_otp_btn", use_container_width=True):
                            pending = st.session_state.pending_registration
                            success, message = add_user(pending['username'], pending['password'], pending['email'])
                            if success:
                                notify("register", "info", "📨 New OTP sent successfully!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("register", "error", f"❌ {message}")
                                st.rerun()

                    with col3:
                        if st.button("↩️ Back", key="back_to_reg_btn", use_container_width=True):
                            del st.session_state.pending_registration
                            st.rerun()

            else:
                # Normal registration form
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>🧾 Register New User</h3>", unsafe_allow_html=True)

                # ── CSS: fixed-height validation slot — zero layout shift ──
                st.markdown("""
                <style>
                /* Outer wrapper: zero document-flow height — no spacing contribution */
                .val-slot {
                    height: 0;
                    overflow: visible;
                    position: relative;
                    margin: 0;
                    padding: 0;
                    line-height: 0;
                }
                /* The badge floats above the next field via negative top offset */
                .val-badge {
                    position: absolute;
                    top: -26px;
                    left: 0;
                    right: 0;
                    display: flex;
                    align-items: center;
                    gap: 6px;
                    padding: 3px 10px;
                    border-radius: 5px;
                    font-size: 0.75rem;
                    font-weight: 500;
                    font-family: var(--font-sans), -apple-system, sans-serif;
                    line-height: 1.3;
                    opacity: 0;
                    transform: translateY(2px);
                    transition: opacity 0.18s ease, transform 0.18s ease;
                    pointer-events: none;
                    white-space: nowrap;
                    overflow: hidden;
                    text-overflow: ellipsis;
                    z-index: 10;
                }
                .val-badge.val-visible {
                    opacity: 1;
                    transform: translateY(0);
                    pointer-events: auto;
                }
                .val-badge.val-success {
                    background: rgba(52,211,153,0.12);
                    border: 1px solid rgba(52,211,153,0.25);
                    color: #6ee7b7;
                }
                .val-badge.val-error {
                    background: rgba(251,113,133,0.12);
                    border: 1px solid rgba(251,113,133,0.25);
                    color: #fca5a5;
                }
                .val-badge.val-warn {
                    background: rgba(251,191,36,0.12);
                    border: 1px solid rgba(251,191,36,0.25);
                    color: #fde68a;
                }
                @keyframes _val_autofade {
                    0%   { opacity: 1; }
                    65%  { opacity: 1; }
                    100% { opacity: 0; }
                }
                .val-badge.val-autofade {
                    animation: _val_autofade 3.5s ease forwards;
                }
                </style>
                """, unsafe_allow_html=True)

                # ── on_change callbacks — DB hits ONLY when field value changes ──
                def _validate_email():
                    val = st.session_state.get("reg_email", "").strip()
                    if not val:
                        st.session_state._email_msg = ("", "")
                        return
                    if not is_valid_email(val):
                        st.session_state._email_msg = ("warn", "⚠️ Invalid email format.")
                    elif email_exists(val):
                        st.session_state._email_msg = ("error", "❌ Email already registered.")
                    else:
                        st.session_state._email_msg = ("success", "✅ Email is available.")

                def _validate_username():
                    val = st.session_state.get("reg_user", "").strip()
                    if not val:
                        st.session_state._user_msg = ("", "")
                        return
                    if username_exists(val):
                        st.session_state._user_msg = ("error", "❌ Username already exists.")
                    else:
                        st.session_state._user_msg = ("success", "✅ Username is available.")

                def _validate_password():
                    val = st.session_state.get("reg_pass", "")
                    if not val:
                        st.session_state._pass_msg = ("", "")
                        return
                    if not is_strong_password(val):
                        st.session_state._pass_msg = ("warn", "⚠️ Password must be at least 8 characters and strong.")
                    else:
                        st.session_state._pass_msg = ("success", "✅ Strong password.")

                # Initialise message state once
                if "_email_msg" not in st.session_state:
                    st.session_state._email_msg = ("", "")
                if "_user_msg" not in st.session_state:
                    st.session_state._user_msg = ("", "")
                if "_pass_msg" not in st.session_state:
                    st.session_state._pass_msg = ("", "")

                def _render_val_msg(state_key):
                    """Render a compact validation badge in a fixed-height slot — zero layout shift."""
                    kind, text = st.session_state.get(state_key, ("", ""))
                    if not kind or not text:
                        # Always emit the fixed-height wrapper, badge stays invisible
                        st.markdown(
                            '<div class="val-slot"><div class="val-badge"></div></div>',
                            unsafe_allow_html=True
                        )
                        return
                    type_class = {"warn": "val-warn", "error": "val-error", "success": "val-success"}.get(kind, "val-warn")
                    # Success messages auto-fade; errors/warnings stay visible
                    fade_class = " val-autofade" if kind == "success" else ""
                    st.markdown(
                        f'<div class="val-slot"><div class="val-badge {type_class} val-visible{fade_class}">{text}</div></div>',
                        unsafe_allow_html=True
                    )

                # ── Inputs wired to on_change — NO inline DB calls ──
                new_email = st.text_input(
                    "📧 Email", key="reg_email",
                    placeholder="your@email.com",
                    on_change=_validate_email
                )
                _render_val_msg("_email_msg")

                new_user = st.text_input(
                    "👤 Username", key="reg_user",
                    on_change=_validate_username
                )
                _render_val_msg("_user_msg")

                new_pass = st.text_input(
                    "🔑 Password", type="password", key="reg_pass",
                    on_change=_validate_password
                )
                st.caption("Password must be at least 8 characters, include uppercase, lowercase, number, and special character.")
                _render_val_msg("_pass_msg")

                # Render notification area (reserves space)
                render_notification("register")

                if st.button("📧 Register & Send OTP", key="register_btn", use_container_width=True):
                    if new_email.strip() and new_user.strip() and new_pass.strip():
                        # Validate before attempting registration
                        if not is_valid_email(new_email.strip()):
                            notify("register", "warning", "⚠️ Invalid email format.")
                            st.rerun()
                        elif email_exists(new_email.strip()):
                            notify("register", "error", "🚫 Email already registered.")
                            st.rerun()
                        elif username_exists(new_user.strip()):
                            notify("register", "error", "🚫 Username already exists.")
                            st.rerun()
                        else:
                            success, message = add_user(new_user.strip(), new_pass.strip(), new_email.strip())
                            if success:
                                notify("register", "success", message)
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("register", "error", message)
                                st.rerun()
                    else:
                        notify("register", "warning", "⚠️ Please fill in all fields (email, username, and password).")
                        st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)

    st.stop()

# ------------------- AFTER LOGIN -------------------
if st.session_state.get("authenticated"):
    st.markdown(f"""
    <div class="welcome-banner">
        <div>
            <div class="welcome-title">Welcome back, <span class="welcome-username">{st.session_state.username}</span> 👋</div>
            <div class="welcome-subtitle">HIRELYZER — AI-Powered Resume Intelligence Platform</div>
        </div>
        <div style="display:flex; align-items:center; gap:8px;">
            <div style="
                background: linear-gradient(135deg, rgba(52,211,153,0.15) 0%, rgba(52,211,153,0.06) 100%);
                border: 1px solid rgba(52,211,153,0.25);
                border-radius: 99px;
                padding: 5px 14px;
                font-size: 0.75rem;
                font-weight: 600;
                color: #6ee7b7;
                letter-spacing: 0.04em;
                text-transform: uppercase;
                font-family: var(--font-sans);
            ">● Live</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # 🔓 LOGOUT BUTTON
    if st.button("🚪 Logout"):
        log_user_action(st.session_state.get("username", "unknown"), "logout")

        # ✅ Clear all session keys safely
        for key in list(st.session_state.keys()):
            del st.session_state[key]

        st.success("✅ Logged out successfully.")
        st.rerun()  # Force rerun to prevent stale UI

    # 🔑 GROQ API KEY SECTION (SIDEBAR)
    st.sidebar.markdown("""
    <p style='
        font-size: 0.72rem;
        font-weight: 700;
        letter-spacing: 0.10em;
        text-transform: uppercase;
        color: #4a5568;
        border-bottom: 1px solid rgba(255,255,255,0.06);
        padding-bottom: 8px;
        margin-bottom: 12px;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
    '>🔑 Groq API Key</p>
    """, unsafe_allow_html=True)

    # ✅ Load saved key from DB (cached — won't re-query on every rerun)
    saved_key = _cached_user_api_key(st.session_state.username)
    masked_preview = f"****{saved_key[-6:]}" if saved_key else ""

    user_api_key_input = st.sidebar.text_input(
        "Your Groq API Key (Optional)",
        placeholder=masked_preview,
        type="password"
    )

    # ✅ Save or reuse key — guarded so save_user_api_key only fires once per
    #    new value, not on every rerun while the field holds a value.
    if user_api_key_input:
        if user_api_key_input != st.session_state.get("_last_saved_api_key"):
            save_user_api_key(st.session_state.username, user_api_key_input)
            st.session_state["_last_saved_api_key"] = user_api_key_input
            _cached_user_api_key.clear()  # bust cache so next read gets new value
            st.sidebar.success("✅ New key saved and in use.")
        st.session_state["user_groq_key"] = user_api_key_input
    elif saved_key:
        st.session_state["user_groq_key"] = saved_key
        st.sidebar.info(f"ℹ️ Using your previously saved API key ({masked_preview})")
    else:
        st.sidebar.warning("⚠ Using shared admin key with possible usage limits")

    # 🧹 Clear saved key
    if st.sidebar.button("🗑️ Clear My API Key"):
        st.session_state["user_groq_key"] = None
        st.session_state.pop("_last_saved_api_key", None)
        save_user_api_key(st.session_state.username, None)
        _cached_user_api_key.clear()
        st.sidebar.success("✅ Cleared saved Groq API key. Now using shared admin key.")

if st.session_state.username == "admin":
    st.markdown("""
    <div class="admin-header">
        <h2>⬡ Admin Control Panel</h2>
    </div>
    """, unsafe_allow_html=True)

    # Metrics row — cached, no Supabase hit on every rerun
    _reg_users, _logins_today, _logs = _cached_admin_metrics()
    col1, col2 = st.columns(2)
    with col1:
        st.metric("👤 Total Registered Users", _reg_users)
    with col2:
        st.metric("📅 Logins Today (IST)", _logins_today)

    # Removed API key usage section (no longer tracked)
    # Activity log
    st.markdown("<p class='section-label'>📋 Activity Log</p>", unsafe_allow_html=True)
    logs = _logs
    if logs:
        st.dataframe(
            {
                "Username": [log[0] for log in logs],
                "Action": [log[1] for log in logs],
                "Timestamp": [log[2] for log in logs]
            },
            use_container_width=True
        )
    else:
        st.info("No logs found yet.")

    st.divider()
    st.info("ℹ️ Data is stored in Supabase PostgreSQL. Use the Admin DB View tab to export records as CSV.")
# Always-visible tabs
tab_labels = [
    "📊 Dashboard",
    "🧾 Resume Builder",
    "💼 Job Search",
    "📚 Course Recommendation"
]

# Add Admin tab only for admin user
if st.session_state.username == "admin":
    tab_labels.append("📁 Admin DB View")

# Create tabs dynamically
tabs = st.tabs(tab_labels)

# Unpack first four (always exist)
tab1, tab2, tab3, tab4 = tabs[:4]

# Handle optional admin tab
tab5 = tabs[4] if len(tabs) > 4 else None
with tab1:
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Orbitron', sans-serif;
        background-color: #0b0c10;
        color: #c5c6c7;
        scroll-behavior: smooth;
    }

    /* ---------- SCROLLBAR ---------- */
    ::-webkit-scrollbar { width: 8px; }
    ::-webkit-scrollbar-track { background: #1f2833; }
    ::-webkit-scrollbar-thumb { background: #00ffff; border-radius: 4px; }

    /* ---------- BANNER ---------- */
    .banner-container {
        width: 100%;
        height: 80px;
        background: linear-gradient(90deg, #000428, #004e92);
        border-bottom: 2px solid cyan;
        overflow: hidden;
        display: flex;
        align-items: center;
        justify-content: flex-start;
        position: relative;
        margin-bottom: 20px;
        border-radius: 12px;
        backdrop-filter: blur(14px);
    }
    .pulse-bar {
        position: absolute;
        display: flex;
        align-items: center;
        font-size: 22px;
        font-weight: bold;
        color: #00ffff;
        white-space: nowrap;
        animation: glideIn 12s linear infinite;
        text-shadow: 0 0 10px #00ffff;
    }
    .pulse-bar .bar {
        width: 10px;
        height: 30px;
        margin-right: 10px;
        background: #00ffff;
        box-shadow: 0 0 8px cyan;
        animation: pulse 1s ease-in-out infinite;
    }
    @keyframes glideIn {
        0% { left: -50%; opacity: 0; }
        10% { opacity: 1; }
        90% { opacity: 1; }
        100% { left: 110%; opacity: 0; }
    }
    @keyframes pulse {
        0%, 100% { height: 20px; background-color: #00ffff; }
        50% { height: 40px; background-color: #ff00ff; }
    }

    /* ---------- HEADER ---------- */
    .header {
        font-size: 28px;
        font-weight: bold;
        text-align: center;
        text-transform: uppercase;
        letter-spacing: 2px;
        padding: 20px 30px;  /* ✅ More spacing inside the bar */
        color: #00ffff;
        text-shadow: 0px 0px 10px #00ffff;
        position: relative;
        overflow: hidden;
        border-radius: 14px;
        background: rgba(10,20,40,0.35);
        backdrop-filter: blur(14px);
        border: 1px solid rgba(0,200,255,0.5);
        box-shadow: 0 0 12px rgba(0,200,255,0.25);
    }
    .header::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.18) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .header:hover::before { left: 100%; top: 100%; }

    /* ---------- SHIMMER (COMMON) ---------- */
    .shimmer::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .shimmer:hover::before { left: 100%; top: 100%; }

    /* ---------- FILE UPLOADER ---------- */
    .stFileUploader > div > div {
        border: 1px solid rgba(0,200,255,0.5);
        border-radius: 14px;
        background: rgba(10,20,40,0.35);
        backdrop-filter: blur(14px);
        color: #cce6ff;
        box-shadow: 0 0 12px rgba(0,200,255,0.3);
        position: relative;
        overflow: hidden;
    }
    .stFileUploader > div > div::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stFileUploader > div > div:hover::before { left: 100%; top: 100%; }

    /* ---------- BUTTONS ---------- */
    .stButton > button {
        position: relative;
        overflow: hidden;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        color: #e6f7ff;
        border-radius: 14px;
        padding: 10px 20px;
        font-size: 16px;
        font-weight: 500;
        text-transform: uppercase;
        backdrop-filter: blur(16px);
        box-shadow: 0 0 12px rgba(0,200,255,0.35),
                    inset 0 0 20px rgba(0,200,255,0.05);
        transition: all 0.3s ease-in-out;
    }
    .stButton > button::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stButton > button:hover::before { left: 100%; top: 100%; }

    /* ---------- INPUTS ---------- */
    .stTextInput > div > input,
    .stTextArea > div > textarea {
        position: relative;
        overflow: hidden;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        border-radius: 14px;
        color: #e6f7ff;
        padding: 10px;
        backdrop-filter: blur(16px);
        box-shadow: 0 0 12px rgba(0,200,255,0.3),
                    inset 0 0 15px rgba(0,200,255,0.05);
        transition: all 0.3s ease-in-out;
    }

    /* ---------- CHAT MESSAGES ---------- */
    .stChatMessage {
        position: relative;
        overflow: hidden;
        font-size: 18px;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.5);
        border-radius: 14px;
        padding: 14px;
        color: #e6f7ff;
        text-shadow: 0 0 6px rgba(0,200,255,0.7);
        box-shadow: 0 0 12px rgba(0,200,255,0.3),
                    inset 0 0 15px rgba(0,200,255,0.05);
    }
    .stChatMessage::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stChatMessage:hover::before { left: 100%; top: 100%; }

    /* ---------- METRICS ---------- */
    .stMetric {
        position: relative;
        overflow: hidden;
        background-color: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        border-radius: 14px;
        padding: 15px;
        box-shadow: 0 0 12px rgba(0,200,255,0.35),
                    inset 0 0 20px rgba(0,200,255,0.05);
        text-align: center;
    }
    .stMetric::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stMetric:hover::before { left: 100%; top: 100%; }

    /* ---------- MOBILE ---------- */
    @media (max-width: 768px) {
        .pulse-bar { font-size: 16px; }
        .header { font-size: 20px; }
    }
    </style>

    <!-- Banner -->
    <div class="banner-container">
        <div class="pulse-bar">
            <div class="bar"></div>
            <div>HIRELYZER - Elevate Your Resume Analysis</div>
        </div>
    </div>

    <!-- Header -->
    <div class="header">💼 HIRELYZER - AI BASED ETHICAL RESUME ANALYZER</div>
    """, unsafe_allow_html=True)

# Load environment variables
load_dotenv()

# Detect Device
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
torch.backends.cudnn.benchmark = True
working_dir = os.path.dirname(os.path.abspath(__file__))

# ------------------- Lazy Initialization -------------------
@st.cache_resource(show_spinner=False)
def get_easyocr_reader():
    import easyocr
    return easyocr.Reader(["en"], gpu=torch.cuda.is_available())

@st.cache_data(show_spinner=False)
def ensure_nltk():
    import nltk
    nltk.download('wordnet', quiet=True)
    return WordNetLemmatizer()

lemmatizer = ensure_nltk()
reader = get_easyocr_reader()

def generate_docx(text, filename="bias_free_resume.docx"):
    doc = Document()

    # ── Page margins (standard resume: 1 inch all sides) ──
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

    # ── Document title heading ──
    title = doc.add_heading('Bias-Free Resume', 0)
    title.alignment = 1  # center
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0x2F, 0x4F, 0x6F)
    title_run.font.size = Pt(18)

    doc.add_paragraph()  # spacer

    # ── Process text: detect section headers and bullet points ──
    lines = text.strip().split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Section headers (emoji + CAPS or all-caps lines)
        if (stripped.isupper() and len(stripped) > 3) or \
           any(stripped.startswith(e) for e in ['🏷️','📞','📧','📍','🔗','🌐','✍️','🛠️','💼','🧑‍💼','📂','🎓','🏫','🤝','🌟','🎯']):
            p = doc.add_heading(stripped, level=2)
            p.runs[0].font.color.rgb = RGBColor(0x2F, 0x4F, 0x6F)
            p.runs[0].font.size = Pt(12)
            continue

        # Bullet points
        if stripped.startswith(('•', '-', '*')):
            content = stripped.lstrip('•-* ').strip()
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(content)
            run.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(3)
            continue

        # Regular paragraph
        p = doc.add_paragraph(stripped)
        p.runs[0].font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(4)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Extract text from PDF
def _extract_page_text_smart(page) -> str:
    """
    Extract text from a single fitz page in correct reading order,
    handling both simple single-column and complex multi-column / graphic
    resume layouts (like sidebar designs, blue-header templates, etc.).

    Strategy:
      1. Pull raw text blocks with full bounding-box coordinates.
      2. Detect whether the page has a multi-column layout by checking
         whether meaningful content exists in both the left (<45%) and
         right (>52%) horizontal zones.
      3. Single-column  → sort blocks top-to-bottom (y0), then left-to-right.
      4. Multi-column   → split blocks into left/right columns,
                          sort each column top-to-bottom independently,
                          then concatenate left column first (the main body
                          on most sidebar resume designs sits on the right,
                          but the name/header almost always spans full width
                          or sits at the very top — so we sort by y0 first
                          for the header region, then by column).
    This ensures the candidate name — which is always the topmost block —
    is the very first text we see regardless of layout.
    """
    blocks = page.get_text("blocks")  # (x0, y0, x1, y1, text, block_no, block_type)
    page_width = page.rect.width

    # Filter: only text blocks (block_type == 0) with real content
    text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]

    if not text_blocks:
        return ""

    # ── Detect multi-column layout ────────────────────────────────────────
    x_starts = [b[0] for b in text_blocks if len(b[4].strip()) > 10]
    left_zone  = [x for x in x_starts if x < page_width * 0.45]
    right_zone = [x for x in x_starts if x > page_width * 0.52]
    is_multicolumn = len(left_zone) >= 3 and len(right_zone) >= 3

    if not is_multicolumn:
        # ── Single-column: simple top-to-bottom sort ──────────────────────
        sorted_blocks = sorted(text_blocks, key=lambda b: (round(b[1] / 10) * 10, b[0]))
        return "\n".join(b[4].strip() for b in sorted_blocks)

    # ── Multi-column: split into header + left + right zones ─────────────
    # Blocks in the top 15% of page height are "header" — name, title, etc.
    # They are sorted purely by y0 so the name always comes first.
    page_height  = page.rect.height
    header_zone  = page_height * 0.15

    header_blocks = [b for b in text_blocks if b[1] < header_zone]
    body_blocks   = [b for b in text_blocks if b[1] >= header_zone]

    # Sort header top-to-bottom
    header_sorted = sorted(header_blocks, key=lambda b: (b[1], b[0]))

    # Split body into left / right columns and sort each top-to-bottom
    left_blocks  = sorted(
        [b for b in body_blocks if b[0] < page_width * 0.48],
        key=lambda b: b[1]
    )
    right_blocks = sorted(
        [b for b in body_blocks if b[0] >= page_width * 0.48],
        key=lambda b: b[1]
    )

    # Concatenate: header → left column → right column
    all_sorted = header_sorted + left_blocks + right_blocks
    return "\n".join(b[4].strip() for b in all_sorted)


def extract_text_from_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        text_list = []
        for page in doc:
            page_text = _extract_page_text_smart(page)
            if page_text.strip():
                text_list.append(page_text)
        doc.close()
        return text_list if text_list else extract_text_from_images(file_path)
    except Exception as e:
        st.error(f"⚠ Error extracting text: {e}")
        return []

def extract_text_from_images(pdf_path):
    try:
        images = convert_from_path(pdf_path, dpi=150, first_page=1, last_page=5)
        return ["\n".join(reader.readtext(np.array(img), detail=0)) for img in images]
    except Exception as e:
        st.error(f"⚠ Error extracting from image: {e}")
        return []

def safe_extract_text(uploaded_file):
    """
    Safely extracts text from uploaded file.
    Prevents app crash if file is not a resume or unreadable.
    """
    try:
        # Save uploaded file to a temp location
        temp_path = f"/tmp/{uploaded_file.name}"
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # Try PDF text extraction
        text_list = extract_text_from_pdf(temp_path)

        # If nothing readable found
        if not text_list or all(len(t.strip()) == 0 for t in text_list):
            st.warning("⚠️ This file doesn't look like a resume or contains no readable text.")
            return None

        return "\n".join(text_list)

    except Exception as e:
        st.error(f"⚠️ Could not process this file: {e}")
        return None


# ============================================================
# 🏷️ Deterministic Candidate Name Extractor
# ============================================================



# ============================================================
# 📐 Industry-Standard Resume Format Checker (v2 — Enhanced)
# ============================================================

def _detect_multicolumn_pdf(pdf_path: str) -> bool:
    """
    Detect multi-column layout by analysing raw text-block x-coordinates
    from the first page of the PDF via PyMuPDF.
    Returns True if two or more distinct horizontal content zones are found.
    """
    try:
        doc = fitz.open(pdf_path)
        page = doc[0]
        blocks = page.get_text("blocks")   # (x0, y0, x1, y1, text, block_no, block_type)
        page_width = page.rect.width
        doc.close()

        # Only consider blocks with meaningful text content
        x_starts = [b[0] for b in blocks if len(b[4].strip()) > 10]
        if len(x_starts) < 6:
            return False

        # Split the page width into left zone (< 45 %) and right zone (> 52 %)
        left_zone  = [x for x in x_starts if x < page_width * 0.45]
        right_zone = [x for x in x_starts if x > page_width * 0.52]

        # Multi-column confirmed when both zones carry real content
        return len(left_zone) >= 3 and len(right_zone) >= 3
    except Exception:
        return False


def check_resume_format(text: str, num_pages: int = 1, pdf_path: str = None) -> dict:
    """
    Evaluates a resume against industry-standard ATS formatting rules.

    Scoring model (100 pts total, deduction-based):
      Section presence       — up to −42 pts  (critical ATS sections)
      Contact completeness   — up to −8  pts
      Resume length          — up to −14 pts
      Action verb quality    — up to −8  pts
      Quantified achievements— up to −8  pts
      ATS red flags          — up to −14 pts  (multi-column, dates, objective)
      Bonus credits          — up to +4  pts  (certifications, portfolio)

    Returns a dict compatible with all existing callers (same keys as v1).
    """
    issues  = []
    passes  = []
    deductions = 0
    bonuses    = 0

    text_lower = text.lower() if text else ""

    # ══════════════════════════════════════════════════════════════════════
    # 1. CRITICAL SECTION PRESENCE  (max −42 pts)
    #    Weights reflect real ATS auto-reject risk, not equal treatment.
    # ══════════════════════════════════════════════════════════════════════
    section_checks = {
        # (label, detected, deduction_if_missing)
        "Contact / Email": (
            bool(re.search(r'[\w.+-]+@[\w.-]+\.[a-z]{2,}', text or "")),
            15,   # ATS hard-rejects without contact info
        ),
        "Phone Number": (
            bool(re.search(r'(\+?\d[\d\s\-\(\)]{7,}\d)', text or "")),
            6,
        ),
        "Experience": (
            any(w in text_lower for w in [
                "experience", "employment", "work history", "career",
                "professional experience", "work experience",
                "positions held", "relevant experience", "professional background",
            ]),
            12,   # Core content section — high ATS weight
        ),
        "Education": (
            any(w in text_lower for w in [
                "education", "university", "college", "degree",
                "bachelor", "master", "b.tech", "b.sc", "m.sc",
                "mca", "bca", "phd", "diploma", "high school",
                "graduated", "pursuing",
            ]),
            4,
        ),
        "Skills": (
            any(w in text_lower for w in [
                "skills", "technologies", "tech stack",
                "competencies", "proficiencies", "tools",
                "technical skills", "core competencies",
            ]),
            3,
        ),
        "Summary / Profile": (
            any(w in text_lower for w in [
                "summary", "objective", "profile",
                "about me", "overview", "professional summary",
                "career objective", "personal statement",
            ]),
            2,
        ),
    }

    for section_name, (present, penalty) in section_checks.items():
        if present:
            passes.append(f"Section present: {section_name}")
        else:
            issues.append(
                f"Missing section: '{section_name}' — "
                f"ATS {'will likely reject' if penalty >= 10 else 'may penalise'} without it"
            )
            deductions += penalty

    # ══════════════════════════════════════════════════════════════════════
    # 2. CONTACT COMPLETENESS  (max −8 pts)
    # ══════════════════════════════════════════════════════════════════════
    if re.search(r'linkedin\.com/in/[\w\-]+', text_lower):
        passes.append("LinkedIn profile URL detected")
    else:
        issues.append("No LinkedIn URL — recruiters expect it; many ATS rank it as a signal")
        deductions += 5

    if re.search(r'github\.com/[\w\-]+', text_lower):
        passes.append("GitHub profile URL detected")
        bonuses += 1   # bonus: shows technical proof of work
    elif re.search(r'(portfolio|behance\.net|dribbble\.com|leetcode\.com|kaggle\.com)', text_lower):
        passes.append("Portfolio / professional profile URL detected")
        bonuses += 1
    else:
        issues.append("No GitHub or portfolio URL — especially important for technical roles")
        deductions += 3

    # ══════════════════════════════════════════════════════════════════════
    # 3. RESUME LENGTH  (max −14 pts)
    #    Sweet spot: 400–900 words for most roles.
    # ══════════════════════════════════════════════════════════════════════
    word_count = len(text.split()) if text else 0

    if word_count < 150:
        issues.append(
            f"Resume critically short ({word_count} words) — "
            "ATS expects 400–900 words; this will likely be filtered out"
        )
        deductions += 12
    elif word_count < 400:
        issues.append(
            f"Resume too short ({word_count} words) — "
            "aim for 400–900 words with detailed experience and skills"
        )
        deductions += 7
    elif word_count > 1400:
        issues.append(
            f"Resume too long ({word_count} words) — "
            "trim to under 1,000 words; ATS and recruiters prefer concise resumes"
        )
        deductions += 5
    elif word_count > 1000:
        issues.append(
            f"Resume slightly long ({word_count} words) — "
            "consider tightening to under 1,000 words"
        )
        deductions += 2
    else:
        passes.append(f"Optimal length ({word_count} words — within 400–1,000 word sweet spot)")

    if num_pages > 2:
        issues.append(
            f"Resume is {num_pages} pages — "
            "ATS industry standard is 1–2 pages; longer resumes are often truncated"
        )
        deductions += 4
    elif num_pages == 2:
        passes.append("Page count acceptable (2 pages — standard for 5+ years experience)")
    else:
        passes.append("Page count ideal (1 page — strong for early-career candidates)")

    # ══════════════════════════════════════════════════════════════════════
    # 4. ACTION VERB QUALITY  (max −8 pts)
    #    Expanded to 54 verbs across all common resume categories.
    # ══════════════════════════════════════════════════════════════════════
    strong_verbs = [
        # Engineering / technical
        "architected", "engineered", "designed", "deployed", "optimized", "automated",
        "built", "launched", "developed", "implemented", "integrated", "configured",
        "migrated", "refactored", "debugged", "benchmarked", "containerized", "scaled",
        "maintained", "upgraded", "tested", "validated",
        # Leadership / management
        "led", "managed", "directed", "oversaw", "supervised", "coordinated",
        "spearheaded", "mentored", "trained", "guided", "facilitated",
        # Business / impact
        "reduced", "increased", "improved", "accelerated", "streamlined", "transformed",
        "negotiated", "established", "executed", "delivered", "created",
        "resolved", "analyzed", "collaborated", "authored", "published",
        # Data / research
        "researched", "evaluated", "identified", "modelled", "forecasted",
        "presented", "reported", "drafted",
    ]
    found_verbs = [v for v in strong_verbs if re.search(rf'\b{v}\b', text_lower)]
    verb_count = len(found_verbs)

    if verb_count == 0:
        issues.append(
            "No strong action verbs found — ATS and recruiters expect bullet points "
            "starting with verbs like 'Engineered', 'Led', 'Optimized'"
        )
        deductions += 8
    elif verb_count < 3:
        issues.append(
            f"Weak action verb usage ({verb_count} found) — "
            "aim for 5+ distinct strong verbs across experience bullet points"
        )
        deductions += 5
    elif verb_count < 5:
        issues.append(
            f"Limited action verb variety ({verb_count} found) — "
            "diversify verbs to better demonstrate range of contributions"
        )
        deductions += 2
    else:
        passes.append(f"Strong action verb usage ({verb_count} distinct verbs detected)")

    # ══════════════════════════════════════════════════════════════════════
    # 5. QUANTIFIED ACHIEVEMENTS  (max −8 pts)
    #    Broader pattern set captures $, %, x, K, M, large numbers, etc.
    # ══════════════════════════════════════════════════════════════════════
    quant_patterns = []

    # Percentage metrics: 35%, 2.5 percent
    quant_patterns += re.findall(
        r'\b\d+[\.,]?\d*\s*(%|percent)\b', text_lower
    )
    # Multiplier / scale: 10x, 3 times
    quant_patterns += re.findall(
        r'\b\d+[\.,]?\d*\s*(x|times)\b', text_lower
    )
    # Counts with units: 10K users, 500 clients, 3 projects, 50 hours
    quant_patterns += re.findall(
        r'\b\d+[,.]?\d*\s*(k|m)?\s*(users|clients|customers|projects|tickets|'
        r'requests|transactions|queries|hrs|hours|days|weeks|months|years|'
        r'engineers|developers|members|students|candidates|submissions)\b',
        text_lower
    )
    # Technical metrics: 200ms, 50GB, 1TB
    quant_patterns += re.findall(
        r'\b\d+[\.,]?\d*\s*(ms|gb|tb|mb|rpm|rps|qps|wpm|tps)\b', text_lower
    )
    # Dollar amounts: $50K, $1.2M, $500
    quant_patterns += re.findall(
        r'\$\s*\d+[\d,.]*\s*[kKmMbB]?\b', text_lower
    )
    # Large bare numbers (10,000+) — likely meaningful scale references
    quant_patterns += re.findall(
        r'\b\d{1,3}[,]\d{3}\b', text_lower
    )
    # Qualitative scale indicators
    quant_patterns += re.findall(
        r'\b(doubled|tripled|halved|10x|100x)\b', text_lower
    )
    # "3+ years" style
    quant_patterns += re.findall(
        r'\b\d+\+\s*(years|yrs|months)\b', text_lower
    )

    metric_count = len(quant_patterns)
    if metric_count == 0:
        issues.append(
            "No quantified achievements detected — add measurable impact "
            "(e.g., 'reduced latency by 35%', 'served 10K users', 'saved $50K annually')"
        )
        deductions += 8
    elif metric_count < 3:
        issues.append(
            f"Few quantified achievements ({metric_count} found) — "
            "aim for 4+ metrics across your experience to demonstrate concrete impact"
        )
        deductions += 4
    else:
        passes.append(f"Quantified achievements present ({metric_count} metrics detected)")

    # ══════════════════════════════════════════════════════════════════════
    # 6. ATS RED FLAGS  (max −14 pts)
    # ══════════════════════════════════════════════════════════════════════

    # 6a. Multi-column layout detection
    #     Priority: use real PDF block coordinates if pdf_path available,
    #     fall back to tab-character heuristic for plain-text paths.
    multicolumn_detected = False
    if pdf_path:
        try:
            multicolumn_detected = _detect_multicolumn_pdf(pdf_path)
        except Exception:
            multicolumn_detected = False
    if not multicolumn_detected and text:
        # Fallback heuristic: heavy tab usage OR many pipe characters
        multicolumn_detected = (
            text.count('\t') > 8 or
            text.count('|') > 12
        )

    if multicolumn_detected:
        issues.append(
            "Multi-column or table layout detected — "
            "many ATS parsers read columns out of order, scrambling your resume content; "
            "use a single-column layout"
        )
        deductions += 7
    else:
        passes.append("Single-column layout detected — ATS-safe structure")

    # 6b. Outdated 'Objective' section
    if "objective" in text_lower and "summary" not in text_lower and "professional summary" not in text_lower:
        issues.append(
            "Uses 'Objective' section — this is outdated; "
            "replace with a modern 'Professional Summary' (2–3 targeted sentences)"
        )
        deductions += 3

    # 6c. Employment dates
    has_dates = bool(re.search(r'\b(19|20)\d{2}\b', text or ""))
    if not has_dates:
        issues.append(
            "No employment dates detected — "
            "ATS requires dates to build a timeline; "
            "add month/year ranges (e.g., 'Jan 2021 – Mar 2023')"
        )
        deductions += 5
    else:
        passes.append("Employment dates detected — ATS can parse your timeline")

    # 6d. Special characters / encoding issues that confuse ATS parsers
    if text:
        special_char_count = len(re.findall(r'[^\x00-\x7F]', text))
        ratio = special_char_count / max(len(text), 1)
        if ratio > 0.04:
            issues.append(
                f"High non-ASCII character density ({special_char_count} chars) — "
                "special characters from stylised fonts or copy-paste can corrupt ATS parsing"
            )
            deductions += 3
        else:
            passes.append("Character encoding looks ATS-safe (low non-ASCII density)")

    # 6e. Excessive repetition of buzzwords (keyword stuffing signal)
    buzzwords = ["synergy", "passionate", "hardworking", "go-getter", "think outside the box",
                 "detail-oriented", "team player", "results-driven", "dynamic", "proactive"]
    stuffed = [bw for bw in buzzwords if text_lower.count(bw) >= 2]
    if len(stuffed) >= 2:
        issues.append(
            f"Possible keyword stuffing detected ({', '.join(stuffed)}) — "
            "overused buzzwords reduce credibility; replace with concrete examples"
        )
        deductions += 2

    # ══════════════════════════════════════════════════════════════════════
    # 7. BONUS CREDITS  (up to +4 pts)
    #    Reward genuine ATS positive signals.
    # ══════════════════════════════════════════════════════════════════════

    # Certifications section
    if any(w in text_lower for w in [
        "certification", "certified", "certificate", "aws certified",
        "google certified", "microsoft certified", "pmp", "cpa",
        "cissp", "ceh", "comptia", "coursera", "udemy", "edx",
    ]):
        passes.append("Certifications / credentials detected — strong ATS positive signal")
        bonuses += 1

    # Projects section
    if any(w in text_lower for w in [
        "projects", "personal projects", "side projects",
        "open source", "github.com", "hackathon",
    ]):
        passes.append("Projects section detected — demonstrates initiative beyond job roles")
        bonuses += 1

    # Consistent date format (month year)
    month_year_dates = re.findall(
        r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*(19|20)\d{2}\b',
        text_lower
    )
    if len(month_year_dates) >= 2:
        passes.append("Consistent Month-Year date format detected — preferred by ATS parsers")
        bonuses += 1

    # ══════════════════════════════════════════════════════════════════════
    # 8. FINAL SCORE CALCULATION
    # ══════════════════════════════════════════════════════════════════════
    raw_score = max(0, min(100, 100 - deductions + bonuses))

    if raw_score >= 90:
        letter_grade = "A+"
        label = "ATS-Optimized"
    elif raw_score >= 80:
        letter_grade = "A"
        label = "Excellent Format"
    elif raw_score >= 70:
        letter_grade = "B+"
        label = "Good Format"
    elif raw_score >= 60:
        letter_grade = "B"
        label = "Acceptable"
    elif raw_score >= 45:
        letter_grade = "C"
        label = "Needs Work"
    else:
        letter_grade = "D"
        label = "Poor — Major Issues"

    return {
        "format_score":  raw_score,
        "letter_grade":  letter_grade,
        "label":         label,
        "issues":        issues,
        "passes":        passes,
        "word_count":    word_count,
        # Extended breakdown (available to callers that want sub-scores)
        "deductions":    deductions,
        "bonuses":       bonuses,
        "verb_count":    verb_count,
        "metric_count":  metric_count,
        "multicolumn":   multicolumn_detected,
    }

# Detect bias in resume
# Predefined gender-coded word lists
gender_words = {
    "masculine": [
        # Dominance / aggression-coded
        "active", "aggressive", "ambitious", "assertive", "autonomous", "boast", "bold",
        "challenging", "competitive", "confident", "courageous", "decisive", "determined", "dominant", "driven",
        "dynamic", "forceful", "independent", "individualistic", "intellectual", "lead", "leader", "objective",
        "outspoken", "persistent", "principled", "proactive", "resilient", "self-reliant", "self-sufficient",
        "strong", "superior", "tenacious", "guru", "tech guru", "technical guru", "visionary", "manpower",
        "strongman", "command", "assert", "headstrong", "rockstar", "superstar", "go-getter", "trailblazer",
        "results-driven", "fast-paced", "determination", "competitive spirit",
        # Additional research-backed masculine-coded terms (Gaucher et al., 2011)
        "analytical", "backbone", "challenge", "champion", "combat", "conquer", "courageous",
        "crusade", "debate", "fearless", "fight", "grit", "hustle", "impact", "ninja",
        "power", "ruthless", "self-starter", "sharp", "warrior", "win", "wrestler",
        "alpha", "beast", "brutally honest", "cutting-edge", "dominate", "edge", "elite",
        "fearless", "grind", "hardcore", "hero", "high-performance", "intense",
        "kill it", "relentless", "savage", "slayer", "tiger", "tough", "uncompromising"
    ],
    
    "feminine": [
        # Communal / warmth-coded
        "affectionate", "agreeable", "attentive", "collaborative", "committed", "compassionate", "considerate",
        "cooperative", "dependable", "dependent", "emotional", "empathetic", "enthusiastic", "friendly", "gentle",
        "honest", "inclusive", "interpersonal", "kind", "loyal", "modest", "nurturing", "pleasant", "polite",
        "sensitive", "supportive", "sympathetic", "tactful", "tender", "trustworthy", "understanding", "warm",
        "yield", "adaptable", "communal", "helpful", "dedicated", "respectful", "nurture", "sociable",
        "relationship-oriented", "team player", "people-oriented", "empathetic listener",
        "gentle communicator", "open-minded",
        # Additional research-backed feminine-coded terms
        "balance", "caring", "child-friendly", "connect", "connection", "flexible hours",
        "harmony", "heart", "humanize", "mindful", "patience", "patient", "peace",
        "personal touch", "responsive", "share", "sharing", "together", "unite",
        "welcoming", "wholesome", "connect with", "feeling", "feelings", "giving back",
        "heartfelt", "humanity", "inspire", "inspired", "passion", "passionate",
        "personable", "relate", "relatable", "soften", "soft skills", "spread",
        "thrive", "togetherness", "transparent", "uplift", "vulnerable"
    ]
}

def detect_bias(text):
    # Split into sentences using simple delimiters
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())

    masc_set, fem_set = set(), set()
    masculine_found, feminine_found = [] , []

    masculine_words = sorted(gender_words["masculine"], key=len, reverse=True)
    feminine_words = sorted(gender_words["feminine"], key=len, reverse=True)

    for sent in sentences:
        sent_text = sent.strip()
        sent_lower = sent_text.lower()
        matched_spans = []

        def is_overlapping(start, end):
            return any(start < e and end > s for s, e in matched_spans)

        # 🔵 Highlight masculine words in blue
        for word in masculine_words:
            pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
            for match in pattern.finditer(sent_lower):
                start, end = match.span()
                if not is_overlapping(start, end):
                    matched_spans.append((start, end))
                    key = (word.lower(), sent_text)
                    if key not in masc_set:
                        masc_set.add(key)
                        highlighted = re.sub(
                            rf'\b({re.escape(word)})\b',
                            r'<span style="color:blue;">\1</span>',
                            sent_text,
                            flags=re.IGNORECASE
                        )
                        masculine_found.append({
                            "word": word,
                            "sentence": highlighted
                        })

        # 🔴 Highlight feminine words in red
        for word in feminine_words:
            pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
            for match in pattern.finditer(sent_lower):
                start, end = match.span()
                if not is_overlapping(start, end):
                    matched_spans.append((start, end))
                    key = (word.lower(), sent_text)
                    if key not in fem_set:
                        fem_set.add(key)
                        highlighted = re.sub(
                            rf'\b({re.escape(word)})\b',
                            r'<span style="color:red;">\1</span>',
                            sent_text,
                            flags=re.IGNORECASE
                        )
                        feminine_found.append({
                            "word": word,
                            "sentence": highlighted
                        })

    masc = len(masculine_found)
    fem = len(feminine_found)
    total = masc + fem
    bias_score = min(total / 20, 1.0) if total > 0 else 0.0

    return round(bias_score, 2), masc, fem, masculine_found, feminine_found

replacement_mapping = {
    "masculine": {
        "active": "engaged",
        "aggressive": "proactive",
        "ambitious": "motivated",
        "analytical": "detail-oriented",
        "assertive": "direct",
        "autonomous": "self-directed",
        "boast": "highlight",
        "bold": "confident",
        "challenging": "demanding",
        "competitive": "goal-oriented",
        "confident": "self-assured",
        "courageous": "bold",
        "decisive": "action-oriented",
        "determined": "focused",
        "dominant": "influential",
        "driven": "committed",
        "dynamic": "adaptable",
        "forceful": "persuasive",
        "guru": "technical expert",
        "independent": "self-sufficient",
        "individualistic": "self-motivated",
        "intellectual": "knowledgeable",
        "lead": "guide",
        "leader": "team lead",
        "objective": "unbiased",
        "outspoken": "expressive",
        "persistent": "tenacious",
        "principled": "ethical",
        "proactive": "initiative-taking",
        "resilient": "adaptable",
        "self-reliant": "resourceful",
        "self-sufficient": "capable",
        "strong": "capable",
        "superior": "exceptional",
        "tenacious": "determined",
        "technical guru": "technical expert",
        "visionary": "forward-thinking",
        "manpower": "workforce",
        "strongman": "resilient individual",
        "command": "direct",
        "assert": "state clearly",
        "headstrong": "determined",
        "rockstar": "top performer",
        "superstar": "outstanding contributor",
        "go-getter": "initiative-taker",
        "trailblazer": "innovator",
        "results-driven": "outcome-focused",
        "fast-paced": "dynamic",
        "determination": "commitment",
        "competitive spirit": "goal-oriented mindset",
        # New additions
        "ninja": "specialist",
        "warrior": "dedicated professional",
        "alpha": "senior",
        "beast": "high performer",
        "dominate": "excel in",
        "elite": "high-performing",
        "relentless": "persistent",
        "savage": "highly skilled",
        "hustle": "work efficiently",
        "grit": "resilience",
        "hardcore": "rigorous",
        "hero": "key contributor",
        "ruthless": "highly focused",
        "kill it": "excel",
        "champion": "advocate",
        "conquer": "achieve",
        "fight": "address",
        "win": "achieve success",
        "crush": "exceed targets",
        "unstoppable": "highly motivated",
        "fearless": "courageous",
        "power": "capability",
        "backbone": "core strength",
        "sharp": "perceptive"
    },
    
    "feminine": {
        "affectionate": "approachable",
        "agreeable": "cooperative",
        "attentive": "observant",
        "collaborative": "team-oriented",
        "collaborate": "team-oriented",
        "collaborated": "worked together",
        "committed": "dedicated",
        "compassionate": "caring",
        "considerate": "thoughtful",
        "cooperative": "supportive",
        "dependable": "reliable",
        "dependent": "team-oriented",
        "emotional": "passionate",
        "empathetic": "perceptive",
        "enthusiastic": "energized",
        "gentle": "respectful",
        "honest": "transparent",
        "inclusive": "open-minded",
        "interpersonal": "people-focused",
        "kind": "respectful",
        "loyal": "dedicated",
        "modest": "measured",
        "nurturing": "supportive",
        "pleasant": "professional",
        "polite": "courteous",
        "sensitive": "perceptive",
        "supportive": "enabling",
        "sympathetic": "understanding",
        "tactful": "diplomatic",
        "tender": "considerate",
        "trustworthy": "reliable",
        "understanding": "empathetic",
        "warm": "welcoming",
        "yield": "adjust",
        "adaptable": "flexible",
        "communal": "team-centered",
        "helpful": "contributive",
        "dedicated": "committed",
        "respectful": "professional",
        "nurture": "develop",
        "sociable": "collegial",
        "relationship-oriented": "team-focused",
        "team player": "collaborative member",
        "people-oriented": "stakeholder-focused",
        "empathetic listener": "active listener",
        "gentle communicator": "considerate communicator",
        "open-minded": "inclusive",
        # New additions
        "passionate": "highly motivated",
        "inspired": "driven by purpose",
        "inspire": "motivate",
        "vulnerable": "transparent",
        "heartfelt": "sincere",
        "harmony": "alignment",
        "caring": "attentive",
        "patient": "thorough",
        "wholesome": "balanced",
        "togetherness": "team cohesion",
        "soft skills": "professional competencies",
        "personal touch": "tailored approach",
        "feeling": "assessment",
        "feelings": "perspectives",
        "transparent": "accountable",
        "uplift": "elevate",
        "thrive": "excel",
        "welcoming": "inclusive",
        "relatable": "accessible",
        "connect": "engage",
        "together": "collaboratively",
        "sharing": "distributing",
        "mindful": "deliberate",
        "balance": "manage effectively"
    }
}

def rewrite_and_optimize_resume(text, replacement_mapping, user_location):
    """
    ⚡ MERGED FUNCTION — replaces rewrite_text_with_llm() + optimize_resume_to_json()
    Single LLM call that returns BOTH:
      - rewritten_text : plain-text ATS-optimised resume + job title suggestions (for UI display)
      - json_str       : strict JSON object (for DOCX generation)
    Saves 1 API key call per resume analysis (6 → 5 calls total).
    """

    formatted_mapping = "\n".join(
        [f'- "{key}" → "{value}"' for key, value in replacement_mapping.items()]
    )

    prompt = f"""You are an enterprise-grade ATS resume optimization engine and bias-removal specialist.

Your task is to process the resume below and return TWO outputs in a single response, separated by the exact delimiter shown.

════════════════════════════════════════════════════════
OUTPUT STRUCTURE (return EXACTLY this — no deviation):
════════════════════════════════════════════════════════

===REWRITTEN_RESUME_START===
<full plain-text ATS-optimised resume here>
<followed by job title suggestions block>
===REWRITTEN_RESUME_END===

===JSON_START===
<strict JSON object here — no markdown fences, no explanation>
===JSON_END===

════════════════════════════════════════════════════════
PART 1 — PLAIN TEXT REWRITE (inside REWRITTEN_RESUME tags)
════════════════════════════════════════════════════════

ABSOLUTE RULES — NEVER VIOLATE:
• DO NOT fabricate companies, job titles, degrees, institutions, or dates.
• DO NOT invent statistics or metrics not implied by the resume content.
• DO NOT add certifications, tools, or skills absent from the resume.
• DO NOT use personal pronouns (I, my, me, we, our) anywhere.
• DO NOT repeat the same phrase or word across multiple sections.
• EVERY section must contain unique, non-overlapping content.

YOU MAY:
✓ Strengthen bullet points with stronger action verbs and tighter phrasing.
✓ Reconstruct missing sections when clear evidence exists in the resume.
✓ Consolidate skills scattered across experience/projects into the Skills section.
✓ Infer tool proficiency when strongly implied (e.g., "built Flask API" → Python, Flask).
✓ Add plausible impact framing using "~" when the role implies measurable output.

SECTION ORDER: Contact Header → Professional Summary → Core Skills →
Work Experience → Projects → Education → Certifications & Links

CONTACT HEADER: Full Name | Job Title | Email | Phone | Location | LinkedIn URL | GitHub/Portfolio URL

PROFESSIONAL SUMMARY (2–3 sentences, max 60 words):
  Sentence 1: [Seniority level] + [core domain] + [years of experience]
  Sentence 2: [Top 2–3 specific technical or functional strengths]
  Sentence 3: [Career value proposition — what the candidate delivers]

CORE SKILLS: labeled lines — Technical Skills: [...] and Professional Skills: [...]

WORK EXPERIENCE (reverse chronological):
  Job Title | Company Name | MMM YYYY – MMM YYYY
  [1-sentence role scope]
  • [Action Verb] + [Task] + [Technology] + [Quantified impact]
  (3–5 bullets per role)
  Strong verbs ONLY: Architected, Engineered, Developed, Implemented, Optimized, Automated,
  Spearheaded, Deployed, Designed, Reduced, Increased, Streamlined, Led, Built.
  NEVER: helped, assisted, worked on, involved in, responsible for.

PROJECTS: Name | Tech Stack | Duration
  [1-sentence purpose]
  • [Achievement bullet with action verb and metric]
  (3–5 bullets)

EDUCATION: Degree, Major | Institution | Graduation Year
CERTIFICATIONS: • Name | Issuing Body | MMM YYYY

ATS FORMATTING:
• Single-column structure — no tables, columns, text boxes.
• Bullet points: "•" only. Section headings: ALL CAPS.
• No emojis, no personal pronouns.

BIAS REPLACEMENT RULES — APPLY EXACTLY:
{formatted_mapping}

MANDATORY JOB TITLE SUGGESTIONS (append after the resume text):

### 🎯 Suggested Job Titles (Based on Resume)

Provide EXACTLY 5 job titles suited for a candidate in {user_location}.
FORMAT:
1. **[Job Title]** — [Specific reason tied to resume evidence]
🔗 https://www.linkedin.com/jobs/search/?keywords=[URL+encoded+title]&location={urllib.parse.quote(user_location)}

════════════════════════════════════════════════════════
PART 2 — JSON OBJECT (inside JSON tags)
════════════════════════════════════════════════════════

Return ONLY valid JSON. No preamble, no explanation, no markdown fences.

CONTENT REWRITING — same ATS rules as Part 1 apply to all bullet fields.
Strong verbs only. Quantified impact. No pronouns. No repetition across sections.

RETURN ONLY THIS EXACT JSON STRUCTURE:
{{
  "contact": {{
    "name": "",
    "title": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin": "",
    "github": "",
    "portfolio": ""
  }},
  "summary": "",
  "skills": [],
  "soft_skills": [],
  "languages": [],
  "interests": [],
  "experience": [
    {{
      "role": "",
      "company": "",
      "duration": "",
      "description": "",
      "bullets": []
    }}
  ],
  "projects": [
    {{
      "name": "",
      "duration": "",
      "tech_stack": "",
      "url": "",
      "description": "",
      "bullets": []
    }}
  ],
  "education": [
    {{
      "degree": "",
      "institution": "",
      "year": "",
      "bullets": []
    }}
  ],
  "certifications": [
    {{
      "name": "",
      "issuer": "",
      "duration": ""
    }}
  ],
  "additional": [
    {{
      "name": "",
      "description": "",
      "duration": ""
    }}
  ]
}}

FIELD RULES:
- "skills" = flat array of individual skill strings. Minimum 8. No duplicates.
- "soft_skills" = professional competency phrases. Must NOT duplicate items in "skills".
- "contact.*" = extract exactly as written. Use "" not null for missing fields.
- "summary" = 2–3 sentences, max 60 words, no pronouns.
- "experience[].description" = 1-sentence role scope, unique from bullets.
- "experience[].bullets" = 3–5 bullets each. Strong verb + task + tech + impact.
- "projects[].bullets" = must NOT restate experience bullets.
- "additional" items MUST use object format: {{"name":"","description":"","duration":""}}.
- Missing fields: use "[Not Provided]" for text, [] for arrays.

RESUME TEXT:
\"\"\"{text}\"\"\"
"""

    raw_response = call_llm(prompt, session=st.session_state)

    # ── Parse the two sections out of the combined response ──────────────
    rewritten_text = ""
    json_str = ""

    rewrite_match = re.search(
        r"===REWRITTEN_RESUME_START===(.*?)===REWRITTEN_RESUME_END===",
        raw_response, re.DOTALL
    )
    json_match = re.search(
        r"===JSON_START===(.*?)===JSON_END===",
        raw_response, re.DOTALL
    )

    if rewrite_match:
        rewritten_text = rewrite_match.group(1).strip()
    else:
        # fallback: use everything before JSON block
        rewritten_text = raw_response.split("===JSON_START===")[0].strip()

    if json_match:
        json_str = json_match.group(1).strip()
    else:
        # fallback: try to extract JSON object from anywhere in the response
        json_fallback = re.search(r'\{[\s\S]*\}', raw_response)
        json_str = json_fallback.group(0).strip() if json_fallback else ""

    return rewritten_text, json_str


# ── Thin compatibility wrappers — keep callers working without changes ────────

def rewrite_text_with_llm(text, replacement_mapping, user_location):
    """Compatibility wrapper — calls merged rewrite_and_optimize_resume()."""
    rewritten_text, _ = rewrite_and_optimize_resume(text, replacement_mapping, user_location)
    return rewritten_text


def optimize_resume_to_json(raw_text: str) -> str:
    """Compatibility wrapper — calls merged rewrite_and_optimize_resume()."""
    _, json_str = rewrite_and_optimize_resume(raw_text, {}, "")
    return json_str


def _salvage_additional_str(s):
    """
    Extract name/description/duration from a leaked dict/JSON string.
    Handles JSON double-quoted, Python single-quoted, and mixed formats.
    Returns a clean dict or None.
    """
    import json as _j, re as _r
    if not s or not s.strip():
        return None
    s = s.strip()
    # Attempt 1: valid JSON double-quoted keys
    try:
        sub = _j.loads(s)
        if isinstance(sub, dict):
            n = str(sub.get("name", "") or "").strip()
            d = str(sub.get("description", "") or "").strip()
            r = str(sub.get("duration", "") or "").strip()
            if n or d:
                return {"name": n, "description": d, "duration": r}
    except Exception:
        pass
    # Attempt 2: replace single quotes -> double quotes and try JSON parse
    try:
        # Preserve escaped single quotes, swap bare ones to double quotes
        converted = s.replace("\\'", "\x01").replace("'", '"').replace("\x01", "\\'")
        sub = _j.loads(converted)
        if isinstance(sub, dict):
            n = str(sub.get("name", "") or "").strip()
            d = str(sub.get("description", "") or "").strip()
            r = str(sub.get("duration", "") or "").strip()
            if n or d:
                return {"name": n, "description": d, "duration": r}
    except Exception:
        pass
    # Attempt 3: brute-force regex — key can be single OR double quoted
    def _extract(key, text):
        for q in ('"', "'"):
            pat = q + key + q + r"\s*:\s*" + q + r"(.*?)" + q + r'(?=\s*,\s*[\'"{]|\s*\})'
            m = _r.search(pat, text, _r.DOTALL)
            if m:
                return m.group(1).strip()
        return ""
    n = _extract("name", s)
    d = _extract("description", s)
    r = _extract("duration", s)
    if n or d:
        return {"name": n, "description": d, "duration": r}
    return None



def extract_resume_json(llm_response: str) -> dict:
    """
    Safely extracts and parses JSON from LLM response.
    Handles markdown fences, leading/trailing text, and partial JSON.
    Returns a dict. Falls back to empty skeleton on any parse failure.
    """
    EMPTY = {
        "contact": {
            "name": "", "title": "", "email": "", "phone": "",
            "location": "", "linkedin": "", "github": "", "portfolio": ""
        },
        "summary": "",
        "skills": [],
        "soft_skills": [],
        "languages": [],
        "interests": [],
        "experience": [],
        "projects": [],
        "education": [],
        "certifications": [],
        "additional": [],
    }
    CONTACT_DEFAULTS = {
        "name": "", "title": "", "email": "", "phone": "",
        "location": "", "linkedin": "", "github": "", "portfolio": ""
    }
    if not llm_response:
        return EMPTY
    text = llm_response.strip()
    # Strip markdown fences
    text = re.sub(r'^```(?:json)?\s*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*```$', '', text)
    text = text.strip()
    # Find first { and last }
    start = text.find('{')
    end = text.rfind('}')
    if start == -1 or end == -1:
        return EMPTY
    text = text[start:end + 1]
    try:
        data = json.loads(text)
        # Ensure all top-level keys exist
        for key in EMPTY:
            if key not in data:
                data[key] = EMPTY[key]
        if not isinstance(data.get("contact"), dict):
            data["contact"] = CONTACT_DEFAULTS.copy()
        for field, default in CONTACT_DEFAULTS.items():
            if field not in data["contact"]:
                data["contact"][field] = default

        # ── Contact field rescue: fix all LLM inconsistencies ───────────────
        ct = data["contact"]

        # 1. Null/None/"null"/"undefined" → empty string for ALL contact fields
        for field in list(ct.keys()):
            v = ct[field]
            if v is None or str(v).strip() in ("null", "None", "undefined", "N/A", "n/a"):
                ct[field] = ""

        # 2. LLM put location/linkedin/github at top level instead of inside contact
        for top_key, contact_key in [
            ("location",  "location"),
            ("linkedin",  "linkedin"),
            ("github",    "github"),
            ("portfolio", "portfolio"),
            ("github_url","github"),
            ("linkedin_url","linkedin"),
            ("profile",   "linkedin"),
            ("website",   "portfolio"),
        ]:
            top_val = data.get(top_key, "") or ""
            if top_val and not ct.get(contact_key):
                ct[contact_key] = str(top_val).strip()

        # 3. LLM used alternate key names inside contact block
        ALT_KEYS = {
            "linkedin": ["linkedin_url", "linkedin_profile", "linkedin_profile_url",
                         "linkedIn", "linked_in", "profile_url", "profile"],
            "github":   ["github_url", "github_profile", "github_link",
                         "portfolio_url", "portfolio", "website", "repo"],
            "location": ["city", "address", "city_country", "current_location"],
        }
        for canonical, alts in ALT_KEYS.items():
            if not ct.get(canonical):
                for alt in alts:
                    v = ct.get(alt, "") or ""
                    if v and str(v).strip():
                        ct[canonical] = str(v).strip()
                        break

        # 4. Final null/empty cleanup — ensure every field is a string
        for field in CONTACT_DEFAULTS:
            if not isinstance(ct.get(field), str) or ct[field] is None:
                ct[field] = ""
        # Backfill missing experience fields
        for exp in data.get("experience", []):
            for f in ["role", "company", "duration", "description"]:
                if f not in exp:
                    exp[f] = ""
            if "bullets" not in exp:
                exp["bullets"] = []
        # Backfill missing project fields
        for proj in data.get("projects", []):
            for f in ["name", "duration", "tech_stack", "url", "description"]:
                if f not in proj:
                    proj[f] = ""
            if "bullets" not in proj:
                proj["bullets"] = []
        # Backfill missing education fields
        for edu in data.get("education", []):
            for f in ["degree", "institution", "year"]:
                if f not in edu:
                    edu[f] = ""
            if "bullets" not in edu:
                edu["bullets"] = []
        # Normalise additional — accept dicts, strings, or malformed objects
        raw_add = data.get("additional", [])
        norm_add = []
        for item in raw_add:
            if isinstance(item, dict):
                name = str(item.get("name", "") or "").strip()
                desc = str(item.get("description", "") or "").strip()
                dur  = str(item.get("duration", "") or "").strip()
                if not name and not desc:
                    continue
                norm_add.append({"name": name, "description": desc, "duration": dur})
            elif isinstance(item, str):
                s = item.strip()
                if not s or s in ("[Not Provided]",):
                    continue
                # If it looks like a leaked dict/JSON string, try to parse or salvage it
                if s.startswith("{") or ("name" in s and "description" in s):
                    salvaged = _salvage_additional_str(s)
                    if salvaged:
                        norm_add.append(salvaged)
                    # else discard entirely — do NOT render raw string
                else:
                    norm_add.append({"name": s, "description": "", "duration": ""})
        data["additional"] = norm_add

        # Normalise certifications — accept both flat strings and objects
        raw_certs = data.get("certifications", [])
        norm_certs = []
        for c in raw_certs:
            if isinstance(c, dict):
                norm_certs.append({
                    "name":     c.get("name", ""),
                    "issuer":   c.get("issuer", ""),
                    "duration": c.get("duration", ""),
                })
            elif isinstance(c, str) and c.strip():
                norm_certs.append({"name": c.strip(), "issuer": "", "duration": ""})
        data["certifications"] = norm_certs
        return data
    except (json.JSONDecodeError, ValueError):
        return EMPTY


# ============================================================
# 📄 DOCX TEMPLATE GENERATORS — Three professional styles
# ============================================================

def _val(v) -> str:
    """Return value or 'Not Provided' placeholder — never empty, never None, never '[Not Provided]'."""
    if v is None:
        return "Not Provided"
    s = str(v).strip()
    if not s or s in ("[Not Provided]", "null", "None", "undefined"):
        return "Not Provided"
    return s


def _build_contact_header(doc, data: dict, name_size: int, name_color_rgb: tuple,
                           name_font: str, contact_font: str, contact_color_hex: str,
                           contact_size: int = 9, title_font: str = None,
                           title_size: int = 11, title_color_rgb: tuple = None,
                           separator: str = "  |  ",
                           label_color_hex: str = None,
                           accent_color_hex: str = None):
    """
    Builds the header block matching the exact sample template format:

      Line 1:  FULL NAME  (large bold centered)
      Line 2:  JOB TITLE  (smaller bold/normal centered uppercase)
      Line 3:  email  |  phone  |  LOCATION  |  linkedin_url  |  github_url
               (single pipe-separated line, centered, all fields always present)

    Every field always appears. Missing values show "Not Provided".
    No separate labeled rows. No dividers. Single clean contact line.
    """
    contact = data.get("contact", {})

    # ── Resolve contact color ────────────────────────────────────────────────
    cc = RGBColor(
        int(contact_color_hex[0:2], 16),
        int(contact_color_hex[2:4], 16),
        int(contact_color_hex[4:6], 16),
    )

    # ── ① Name ──────────────────────────────────────────────────────────────
    raw_name = contact.get("name", "") or ""
    name = raw_name if raw_name and raw_name not in ("", "[Not Provided]") else "Your Name"
    p_name = doc.add_paragraph()
    p_name.clear()
    r_name = p_name.add_run(name)
    r_name.bold = True
    r_name.font.size = Pt(name_size)
    r_name.font.name = name_font
    r_name.font.color.rgb = RGBColor(*name_color_rgb)
    p_name.alignment = 1
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(4)

    # ── ② Job Title — always shown, placeholder if missing ──────────────────
    raw_title = contact.get("title", "") or ""
    title_text = raw_title if raw_title and raw_title not in ("", "[Not Provided]") else "Job Title"
    p_title = doc.add_paragraph()
    p_title.clear()
    r_title = p_title.add_run(title_text.upper())
    r_title.font.size = Pt(title_size)
    r_title.font.name = title_font or name_font
    r_title.bold = True
    if title_color_rgb:
        r_title.font.color.rgb = RGBColor(*title_color_rgb)
    p_title.alignment = 1
    p_title.paragraph_format.space_after = Pt(4)

    # ── ③ Single pipe-separated contact line ────────────────────────────────
    # Build parts in order: email | phone | location | linkedin | github
    # Every field always included — "Not Provided" if missing
    def _clean(v):
        """Return value or 'Not Provided' — never None or empty."""
        if v is None:
            return "Not Provided"
        s = str(v).strip()
        return s if s and s not in ("[Not Provided]", "null", "None", "undefined") else "Not Provided"

    email_val    = _clean(contact.get("email", ""))
    phone_val    = _clean(contact.get("phone", ""))
    location_val = _clean(contact.get("location", ""))
    linkedin_val = _clean(contact.get("linkedin", ""))
    github_raw   = contact.get("github", "") or ""
    portfolio_raw= contact.get("portfolio", "") or ""
    github_val   = _clean(github_raw if github_raw and github_raw not in ("", "[Not Provided]") else portfolio_raw)

    parts = [email_val, phone_val, location_val, linkedin_val, github_val]
    contact_line = "  |  ".join(parts)

    p_contact = doc.add_paragraph()
    p_contact.clear()
    r_contact = p_contact.add_run(contact_line)
    r_contact.font.size = Pt(contact_size)
    r_contact.font.name = contact_font
    r_contact.font.color.rgb = cc
    p_contact.alignment = 1
    p_contact.paragraph_format.space_before = Pt(2)
    p_contact.paragraph_format.space_after = Pt(6)

    # ── ④ Thin bottom rule below contact block ────────────────────────────
    # Signals end of header to ATS parsers and improves recruiter readability.
    from docx.oxml import OxmlElement as _OE
    from docx.oxml.ns import qn as _qn
    pPr = p_contact._p.get_or_add_pPr()
    pBdr = _OE('w:pBdr')
    btm = _OE('w:bottom')
    btm.set(_qn('w:val'), 'single')
    btm.set(_qn('w:sz'), '4')
    btm.set(_qn('w:space'), '1')
    # Use accent color if provided, else dark gray
    _border_col = accent_color_hex if accent_color_hex else "555555"
    btm.set(_qn('w:color'), _border_col)
    pBdr.append(btm)
    pPr.append(pBdr)


def _section_heading_bordered(doc, text: str, font_name: str,
                               font_size: int, bold: bool,
                               color_hex: str, border_color: str,
                               border_sz: str = "6",
                               space_before: float = 10, space_after: float = 4,
                               prefix: str = ""):
    """Universal bordered section heading."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    p.clear()
    label = f"{prefix}{text.upper()}" if prefix else text.upper()
    run = p.add_run(label)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), border_sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), border_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _add_bullet(doc, text: str, font_size: int = 10, font_name: str = "Arial",
                indent_left: int = 360, indent_hanging: int = 180,
                color_rgb: tuple = None):
    """
    Add a properly formatted ATS-compliant bullet point paragraph.
    Uses standard hanging indent matching Jobscan/Enhancv output.
    Bullet character is plain Unicode bullet (U+2022) — universally parsed by all ATS.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    # Strip any leading bullet/dash the LLM may have prepended to avoid double-bullets
    clean_text = text.strip()
    for prefix in ("\u2022", "-", "*", "\u25aa", "\u25cf"):
        if clean_text.startswith(prefix):
            clean_text = clean_text[len(prefix):].lstrip()
            break
    p = doc.add_paragraph(style="Normal")
    p.clear()
    run = p.add_run(f"\u2022  {clean_text}")
    run.font.size = Pt(font_size)
    run.font.name = font_name
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(indent_left))
    ind.set(qn('w:hanging'), str(indent_hanging))
    pPr.append(ind)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    return p


def _add_role_line(doc, role: str, company: str, duration: str,
                   font_name: str, role_size: int = 11, meta_size: int = 9,
                   role_color: tuple = (0, 0, 0), company_color: tuple = (74, 74, 74),
                   duration_color: tuple = (128, 128, 128), separator: str = "  —  "):
    """Add role | company | duration header row for experience."""
    p = doc.add_paragraph()
    p.clear()
    if role:
        r1 = p.add_run(role)
        r1.bold = True
        r1.font.size = Pt(role_size)
        r1.font.name = font_name
        r1.font.color.rgb = RGBColor(*role_color)
    if company:
        r2 = p.add_run(f"{separator}{company}")
        r2.font.size = Pt(role_size - 1)
        r2.font.name = font_name
        r2.font.color.rgb = RGBColor(*company_color)
    if duration and duration not in ("", "[Not Provided]"):
        r3 = p.add_run(f"   [{duration}]")
        r3.italic = True
        r3.font.size = Pt(meta_size)
        r3.font.name = font_name
        r3.font.color.rgb = RGBColor(*duration_color)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def _add_project_header(doc, name: str, duration: str, tech_stack: str, url: str,
                         font_name: str, name_size: int = 11,
                         name_color: tuple = (0, 0, 0),
                         meta_color: tuple = (74, 74, 74),
                         url_color: tuple = (30, 58, 95)):
    """Add project name | duration | tech stack | URL header row."""
    p = doc.add_paragraph()
    p.clear()
    if name:
        r1 = p.add_run(name)
        r1.bold = True
        r1.font.size = Pt(name_size)
        r1.font.name = font_name
        r1.font.color.rgb = RGBColor(*name_color)
    if duration and duration not in ("", "[Not Provided]"):
        rd = p.add_run(f"   [{duration}]")
        rd.italic = True
        rd.font.size = Pt(name_size - 2)
        rd.font.name = font_name
        rd.font.color.rgb = RGBColor(128, 128, 128)
    if tech_stack and tech_stack not in ("", "[Not Provided]"):
        r2 = p.add_run(f"  |  Tech: {tech_stack}")
        r2.font.size = Pt(name_size - 2)
        r2.font.name = font_name
        r2.font.color.rgb = RGBColor(*meta_color)
    if url and url not in ("", "[Not Provided]"):
        r3 = p.add_run(f"  |  {url}")
        r3.font.size = Pt(name_size - 2)
        r3.font.name = font_name
        r3.font.color.rgb = RGBColor(*url_color)
        r3.underline = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def _add_education_row(doc, degree: str, institution: str, year: str,
                        edu_bullets: list, font_name: str,
                        degree_size: int = 10, meta_size: int = 9,
                        degree_color: tuple = (0, 0, 0),
                        inst_color: tuple = (74, 74, 74),
                        year_color: tuple = (128, 128, 128)):
    """Add degree | institution | year + optional bullets."""
    p = doc.add_paragraph()
    p.clear()
    if degree:
        r1 = p.add_run(degree)
        r1.bold = True
        r1.font.size = Pt(degree_size)
        r1.font.name = font_name
        r1.font.color.rgb = RGBColor(*degree_color)
    if institution and institution not in ("", "[Not Provided]"):
        r2 = p.add_run(f"  —  {institution}")
        r2.font.size = Pt(degree_size)
        r2.font.name = font_name
        r2.font.color.rgb = RGBColor(*inst_color)
    if year and year not in ("", "[Not Provided]"):
        r3 = p.add_run(f"  ({year})")
        r3.italic = True
        r3.font.size = Pt(meta_size)
        r3.font.name = font_name
        r3.font.color.rgb = RGBColor(*year_color)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    for b in (edu_bullets or []):
        if b and b != "[Not Provided]":
            _add_bullet(doc, b, font_size=degree_size - 1, font_name=font_name)


def _render_additional(doc, data: dict, font_name: str, font_size: int,
                        heading_fn, bullet_fn,
                        name_color_rgb: tuple = (0,0,0),
                        desc_color_rgb: tuple = (80,80,80),
                        dur_color_rgb: tuple = (128,128,128)):
    """
    Render the Additional section from structured objects.
    Each item: bold name + optional [duration] on one line, description below.
    Handles dicts, flat strings, and completely skips raw leaked JSON strings.
    """
    raw = data.get("additional", [])
    # Final safety normalisation at render time
    items = []
    for item in raw:
        if isinstance(item, dict):
            name = str(item.get("name", "") or "").strip()
            desc = str(item.get("description", "") or "").strip()
            dur  = str(item.get("duration", "") or "").strip()
            # Skip if both name and desc are empty or placeholder
            if not name and not desc:
                continue
            if name in ("[Not Provided]", "") and desc in ("[Not Provided]", ""):
                continue
            items.append({"name": name, "description": desc, "duration": dur})
        elif isinstance(item, str):
            s = item.strip()
            if not s or s == "[Not Provided]":
                continue
            # Discard any raw dict/JSON leak silently
            if s.startswith("{") or ("'name'" in s and "'description'" in s) or s.startswith("["):
                continue
            items.append({"name": s, "description": "", "duration": ""})

    if not items:
        return

    heading_fn("Additional")

    for item in items:
        name = item.get("name", "")
        desc = item.get("description", "")
        dur  = item.get("duration", "")

        # Name line: bold name + italic [duration]
        p = doc.add_paragraph()
        p.clear()
        if name and name not in ("[Not Provided]", ""):
            r1 = p.add_run(name)
            r1.bold = True
            r1.font.size = Pt(font_size)
            r1.font.name = font_name
            r1.font.color.rgb = RGBColor(*name_color_rgb)
        if dur and dur not in ("[Not Provided]", ""):
            rd = p.add_run(f"   [{dur}]")
            rd.italic = True
            rd.font.size = Pt(font_size - 1)
            rd.font.name = font_name
            rd.font.color.rgb = RGBColor(*dur_color_rgb)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)

        # Description line (if present)
        if desc and desc not in ("[Not Provided]", ""):
            pd = doc.add_paragraph()
            pd.clear()
            rd2 = pd.add_run(desc)
            rd2.font.size = Pt(font_size - 1)
            rd2.font.name = font_name
            rd2.font.color.rgb = RGBColor(*desc_color_rgb)
            pd.paragraph_format.space_before = Pt(0)
            pd.paragraph_format.space_after = Pt(2)


# ─── MODERN TEMPLATE ──────────────────────────────────────────────────────────
def generate_modern_docx(data: dict) -> BytesIO:
    """
    Modern ATS-Optimized template — single-column, Calibri font, navy headings.
    Strictly follows ATS section ordering used by Workday, Greenhouse, and Lever:
      Header → Professional Summary → Skills → Work Experience →
      Projects → Education → Certifications → Languages → Interests → Additional

    All formatting decisions prioritize machine readability over visual design.
    No tables, no columns, no text boxes — pure linear paragraph flow for ATS parsers.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    NAVY = (0x1E, 0x3A, 0x5F)
    NAVY_HEX = "1E3A5F"
    FONT = "Calibri"
    BODY = 10

    # ── ATS RULE: Single-column header block — centered plain text, no tables ──
    _build_contact_header(
        doc, data,
        name_size=20, name_color_rgb=NAVY, name_font=FONT,
        contact_font=FONT, contact_color_hex="4A4A4A", contact_size=9,
        title_font=FONT, title_size=11, title_color_rgb=NAVY,
        accent_color_hex=NAVY_HEX,
    )

    def _heading(text):
        """
        ATS-standard section heading: ALL CAPS, bold, bottom-bordered.
        Border signals section boundary to ATS parsers without using tables.
        Matches heading labels expected by Workday/Greenhouse parsers.
        """
        _section_heading_bordered(doc, text, font_name=FONT, font_size=BODY,
                                   bold=True, color_hex=NAVY_HEX,
                                   border_color=NAVY_HEX, border_sz="6",
                                   space_before=10, space_after=4)

    def _body_para(text, italic=False, color_rgb=None):
        """Standard body paragraph — consistent 10pt Calibri, minimal spacing."""
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(BODY)
        run.font.name = FONT
        run.italic = italic
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        return p

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: PROFESSIONAL SUMMARY
    # ATS parsers (Workday, iCIMS) actively scan this section for role fit.
    # 2-3 sentences: Role identity + core competencies + value proposition.
    # ══════════════════════════════════════════════════════════════════════
    if data.get("summary") and data["summary"] not in ("", "[Not Provided]"):
        _heading("Professional Summary")
        _body_para(data["summary"])

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: CORE SKILLS
    # ATS keyword scanners parse Skills first for job description matching.
    # Categorized format (Technical / Professional) with pipe-separated values
    # is the industry standard used by Jobscan, Enhancv, and Greenhouse parsers.
    # ══════════════════════════════════════════════════════════════════════
    tech_skills = [s for s in data.get("skills", []) if s and s != "[Not Provided]"]
    soft_skills = [s for s in data.get("soft_skills", []) if s and s != "[Not Provided]"]
    if tech_skills or soft_skills:
        _heading("Core Skills")
        if tech_skills:
            p = doc.add_paragraph()
            p.clear()
            label_run = p.add_run("Technical:  ")
            label_run.bold = True
            label_run.font.size = Pt(BODY)
            label_run.font.name = FONT
            label_run.font.color.rgb = RGBColor(*NAVY)
            # Group into rows of max 6 skills for readability — ATS reads all as flat text
            skills_text = "  |  ".join(tech_skills)
            skills_run = p.add_run(skills_text)
            skills_run.font.size = Pt(BODY)
            skills_run.font.name = FONT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
        if soft_skills:
            p = doc.add_paragraph()
            p.clear()
            label_run = p.add_run("Professional:  ")
            label_run.bold = True
            label_run.font.size = Pt(BODY)
            label_run.font.name = FONT
            label_run.font.color.rgb = RGBColor(*NAVY)
            ss_run = p.add_run("  |  ".join(soft_skills))
            ss_run.font.size = Pt(BODY)
            ss_run.font.name = FONT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: WORK EXPERIENCE
    # Industry-standard layout (Enhancv/Jobscan):
    #   Line 1: Job Title (bold, navy, larger)  |  Company Name
    #   Line 2: MMM YYYY – MMM YYYY  (italic, right-aligned)
    #   Line 3: 1-sentence role scope (italic gray) — optional
    #   Lines 4+: • Bullet points (action verb + task + tech + impact)
    # ATS parsers map Role Title and Company as separate parsed fields.
    # ══════════════════════════════════════════════════════════════════════
    valid_exp = [e for e in data.get("experience", [])
                 if (e.get("role") and e["role"] not in ("", "[Not Provided]"))
                 or (e.get("company") and e["company"] not in ("", "[Not Provided]"))]
    if valid_exp:
        _heading("Work Experience")
        for exp in valid_exp:
            _role     = exp.get("role", "")     if exp.get("role", "")     not in ("", "[Not Provided]") else ""
            _company  = exp.get("company", "")  if exp.get("company", "")  not in ("", "[Not Provided]") else ""
            _duration = exp.get("duration", "") if exp.get("duration", "") not in ("", "[Not Provided]") else ""

            # ── Row 1: Role Title (bold navy) — Company (regular gray) ──────
            p = doc.add_paragraph()
            p.clear()
            if _role:
                r_role = p.add_run(_role)
                r_role.bold = True
                r_role.font.size = Pt(BODY + 1)
                r_role.font.name = FONT
                r_role.font.color.rgb = RGBColor(*NAVY)
            if _company:
                sep = "  \u2014  " if _role else ""   # em-dash separator — ATS-safe
                r_co = p.add_run(f"{sep}{_company}")
                r_co.font.size = Pt(BODY)
                r_co.font.name = FONT
                r_co.font.color.rgb = RGBColor(74, 74, 74)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(0)

            # ── Row 2: Duration (italic, smaller, dark-gray) ─────────────────
            if _duration:
                p_dur = doc.add_paragraph()
                p_dur.clear()
                r_dur = p_dur.add_run(_duration)
                r_dur.italic = True
                r_dur.font.size = Pt(BODY - 1)
                r_dur.font.name = FONT
                r_dur.font.color.rgb = RGBColor(110, 110, 110)
                p_dur.paragraph_format.space_before = Pt(0)
                p_dur.paragraph_format.space_after = Pt(2)

            # ── Row 3: Role scope summary (italic gray) ────────────────────
            if exp.get("description") and exp["description"] not in ("", "[Not Provided]"):
                _body_para(exp["description"], italic=True, color_rgb=(90, 90, 90))

            # ── Rows 4+: Achievement bullets ───────────────────────────────
            for b in exp.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: PROJECTS
    # Industry-standard project entry format (Jobscan/Enhancv):
    #   Line 1: Project Name (bold navy) | [Duration] (italic gray)
    #   Line 2: Tech Stack: ... (smaller, plain) | URL (plain text — no hyperlink)
    #   Line 3: 1-sentence project purpose (italic)
    #   Lines 4+: Achievement bullets
    # ATS parsers cannot follow hyperlinks — URLs must be plain text.
    # ══════════════════════════════════════════════════════════════════════
    valid_proj = [p for p in data.get("projects", [])
                  if p.get("name") and p["name"] not in ("", "[Not Provided]")]
    if valid_proj:
        _heading("Projects")
        for proj in valid_proj:
            # Line 1: Project Name + Duration
            p = doc.add_paragraph()
            p.clear()
            r_name = p.add_run(proj.get("name", ""))
            r_name.bold = True
            r_name.font.size = Pt(BODY + 1)
            r_name.font.name = FONT
            r_name.font.color.rgb = RGBColor(*NAVY)
            if proj.get("duration") and proj["duration"] not in ("", "[Not Provided]"):
                r_dur = p.add_run(f"   [{proj['duration']}]")
                r_dur.italic = True
                r_dur.font.size = Pt(BODY - 1)
                r_dur.font.name = FONT
                r_dur.font.color.rgb = RGBColor(110, 110, 110)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            # Line 2: Tech Stack + URL (plain text, ATS-safe)
            meta_parts = []
            if proj.get("tech_stack") and proj["tech_stack"] not in ("", "[Not Provided]"):
                meta_parts.append(f"Tech: {proj['tech_stack']}")
            if proj.get("url") and proj["url"] not in ("", "[Not Provided]"):
                meta_parts.append(proj["url"])
            if meta_parts:
                p_meta = doc.add_paragraph()
                p_meta.clear()
                r_meta = p_meta.add_run("  |  ".join(meta_parts))
                r_meta.font.size = Pt(BODY - 1)
                r_meta.font.name = FONT
                r_meta.font.color.rgb = RGBColor(74, 74, 74)
                p_meta.paragraph_format.space_before = Pt(0)
                p_meta.paragraph_format.space_after = Pt(2)
            if proj.get("description") and proj["description"] not in ("", "[Not Provided]"):
                _body_para(proj["description"], italic=True, color_rgb=(80, 80, 80))
            for b in proj.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: EDUCATION
    # Two-line entry: Degree + Institution on line 1, Year on line 2.
    # ATS parsers match: "Bachelor", "Master", "B.Tech", "MBA", "Ph.D".
    # ══════════════════════════════════════════════════════════════════════
    valid_edu = [e for e in data.get("education", [])
                 if e.get("degree") or e.get("institution")]
    if valid_edu:
        _heading("Education")
        for edu in valid_edu:
            p = doc.add_paragraph()
            p.clear()
            if edu.get("degree") and edu["degree"] not in ("", "[Not Provided]"):
                r_deg = p.add_run(edu["degree"])
                r_deg.bold = True
                r_deg.font.size = Pt(BODY + 1)
                r_deg.font.name = FONT
                r_deg.font.color.rgb = RGBColor(*NAVY)
            if edu.get("institution") and edu["institution"] not in ("", "[Not Provided]"):
                r_inst = p.add_run(f"  —  {edu['institution']}")
                r_inst.font.size = Pt(BODY)
                r_inst.font.name = FONT
                r_inst.font.color.rgb = RGBColor(74, 74, 74)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            if edu.get("year") and edu["year"] not in ("", "[Not Provided]"):
                p_yr = doc.add_paragraph()
                p_yr.clear()
                r_yr = p_yr.add_run(edu["year"])
                r_yr.italic = True
                r_yr.font.size = Pt(BODY - 1)
                r_yr.font.name = FONT
                r_yr.font.color.rgb = RGBColor(110, 110, 110)
                p_yr.paragraph_format.space_before = Pt(0)
                p_yr.paragraph_format.space_after = Pt(2)
            for b in (edu.get("bullets") or []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY - 1, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: CERTIFICATIONS & LINKS
    # Each cert: Name (bold) | Issuer | Date — one bullet per credential.
    # Followed by LinkedIn / GitHub / Portfolio as plain-text URL bullets.
    # ATS systems (LinkedIn, Workday) match cert keywords directly.
    # ══════════════════════════════════════════════════════════════════════
    valid_certs = [c for c in data.get("certifications", [])
                   if isinstance(c, dict) and c.get("name") and c["name"] not in ("", "[Not Provided]")]
    contact = data.get("contact", {})
    has_links = any(contact.get(k, "") not in ("", "[Not Provided]", "Not Provided")
                    for k in ("linkedin", "github", "portfolio"))
    if valid_certs or has_links:
        _heading("Certifications & Links")
        for cert in valid_certs:
            parts = [cert["name"]]
            if cert.get("issuer") and cert["issuer"] not in ("", "[Not Provided]"):
                parts.append(cert["issuer"])
            if cert.get("duration") and cert["duration"] not in ("", "[Not Provided]"):
                parts.append(cert["duration"])
            _add_bullet(doc, "  |  ".join(parts), font_size=BODY, font_name=FONT)
        # Profile links as plain-text bullets (ATS-safe)
        for label, key in [("LinkedIn", "linkedin"), ("GitHub", "github"), ("Portfolio", "portfolio")]:
            val = contact.get(key, "")
            if val and val not in ("", "[Not Provided]", "Not Provided"):
                _add_bullet(doc, f"{label}: {val}", font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7: LANGUAGES
    # ══════════════════════════════════════════════════════════════════════
    valid_lang = [l for l in data.get("languages", []) if l and l != "[Not Provided]"]
    if valid_lang:
        _heading("Languages")
        _body_para("  |  ".join(valid_lang))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8: INTERESTS
    # ══════════════════════════════════════════════════════════════════════
    valid_int = [i for i in data.get("interests", []) if i and i != "[Not Provided]"]
    if valid_int:
        _heading("Interests")
        _body_para("  |  ".join(valid_int))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9: ADDITIONAL (Awards, Training, Volunteering, Publications)
    # ══════════════════════════════════════════════════════════════════════
    _render_additional(doc, data, font_name=FONT, font_size=BODY,
                       heading_fn=_heading,
                       bullet_fn=lambda t: _add_bullet(doc, t, font_size=BODY, font_name=FONT),
                       name_color_rgb=NAVY, desc_color_rgb=(80, 80, 80))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── MINIMAL TEMPLATE ─────────────────────────────────────────────────────────
def generate_minimal_docx(data: dict) -> BytesIO:
    """
    Minimal ATS-Optimized template — pure black/white Arial, maximum machine readability.
    Highest ATS parse accuracy of all three templates.
    Follows identical section ordering to Modern template for ATS consistency:
      Header → Professional Summary → Skills → Work Experience →
      Projects → Education → Certifications → Languages → Interests → Additional

    No color, no decoration, no graphics — every byte serves ATS keyword matching.
    Preferred by Taleo, SmartRecruiters, and legacy HRIS systems.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(1.1)
        sec.right_margin = Inches(1.1)

    FONT = "Arial"
    BODY = 10
    BLACK_HEX = "000000"
    BLACK = (0, 0, 0)
    DARK_GRAY = (60, 60, 60)
    MID_GRAY = (100, 100, 100)

    # ── ATS RULE: Plain-text header — no color, no decoration ──
    _build_contact_header(
        doc, data,
        name_size=18, name_color_rgb=BLACK, name_font=FONT,
        contact_font=FONT, contact_color_hex="333333", contact_size=9,
        title_font=FONT, title_size=10, title_color_rgb=DARK_GRAY,
    )

    def _heading(text):
        """
        Pure black bold ALL-CAPS heading with bottom rule.
        Maximum compatibility with legacy ATS parsers that strip color.
        """
        _section_heading_bordered(doc, text, font_name=FONT, font_size=BODY,
                                   bold=True, color_hex=BLACK_HEX,
                                   border_color=BLACK_HEX, border_sz="4",
                                   space_before=10, space_after=3)

    def _body_para(text, italic=False):
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(BODY)
        run.font.name = FONT
        run.italic = italic
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        return p

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: PROFESSIONAL SUMMARY
    # ══════════════════════════════════════════════════════════════════════
    if data.get("summary") and data["summary"] not in ("", "[Not Provided]"):
        _heading("Professional Summary")
        _body_para(data["summary"])

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: CORE SKILLS
    # Comma-separated — maximum ATS keyword parser compatibility.
    # Taleo, SmartRecruiters, and legacy HRIS systems parse comma lists best.
    # Labeled "Technical" and "Professional" — matches Greenhouse/Lever field names.
    # ══════════════════════════════════════════════════════════════════════
    tech_skills = [s for s in data.get("skills", []) if s and s != "[Not Provided]"]
    soft_skills = [s for s in data.get("soft_skills", []) if s and s != "[Not Provided]"]
    if tech_skills or soft_skills:
        _heading("Core Skills")
        if tech_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Technical:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT
            skills_run = p.add_run(", ".join(tech_skills))
            skills_run.font.size = Pt(BODY)
            skills_run.font.name = FONT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
        if soft_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Professional:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT
            ss_run = p.add_run(", ".join(soft_skills))
            ss_run.font.size = Pt(BODY)
            ss_run.font.name = FONT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: WORK EXPERIENCE
    # Industry-standard two-line layout for maximum ATS field recognition:
    #   Line 1: Job Title (bold, black)  |  Duration (italic, gray)
    #   Line 2: Company Name (plain)
    #   Line 3: Role scope (italic, optional)
    #   Lines 4+: Bullet achievements
    # Plain black/white — preferred by Taleo and legacy HRIS parsers.
    # ══════════════════════════════════════════════════════════════════════
    valid_exp = [e for e in data.get("experience", [])
                 if (e.get("role") and e["role"] not in ("", "[Not Provided]"))
                 or (e.get("company") and e["company"] not in ("", "[Not Provided]"))]
    if valid_exp:
        _heading("Work Experience")
        for exp in valid_exp:
            _role     = exp.get("role", "")     if exp.get("role", "")     not in ("", "[Not Provided]") else ""
            _company  = exp.get("company", "")  if exp.get("company", "")  not in ("", "[Not Provided]") else ""
            _duration = exp.get("duration", "") if exp.get("duration", "") not in ("", "[Not Provided]") else ""

            # Line 1: Job Title + Duration
            p = doc.add_paragraph()
            p.clear()
            if _role:
                r1 = p.add_run(_role)
                r1.bold = True
                r1.font.size = Pt(BODY + 1)
                r1.font.name = FONT
            if _duration:
                r2 = p.add_run(f"  |  {_duration}")
                r2.italic = True
                r2.font.size = Pt(BODY - 1)
                r2.font.name = FONT
                r2.font.color.rgb = RGBColor(*MID_GRAY)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(1)

            # Line 2: Company Name
            if _company:
                p_co = doc.add_paragraph()
                p_co.clear()
                r_co = p_co.add_run(_company)
                r_co.font.size = Pt(BODY)
                r_co.font.name = FONT
                r_co.font.color.rgb = RGBColor(*DARK_GRAY)
                p_co.paragraph_format.space_before = Pt(0)
                p_co.paragraph_format.space_after = Pt(2)

            if exp.get("description") and exp["description"] not in ("", "[Not Provided]"):
                _body_para(exp["description"], italic=True)
            for b in exp.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: PROJECTS
    # Plain black project header: Name (bold) + [Duration] + tech + URL
    # Each on its own line for maximum ATS field parsing accuracy.
    # ══════════════════════════════════════════════════════════════════════
    valid_proj = [p for p in data.get("projects", [])
                  if p.get("name") and p["name"] not in ("", "[Not Provided]")]
    if valid_proj:
        _heading("Projects")
        for proj in valid_proj:
            # Line 1: Project Name + Duration
            p = doc.add_paragraph()
            p.clear()
            r1 = p.add_run(proj.get("name", ""))
            r1.bold = True
            r1.font.size = Pt(BODY + 1)
            r1.font.name = FONT
            if proj.get("duration") and proj["duration"] not in ("", "[Not Provided]"):
                rd = p.add_run(f"   [{proj['duration']}]")
                rd.italic = True
                rd.font.size = Pt(BODY - 1)
                rd.font.name = FONT
                rd.font.color.rgb = RGBColor(*MID_GRAY)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            # Line 2: Tech Stack + URL
            meta_parts = []
            if proj.get("tech_stack") and proj["tech_stack"] not in ("", "[Not Provided]"):
                meta_parts.append(f"Tech: {proj['tech_stack']}")
            if proj.get("url") and proj["url"] not in ("", "[Not Provided]"):
                meta_parts.append(proj["url"])
            if meta_parts:
                p_meta = doc.add_paragraph()
                p_meta.clear()
                r_meta = p_meta.add_run("  |  ".join(meta_parts))
                r_meta.font.size = Pt(BODY - 1)
                r_meta.font.name = FONT
                r_meta.font.color.rgb = RGBColor(*DARK_GRAY)
                p_meta.paragraph_format.space_before = Pt(0)
                p_meta.paragraph_format.space_after = Pt(2)
            if proj.get("description") and proj["description"] not in ("", "[Not Provided]"):
                _body_para(proj["description"], italic=True)
            for b in proj.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: EDUCATION
    # Degree (bold) — Institution — Year (italic)
    # ATS parsers match: "Bachelor", "Master", "B.Tech", "MBA", "Ph.D".
    # ══════════════════════════════════════════════════════════════════════
    valid_edu = [e for e in data.get("education", [])
                 if e.get("degree") or e.get("institution")]
    if valid_edu:
        _heading("Education")
        for edu in valid_edu:
            p = doc.add_paragraph()
            p.clear()
            if edu.get("degree") and edu["degree"] not in ("", "[Not Provided]"):
                r_deg = p.add_run(edu["degree"])
                r_deg.bold = True
                r_deg.font.size = Pt(BODY + 1)
                r_deg.font.name = FONT
            if edu.get("institution") and edu["institution"] not in ("", "[Not Provided]"):
                r_inst = p.add_run(f"  —  {edu['institution']}")
                r_inst.font.size = Pt(BODY)
                r_inst.font.name = FONT
                r_inst.font.color.rgb = RGBColor(*DARK_GRAY)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            if edu.get("year") and edu["year"] not in ("", "[Not Provided]"):
                p_yr = doc.add_paragraph()
                p_yr.clear()
                r_yr = p_yr.add_run(edu["year"])
                r_yr.italic = True
                r_yr.font.size = Pt(BODY - 1)
                r_yr.font.name = FONT
                r_yr.font.color.rgb = RGBColor(*MID_GRAY)
                p_yr.paragraph_format.space_before = Pt(0)
                p_yr.paragraph_format.space_after = Pt(2)
            for b in (edu.get("bullets") or []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY - 1, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: CERTIFICATIONS & LINKS
    # Plain-text bullet per cert: Name | Issuer | Date
    # Profile URLs added as plain-text bullets — ATS-safe (no hyperlinks).
    # ══════════════════════════════════════════════════════════════════════
    valid_certs = [c for c in data.get("certifications", [])
                   if isinstance(c, dict) and c.get("name") and c["name"] not in ("", "[Not Provided]")]
    contact_min = data.get("contact", {})
    has_links_min = any(contact_min.get(k, "") not in ("", "[Not Provided]", "Not Provided")
                        for k in ("linkedin", "github", "portfolio"))
    if valid_certs or has_links_min:
        _heading("Certifications & Links")
        for cert in valid_certs:
            parts = [cert["name"]]
            if cert.get("issuer") and cert["issuer"] not in ("", "[Not Provided]"):
                parts.append(cert["issuer"])
            if cert.get("duration") and cert["duration"] not in ("", "[Not Provided]"):
                parts.append(cert["duration"])
            _add_bullet(doc, "  |  ".join(parts), font_size=BODY, font_name=FONT)
        for label, key in [("LinkedIn", "linkedin"), ("GitHub", "github"), ("Portfolio", "portfolio")]:
            val = contact_min.get(key, "")
            if val and val not in ("", "[Not Provided]", "Not Provided"):
                _add_bullet(doc, f"{label}: {val}", font_size=BODY, font_name=FONT)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7: LANGUAGES
    # ══════════════════════════════════════════════════════════════════════
    valid_lang = [l for l in data.get("languages", []) if l and l != "[Not Provided]"]
    if valid_lang:
        _heading("Languages")
        _body_para(", ".join(valid_lang))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8: INTERESTS
    # ══════════════════════════════════════════════════════════════════════
    valid_int = [i for i in data.get("interests", []) if i and i != "[Not Provided]"]
    if valid_int:
        _heading("Interests")
        _body_para(", ".join(valid_int))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9: ADDITIONAL
    # ══════════════════════════════════════════════════════════════════════
    _render_additional(doc, data, font_name=FONT, font_size=BODY,
                       heading_fn=_heading,
                       bullet_fn=lambda t: _add_bullet(doc, t, font_size=BODY, font_name=FONT),
                       name_color_rgb=BLACK, desc_color_rgb=(80, 80, 80))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── CREATIVE TEMPLATE ────────────────────────────────────────────────────────
def generate_creative_docx(data: dict) -> BytesIO:
    """
    Executive ATS-Optimized template — teal/dark-navy Calibri, polished yet ATS-safe.
    Replaces design-heavy elements (decorative symbols, multi-font mixing, ◆/▌ glyphs)
    with ATS-safe equivalents that retain visual appeal without breaking parsers.
    Identical section ordering to Modern and Minimal templates for ATS consistency:
      Header → Professional Summary → Skills → Work Experience →
      Projects → Education → Certifications → Languages → Interests → Additional

    ATS COMPLIANCE NOTES:
    - All Unicode decorative characters (◆, ▌, @) removed from section headings/bullets.
    - Single font family (Calibri) used throughout body — multi-font mixing confuses some parsers.
    - Teal color is display-only; ATS parsers read plain text, not colors.
    - Georgia used ONLY for candidate name (header) — never in parseable body sections.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.95)
        sec.right_margin = Inches(0.95)

    TEAL_HEX = "0D7377"
    DARK_HEX = "14213D"
    TEAL = (0x0D, 0x73, 0x77)
    DARK = (0x14, 0x21, 0x3D)
    FONT_NAME = "Georgia"   # Name display only — ATS ignores header fonts
    FONT_BODY = "Calibri"   # Consistent body font — required for ATS parsing
    BODY = 10

    # ── Header: Name in Georgia (display), contact in Calibri (parseable) ──
    _build_contact_header(
        doc, data,
        name_size=22, name_color_rgb=DARK, name_font=FONT_NAME,
        contact_font=FONT_BODY, contact_color_hex="444444", contact_size=9,
        title_font=FONT_BODY, title_size=11, title_color_rgb=TEAL,
        accent_color_hex=TEAL_HEX,
    )

    def _heading(text):
        """
        ATS-safe teal heading — NO prefix symbols (◆/▌ break some ATS parsers).
        Teal color preserved for human readers; ATS reads underlying plain text.
        """
        _section_heading_bordered(doc, text, font_name=FONT_BODY, font_size=BODY + 1,
                                   bold=True, color_hex=TEAL_HEX,
                                   border_color=TEAL_HEX, border_sz="6",
                                   space_before=10, space_after=4)
        # NOTE: prefix="\u258c " removed — block character disrupts ATS text extraction

    def _body_para(text, italic=False, color_rgb=None):
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(BODY)
        run.font.name = FONT_BODY
        run.italic = italic
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        return p

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: PROFESSIONAL SUMMARY
    # "Profile" renamed to "Professional Summary" — standard ATS label.
    # Some older ATS systems (Taleo) fail to map "Profile" to summary field.
    # ══════════════════════════════════════════════════════════════════════
    if data.get("summary") and data["summary"] not in ("", "[Not Provided]"):
        _heading("Professional Summary")
        _body_para(data["summary"], italic=False)  # Not italic — ATS reads italic as emphasis, not content

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: CORE SKILLS
    # Teal labels + pipe-separated values — visually distinctive, ATS-safe.
    # Georgia used ONLY for name; Calibri throughout body for ATS compatibility.
    # ══════════════════════════════════════════════════════════════════════
    tech_skills = [s for s in data.get("skills", []) if s and s != "[Not Provided]"]
    soft_skills = [s for s in data.get("soft_skills", []) if s and s != "[Not Provided]"]
    if tech_skills or soft_skills:
        _heading("Core Skills")
        if tech_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Technical:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT_BODY
            lbl.font.color.rgb = RGBColor(*TEAL)
            skills_run = p.add_run("  |  ".join(tech_skills))
            skills_run.font.size = Pt(BODY)
            skills_run.font.name = FONT_BODY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
        if soft_skills:
            p = doc.add_paragraph()
            p.clear()
            lbl = p.add_run("Professional:  ")
            lbl.bold = True
            lbl.font.size = Pt(BODY)
            lbl.font.name = FONT_BODY
            lbl.font.color.rgb = RGBColor(*TEAL)
            ss_run = p.add_run("  |  ".join(soft_skills))
            ss_run.font.size = Pt(BODY)
            ss_run.font.name = FONT_BODY
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: WORK EXPERIENCE
    # Two-line entry (Enhancv/Jobscan standard):
    #   Line 1: Job Title (bold dark-navy, 11pt)  |  Duration (italic gray)
    #   Line 2: Company Name (teal, 10pt)
    #   Line 3: Role scope (italic, optional)
    #   Lines 4+: Achievement bullets
    # ATS-safe: "@" separator replaced with " — " (em-dash).
    # ══════════════════════════════════════════════════════════════════════
    valid_exp = [e for e in data.get("experience", [])
                 if (e.get("role") and e["role"] not in ("", "[Not Provided]"))
                 or (e.get("company") and e["company"] not in ("", "[Not Provided]"))]
    if valid_exp:
        _heading("Work Experience")
        for exp in valid_exp:
            _role     = exp.get("role", "")     if exp.get("role", "")     not in ("", "[Not Provided]") else ""
            _company  = exp.get("company", "")  if exp.get("company", "")  not in ("", "[Not Provided]") else ""
            _duration = exp.get("duration", "") if exp.get("duration", "") not in ("", "[Not Provided]") else ""

            # Line 1: Job Title + Duration
            p = doc.add_paragraph()
            p.clear()
            if _role:
                r1 = p.add_run(_role)
                r1.bold = True
                r1.font.size = Pt(BODY + 1)
                r1.font.name = FONT_BODY
                r1.font.color.rgb = RGBColor(*DARK)
            if _duration:
                r3 = p.add_run(f"   {_duration}")
                r3.italic = True
                r3.font.size = Pt(BODY - 1)
                r3.font.name = FONT_BODY
                r3.font.color.rgb = RGBColor(110, 110, 110)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(1)

            # Line 2: Company Name (teal)
            if _company:
                p_co = doc.add_paragraph()
                p_co.clear()
                r_co = p_co.add_run(_company)
                r_co.font.size = Pt(BODY)
                r_co.font.name = FONT_BODY
                r_co.font.color.rgb = RGBColor(*TEAL)
                p_co.paragraph_format.space_before = Pt(0)
                p_co.paragraph_format.space_after = Pt(2)

            if exp.get("description") and exp["description"] not in ("", "[Not Provided]"):
                _body_para(exp["description"], italic=True, color_rgb=(80, 80, 80))
            for b in exp.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT_BODY)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: PROJECTS
    # Project Name (bold dark-navy, 11pt) + [Duration] (italic gray)
    # Tech Stack line (teal, 9pt) + URL as plain text (ATS-safe — no hyperlink)
    # ══════════════════════════════════════════════════════════════════════
    valid_proj = [pr for pr in data.get("projects", [])
                  if pr.get("name") and pr["name"] not in ("", "[Not Provided]")]
    if valid_proj:
        _heading("Projects")
        for proj in valid_proj:
            # Line 1: Project Name + Duration
            p = doc.add_paragraph()
            p.clear()
            r1 = p.add_run(proj.get("name", ""))
            r1.bold = True
            r1.font.size = Pt(BODY + 1)
            r1.font.name = FONT_BODY
            r1.font.color.rgb = RGBColor(*DARK)
            if proj.get("duration") and proj["duration"] not in ("", "[Not Provided]"):
                rd = p.add_run(f"   [{proj['duration']}]")
                rd.italic = True
                rd.font.size = Pt(BODY - 1)
                rd.font.name = FONT_BODY
                rd.font.color.rgb = RGBColor(110, 110, 110)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            # Line 2: Tech + URL (plain text, teal color — no underline/hyperlink)
            meta_parts = []
            if proj.get("tech_stack") and proj["tech_stack"] not in ("", "[Not Provided]"):
                meta_parts.append(f"Tech: {proj['tech_stack']}")
            if proj.get("url") and proj["url"] not in ("", "[Not Provided]"):
                meta_parts.append(proj["url"])
            if meta_parts:
                p_meta = doc.add_paragraph()
                p_meta.clear()
                r_meta = p_meta.add_run("  |  ".join(meta_parts))
                r_meta.font.size = Pt(BODY - 1)
                r_meta.font.name = FONT_BODY
                r_meta.font.color.rgb = RGBColor(*TEAL)
                p_meta.paragraph_format.space_before = Pt(0)
                p_meta.paragraph_format.space_after = Pt(2)
            if proj.get("description") and proj["description"] not in ("", "[Not Provided]"):
                _body_para(proj["description"], italic=True, color_rgb=(80, 80, 80))
            for b in proj.get("bullets", []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY, font_name=FONT_BODY)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: EDUCATION
    # Degree (bold dark-navy) — Institution (teal) — Year (italic gray)
    # ══════════════════════════════════════════════════════════════════════
    valid_edu = [e for e in data.get("education", [])
                 if e.get("degree") or e.get("institution")]
    if valid_edu:
        _heading("Education")
        for edu in valid_edu:
            p = doc.add_paragraph()
            p.clear()
            if edu.get("degree") and edu["degree"] not in ("", "[Not Provided]"):
                r1 = p.add_run(edu["degree"])
                r1.bold = True
                r1.font.size = Pt(BODY + 1)
                r1.font.name = FONT_BODY
                r1.font.color.rgb = RGBColor(*DARK)
            if edu.get("institution") and edu["institution"] not in ("", "[Not Provided]"):
                r2 = p.add_run(f"  —  {edu['institution']}")
                r2.font.size = Pt(BODY)
                r2.font.name = FONT_BODY
                r2.font.color.rgb = RGBColor(*TEAL)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            if edu.get("year") and edu["year"] not in ("", "[Not Provided]"):
                p_yr = doc.add_paragraph()
                p_yr.clear()
                r3 = p_yr.add_run(edu["year"])
                r3.italic = True
                r3.font.size = Pt(BODY - 1)
                r3.font.name = FONT_BODY
                r3.font.color.rgb = RGBColor(110, 110, 110)
                p_yr.paragraph_format.space_before = Pt(0)
                p_yr.paragraph_format.space_after = Pt(2)
            for b in (edu.get("bullets") or []):
                if b and b != "[Not Provided]":
                    _add_bullet(doc, b, font_size=BODY - 1, font_name=FONT_BODY)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: CERTIFICATIONS & LINKS
    # Cert bullets: Name | Issuer | Date  (teal heading, plain bullets)
    # Profile URLs as plain-text bullets — no hyperlinks (ATS-safe).
    # ══════════════════════════════════════════════════════════════════════
    valid_certs = [c for c in data.get("certifications", [])
                   if isinstance(c, dict) and c.get("name") and c["name"] not in ("", "[Not Provided]")]
    contact_exec = data.get("contact", {})
    has_links_exec = any(contact_exec.get(k, "") not in ("", "[Not Provided]", "Not Provided")
                         for k in ("linkedin", "github", "portfolio"))
    if valid_certs or has_links_exec:
        _heading("Certifications & Links")
        for cert in valid_certs:
            parts = [cert["name"]]
            if cert.get("issuer") and cert["issuer"] not in ("", "[Not Provided]"):
                parts.append(cert["issuer"])
            if cert.get("duration") and cert["duration"] not in ("", "[Not Provided]"):
                parts.append(cert["duration"])
            _add_bullet(doc, "  |  ".join(parts), font_size=BODY, font_name=FONT_BODY)
        for label, key in [("LinkedIn", "linkedin"), ("GitHub", "github"), ("Portfolio", "portfolio")]:
            val = contact_exec.get(key, "")
            if val and val not in ("", "[Not Provided]", "Not Provided"):
                _add_bullet(doc, f"{label}: {val}", font_size=BODY, font_name=FONT_BODY)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7: LANGUAGES
    # ══════════════════════════════════════════════════════════════════════
    valid_lang = [l for l in data.get("languages", []) if l and l != "[Not Provided]"]
    if valid_lang:
        _heading("Languages")
        _body_para("  |  ".join(valid_lang))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8: INTERESTS
    # ══════════════════════════════════════════════════════════════════════
    valid_int = [i for i in data.get("interests", []) if i and i != "[Not Provided]"]
    if valid_int:
        _heading("Interests")
        _body_para("  |  ".join(valid_int))

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9: ADDITIONAL
    # ══════════════════════════════════════════════════════════════════════
    _render_additional(doc, data, font_name=FONT_BODY, font_size=BODY,
                       heading_fn=_heading,
                       bullet_fn=lambda t: _add_bullet(doc, t, font_size=BODY, font_name=FONT_BODY),
                       name_color_rgb=DARK, desc_color_rgb=(80, 80, 80))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ============================================================
# 🔍 RESUME ANALYSIS MODULE — Job Title Suggestions (UI only)
# ============================================================
# NOTE: rewrite_text_with_llm() above is the Analysis Module.
# It generates job title suggestions + rewritten text for DISPLAY ONLY.
# It is NEVER used for DOCX generation.
# DOCX generation uses optimize_resume_to_json() + extract_resume_json() exclusively.


def rewrite_and_highlight(text, replacement_mapping, user_location):
    highlighted_text = text
    masculine_count, feminine_count = 0, 0
    detected_masculine_words, detected_feminine_words = [], []
    matched_spans = []

    masculine_words = sorted(gender_words["masculine"], key=len, reverse=True)
    feminine_words = sorted(gender_words["feminine"], key=len, reverse=True)

    def span_overlaps(start, end):
        return any(s < end and e > start for s, e in matched_spans)

    # Highlight and count masculine words
    for word in masculine_words:
        pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
        for match in pattern.finditer(highlighted_text):
            start, end = match.span()
            if span_overlaps(start, end):
                continue

            word_match = match.group(0)
            colored = f"<span style='color:blue;'>{word_match}</span>"

            # Replace word in the highlighted text
            highlighted_text = highlighted_text[:start] + colored + highlighted_text[end:]
            shift = len(colored) - len(word_match)
            matched_spans = [(s if s < start else s + shift, e if s < start else e + shift) for s, e in matched_spans]
            matched_spans.append((start, start + len(colored)))

            masculine_count += 1

            # Get sentence context and highlight
            sentence_match = re.search(r'([^.]*?\b' + re.escape(word_match) + r'\b[^.]*\.)', text, re.IGNORECASE)
            if sentence_match:
                sentence = sentence_match.group(1).strip()
                colored_sentence = re.sub(
                    rf'\b({re.escape(word_match)})\b',
                    r"<span style='color:blue;'>\1</span>",
                    sentence,
                    flags=re.IGNORECASE
                )
                detected_masculine_words.append({
                    "word": word_match,
                    "sentence": colored_sentence
                })
            break  # Only one match per word

    # Highlight and count feminine words
    for word in feminine_words:
        pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
        for match in pattern.finditer(highlighted_text):
            start, end = match.span()
            if span_overlaps(start, end):
                continue

            word_match = match.group(0)
            colored = f"<span style='color:red;'>{word_match}</span>"

            # Replace word in the highlighted text
            highlighted_text = highlighted_text[:start] + colored + highlighted_text[end:]
            shift = len(colored) - len(word_match)
            matched_spans = [(s if s < start else s + shift, e if s < start else e + shift) for s, e in matched_spans]
            matched_spans.append((start, start + len(colored)))

            feminine_count += 1

            # Get sentence context and highlight
            sentence_match = re.search(r'([^.]*?\b' + re.escape(word_match) + r'\b[^.]*\.)', text, re.IGNORECASE)
            if sentence_match:
                sentence = sentence_match.group(1).strip()
                colored_sentence = re.sub(
                    rf'\b({re.escape(word_match)})\b',
                    r"<span style='color:red;'>\1</span>",
                    sentence,
                    flags=re.IGNORECASE
                )
                detected_feminine_words.append({
                    "word": word_match,
                    "sentence": colored_sentence
                })
            break  # Only one match per word

    # ⚡ Single LLM call — returns BOTH plain-text rewrite (for UI) AND JSON (for DOCX)
    # Replaces the old rewrite_text_with_llm call which discarded the JSON half.
    rewritten_text, json_str = rewrite_and_optimize_resume(
        text,
        replacement_mapping["masculine"] | replacement_mapping["feminine"],
        user_location
    )

    # Return json_str as 7th value so the caller can reuse it directly
    # without triggering a second optimize_resume_to_json() LLM call.
    return highlighted_text, rewritten_text, masculine_count, feminine_count, detected_masculine_words, detected_feminine_words, json_str

# ✅ Enhanced Grammar evaluation using LLM with suggestions
def get_grammar_score_with_llm(text, max_score=5):
    grammar_prompt = f"""
You are a senior HR language quality specialist and professional resume reviewer with 15+ years of experience evaluating resumes for Fortune 500 companies.

Analyze the following resume text across FIVE dimensions and provide an overall language quality score:

**EVALUATION DIMENSIONS:**
1. **Grammar & Mechanics** — Correct grammar, punctuation, subject-verb agreement, tense consistency
2. **Clarity & Conciseness** — Ideas expressed directly; no filler words or redundancy
3. **Professional Tone** — Appropriate formality, no informal slang or casual phrasing
4. **Action Verb Usage** — Starts bullet points with strong, quantifiable action verbs (e.g., "Led", "Engineered", "Reduced")
5. **ATS Language Alignment** — Industry-standard terminology, keyword density, no keyword stuffing

**SCORING SCALE (out of {max_score}):**
- {max_score}: Exceptional — Flawless grammar, powerful action verbs, crystal-clear and professional throughout
- {max_score-1}: Very Good — Minor stylistic issues; highly professional and readable
- {max_score-2}: Good — Some grammar or clarity issues but largely professional and effective
- {max_score-3}: Fair — Noticeable grammar, tone, or clarity problems that could affect readability
- {max_score-4}: Poor — Multiple errors affecting professional impression; needs significant editing
- 0-1: Very Poor — Significant language issues that would cause ATS rejection or recruiter dismissal

**IMPORTANT:** Be balanced — a technically competent resume with minor grammar issues should not be harshly penalized. Focus on overall professional impression.

Return EXACTLY in this format (no extra text):

Score: <number>
Feedback: <single sentence summarizing overall language quality and tone>
Suggestions:
- <Actionable suggestion 1 with example if helpful>
- <Actionable suggestion 2 with example if helpful>
- <Actionable suggestion 3 with example if helpful>
- <Actionable suggestion 4 with example if helpful>
- <Actionable suggestion 5 with example if helpful>

---
{text}
---
"""

    response = call_llm(grammar_prompt, session=st.session_state).strip()
    score_match = re.search(r"Score:\s*(\d+)", response)
    feedback_match = re.search(r"Feedback:\s*(.+)", response)
    suggestions = re.findall(r"- (.+)", response)

    score = int(score_match.group(1)) if score_match else max(0, min(max_score, max(3, max_score - 2)))  # Generous default, clamped to max_score
    feedback = feedback_match.group(1).strip() if feedback_match else "Language quality appears adequate for professional communication."
    return score, feedback, suggestions

# ✅ Main ATS Evaluation Function
def ats_percentage_score(
    resume_text,
    job_description,
    job_title="Unknown",
    logic_profile_score=None,
    edu_weight=20,
    exp_weight=35,
    skills_weight=30,
    lang_weight=5,
    keyword_weight=10,
    format_data=None,   # ← NEW: pass pre-computed format check result
):
    import datetime

    # ⚡ MERGED: detect both domains in a single LLM call (saves 1 API call vs 2 separate detect_domain_llm calls)
    _valid_domains = [
        "Data Science", "AI/Machine Learning", "UI/UX Design", "Mobile Development",
        "Frontend Development", "Backend Development", "Full Stack Development", "Cybersecurity",
        "Cloud Engineering", "DevOps/Infrastructure", "Quality Assurance", "Game Development",
        "Blockchain Development", "Embedded Systems", "System Architecture", "Database Management",
        "Networking", "Site Reliability Engineering", "Product Management", "Project Management",
        "Business Analysis", "Technical Writing", "Digital Marketing", "E-commerce", "Fintech",
        "Healthcare Tech", "EdTech", "IoT Development", "AR/VR Development", "Technical Sales",
        "Agile Coaching", "Software Engineering"
    ]
    _domain_list = ", ".join(_valid_domains)
    _domain_prompt = f"""Classify the two texts below into professional domains.
Return ONLY this exact format on one line, nothing else:
RESUME_DOMAIN | JOB_DOMAIN

Choose each domain from ONLY this list: {_domain_list}

RESUME TEXT (first 600 chars):
{resume_text[:600]}

JOB TEXT (first 600 chars):
{job_description[:600]}
"""
    try:
        _domain_raw = call_llm(_domain_prompt, session=st.session_state).strip()
        _parts = [p.strip() for p in _domain_raw.split("|")]
        resume_domain = _parts[0] if len(_parts) == 2 and _parts[0] in _valid_domains else "Software Engineering"
        job_domain    = _parts[1] if len(_parts) == 2 and _parts[1] in _valid_domains else "Software Engineering"
    except Exception:
        resume_domain = "Software Engineering"
        job_domain    = "Software Engineering"

    similarity_score = get_domain_similarity(resume_domain, job_domain)

    # Grammar defaults — overwritten by values parsed from the ATS prompt response below
    grammar_score       = max(0, min(lang_weight, lang_weight - 2))
    grammar_feedback    = "Language quality appears adequate for professional communication."
    grammar_suggestions = []

    # ✅ Balanced domain penalty
    MAX_DOMAIN_PENALTY = 15
    domain_penalty = round((1 - similarity_score) * MAX_DOMAIN_PENALTY)

    # ✅ Optional profile score note
    logic_score_note = (
        f"\n\nOptional Note: The system also calculated a logic-based profile score of {logic_profile_score}/100 "
        f"based on resume length, experience, and skills."
        if logic_profile_score else ""
    )

    # ✅ FIXED: Stable education scoring with 2025 cutoff
    current_year = datetime.datetime.now().year

    current_month = datetime.datetime.now().month

    # ✅ UPDATED: Stable education scoring with priority degrees minimum
    prompt = f"""
You are a senior ATS (Applicant Tracking System) Evaluator and Technical Recruiter with 15+ years of experience at top-tier tech firms.
Your evaluation must be rigorous, consistent, evidence-based, and match industry-standard hiring benchmarks.

You specialize in: AI/ML, Blockchain, Cloud Computing, Data Engineering, Software Development, DevOps, and Cybersecurity roles.

═══════════════════════════════════════════════════
🎯 EVALUATION PHILOSOPHY
═══════════════════════════════════════════════════
- Score based on EVIDENCE found in the resume — not assumptions
- Reward quantified achievements (numbers, percentages, scale)
- Credit projects, GitHub, hackathons, Kaggle, open-source contributions, certifications
- Penalize vague claims without evidence ("good communication skills")
- Recognize career stage: entry-level vs senior vs lead
- Prioritize recency: skills/experience from the last 3 years matter most
- Be encouraging but calibrated: do not inflate scores without evidence

═══════════════════════════════════════════════════
📐 SCORING FRAMEWORK
═══════════════════════════════════════════════════

**🎓 Education Score ({edu_weight} points max):**

PRIORITY RULE — Minimum {int(edu_weight * 0.75)} pts for these degrees (completed OR pursuing):
  • BSc/MSc Computer Science or Mathematics
  • MCA (Master of Computer Applications)  
  • BE/BTech Computer Science or IT
  • BCA + MCA combination

DATE PARSING (STRICT — Non-negotiable):
  • End year < 2025 → ✅ COMPLETED (hardcoded cutoff)
  • End year = 2025 → ✅ COMPLETED
  • End year > 2025 → 🔄 ONGOING
  • Keywords "pursuing", "in progress", "currently enrolled" → 🔄 ONGOING
  • Keywords "graduated", "completed", "finished" → ✅ COMPLETED
  • If end year < 2025, ALWAYS mark completed regardless of text

Scoring bands:
  • {int(edu_weight * 0.90)}–{edu_weight}: Outstanding — completed highly relevant degree + exceptional academic record
  • {int(edu_weight * 0.75)}–{int(edu_weight * 0.85)}: Excellent — priority degree (completed or ongoing), good standing
  • {int(edu_weight * 0.60)}–{int(edu_weight * 0.70)}: Very Good — related STEM/technical degree
  • {int(edu_weight * 0.45)}–{int(edu_weight * 0.55)}: Good — partially related degree with transferable foundation
  • {int(edu_weight * 0.30)}–{int(edu_weight * 0.40)}: Fair — unrelated degree with relevant self-learning evidence
  • {int(edu_weight * 0.15)}–{int(edu_weight * 0.25)}: Basic — minimal or no degree information
  • 0–{int(edu_weight * 0.10)}: Insufficient — no education details at all

**💼 Experience Score ({exp_weight} points max):**

Evaluate: years of relevant experience, role seniority, domain fit, impact, leadership, quantification.

  • {int(exp_weight * 0.91)}–{exp_weight}: Exceptional — exceeds requirements; strong leadership; quantified high-impact results
  • {int(exp_weight * 0.80)}–{int(exp_weight * 0.89)}: Excellent — meets/exceeds years; strong domain fit; clear achievements
  • {int(exp_weight * 0.69)}–{int(exp_weight * 0.77)}: Very Good — adequate years; good domain fit; solid responsibilities
  • {int(exp_weight * 0.57)}–{int(exp_weight * 0.66)}: Good — reasonable experience; relevant domain; some achievements
  • {int(exp_weight * 0.43)}–{int(exp_weight * 0.54)}: Fair — some gaps but shows clear potential and transferable skills
  • {int(exp_weight * 0.29)}–{int(exp_weight * 0.40)}: Basic — limited experience but relevant direction shown
  • {int(exp_weight * 0.14)}–{int(exp_weight * 0.26)}: Entry Level — minimal experience; strong potential only
  • 0–{int(exp_weight * 0.11)}: Insufficient — major gaps; no transferable evidence

NOTE: Internships, freelance projects, and open-source contributions count as valid experience.

**🛠️ Skills Score ({skills_weight} points max):**

Match each listed skill against job description requirements. Reward:
  • Hard skills: programming languages, frameworks, tools, platforms
  • Certifications: AWS, GCP, Azure, Kubernetes, Terraform, etc.
  • Emerging skills: LLMs, GenAI, Vector DBs, Web3, MLOps, DeFi, Smart Contracts

  • {int(skills_weight * 0.93)}–{skills_weight}: Outstanding — 90%+ required skills; expert proficiency; recent hands-on usage
  • {int(skills_weight * 0.80)}–{int(skills_weight * 0.90)}: Excellent — 80%+ required skills; advanced proficiency
  • {int(skills_weight * 0.67)}–{int(skills_weight * 0.77)}: Very Good — 70%+ required skills; competent usage
  • {int(skills_weight * 0.53)}–{int(skills_weight * 0.63)}: Good — 60%+ required skills; working knowledge
  • {int(skills_weight * 0.40)}–{int(skills_weight * 0.50)}: Fair — 50%+ skills OR strong foundational skills
  • {int(skills_weight * 0.27)}–{int(skills_weight * 0.37)}: Basic — 40%+ skills; clear learning trajectory
  • {int(skills_weight * 0.13)}–{int(skills_weight * 0.23)}: Limited — 30%+ skills; self-learning evident
  • 0–{int(skills_weight * 0.10)}: Insufficient — fewer than 30% required skills

**🔑 Keyword Score ({keyword_weight} points max):**

Systematically extract ALL critical terms from the job description:
technical tools, frameworks, methodologies, role titles, industry terms, certification names.
Compare against resume. Credit synonyms and equivalent terms.

  • {int(keyword_weight * 0.90)}–{keyword_weight}: Excellent — 85%+ critical terms; strong industry vocabulary
  • {int(keyword_weight * 0.80)}: Very Good — 75%+ critical terms
  • {int(keyword_weight * 0.60)}–{int(keyword_weight * 0.70)}: Good — 65%+ critical terms
  • {int(keyword_weight * 0.40)}–{int(keyword_weight * 0.50)}: Fair — 50%+ critical terms
  • {int(keyword_weight * 0.20)}–{int(keyword_weight * 0.30)}: Basic — 35%+ critical terms
  • {int(keyword_weight * 0.10)}: Limited — 20%+ critical terms
  • 0: Poor — fewer than 20% critical terms

═══════════════════════════════════════════════════
📋 REQUIRED OUTPUT FORMAT
═══════════════════════════════════════════════════

Follow this EXACT structure. Do not skip any section:

### 🏷️ Candidate Name
<Copy the candidate's full name EXACTLY as it appears in the resume — character by character. Do NOT correct spelling, do NOT infer from context, do NOT paraphrase. Look at the very top of the resume (header/contact section). Output ONLY the name, nothing else. If you cannot find a name, write: Not Found>

### 🏫 Education Analysis
**Score:** <0–{edu_weight}> / {edu_weight}

**Scoring Rationale:**
- Degree Level & Relevance: <Does it qualify for minimum {int(edu_weight * 0.75)}-pt rule? Which degree?>
- Completion Status: <Apply strict 2025 cutoff rule; state year and final status>
- Academic Quality Indicators: <GPA, honors, relevant coursework if mentioned>
- **Score Justification:** <Explain exact score with evidence from resume>

### 💼 Experience Analysis
**Score:** <0–{exp_weight}> / {exp_weight}

**Experience Breakdown:**
- Total Years of Relevant Experience: <X years — include internships, freelance, open-source>
- Role Progression & Seniority: <Entry → Mid → Senior trajectory>
- Domain Alignment: <How well does background match job domain?>
- Quantified Achievements: <List metrics found: % improvement, $ savings, users served, etc.>
- Leadership & Ownership Evidence: <Managed teams? Led projects? Mentored?>
- Technology Currency: <Are skills/tools recent and relevant (last 3 years)?>
- **Score Justification:** <Explain score with specific resume evidence>

### 🛠 Skills Analysis
**Score:** <0–{skills_weight}> / {skills_weight}

**Skills Assessment:**
- Core Technical Skills Matched: <List matched skills with evidence>
- Emerging/Cutting-Edge Skills: <LLMs, GenAI, Web3, MLOps, Cloud, etc.>
- Certifications Detected: <List any certifications found>
- Soft Skills with Evidence: <Only count if backed by concrete examples>
- Proficiency Depth: <Surface knowledge vs. demonstrated project usage>

**Skills Gaps (Development Opportunities):**
- <Gap 1 — specific missing skill from job description>
- <Gap 2 — specific missing skill>
- <Gap 3 — specific missing skill>
- <Gap 4 — specific missing skill>
- <Gap 5 — specific missing skill>

**Score Justification:** <Explain with matched vs. required skills ratio>

### 🗣 Language Quality Analysis
**Score:** <evaluate and provide a score 0–{lang_weight}> / {lang_weight}
**Grammar & Professional Tone:** <single sentence summarising overall language quality>
**Suggestions:**
- <Actionable suggestion 1>
- <Actionable suggestion 2>
- <Actionable suggestion 3>
- <Actionable suggestion 4>
- <Actionable suggestion 5>
**Assessment:** <Specific feedback on action verb usage, clarity, tense consistency, and ATS language>

SCORING SCALE for language ({lang_weight} pts max):
- {lang_weight}: Exceptional — Flawless grammar, powerful action verbs, crystal-clear and professional throughout
- {lang_weight-1}: Very Good — Minor stylistic issues; highly professional and readable
- {lang_weight-2}: Good — Some grammar or clarity issues but largely professional
- {lang_weight-3}: Fair — Noticeable grammar or clarity problems
- 0-1: Poor — Significant language issues

### 🔑 Keyword Analysis
**Score:** <0–{keyword_weight}> / {keyword_weight}

**Keyword Assessment:**
- Industry Terminology Match: <Percentage and specific matches found>
- Role-Specific Keywords Present: <List matched keywords>
- Technical Vocabulary: <Tools, frameworks, platforms found in both>
- Keyword Density Quality: <Natural integration vs. stuffing>

**Keyword Enhancement Opportunities:**
- <Critical keyword 1 from job description — not in resume>
- <Critical keyword 2>
- <Critical keyword 3>
- <Critical keyword 4>
- <Critical keyword 5>
- <Critical keyword 6>
- <Critical keyword 7>
- <Critical keyword 8>

**Score Justification:** <Evidence-based explanation>

### 📐 Format & ATS Compatibility Analysis
**Format Score:** {format_data.get("format_score", "N/A") if format_data else "N/A"} / 100  
**Format Grade:** {format_data.get("letter_grade", "N/A") if format_data else "N/A"} — {format_data.get("label", "") if format_data else ""}

⚠️ IMPORTANT: The Format Score and Format Grade above are SYSTEM-COMPUTED and LOCKED. Do NOT change these numbers. Only fill in the narrative fields below.

**Structural Assessment:**
- Section Completeness: <narrative only — do NOT include a score>
- Contact Block: <narrative only>
- Resume Length: {f"{format_data.get('word_count', 'N/A')} words — " + ("Optimal" if 300 <= (format_data.get('word_count') or 0) <= 1000 else "Too short" if (format_data.get('word_count') or 0) < 300 else "Too long") if format_data else "N/A"}
- Action Verb Strength: <narrative only>
- Quantification Quality: <narrative only>
- ATS Red Flags: <narrative only>

**Format Issues Detected:**
{chr(10).join(f"- {issue}" for issue in (format_data.get("issues", []) or ["No issues detected"])) if format_data else "- Format data not available"}

**Format Strengths:**
{chr(10).join(f"- {p}" for p in (format_data.get("passes", []) or ["No specific passes noted"])) if format_data else "- Format data not available"}

**Improvement Recommendations:**
- <Top format fix 1 — specific and actionable>
- <Top format fix 2>
- <Top format fix 3>

### ✅ Final Assessment

**Overall Evaluation:**
<5–7 sentences covering: candidate's unique value proposition, strongest evidence-backed qualifications, key gaps, culture/team fit signals, and a clear hire/interview recommendation>

**Top 3 Strengths (with evidence):**
1. <Strength 1 — backed by resume evidence>
2. <Strength 2 — backed by resume evidence>
3. <Strength 3 — backed by resume evidence>

**Top 3 Development Areas:**
1. <Gap 1 framed as a growth opportunity>
2. <Gap 2 framed as a growth opportunity>
3. <Gap 3 framed as a growth opportunity>

**Hiring Recommendation:** <Strongly Recommend / Recommend / Recommend with Reservations / Do Not Recommend> — <2-sentence reasoning>

---

**EVALUATION CONTEXT:**
- Current Date: {datetime.datetime.now().strftime('%B %Y')} (Year: {current_year}, Month: {current_month})
- Resume Domain Detected: {resume_domain}
- Target Job Domain: {job_domain}
- Domain Similarity Score: {similarity_score:.2f}/1.0
- Domain Mismatch Penalty Applied: {domain_penalty}/{MAX_DOMAIN_PENALTY} pts

---

📄 **JOB DESCRIPTION:**
{job_description}

📄 **RESUME TEXT:**
{resume_text}

{logic_score_note}
"""
   
   
    ats_result = call_llm(prompt, session=st.session_state).strip()

    # ── CRITICAL: Overwrite any LLM-modified Format Score/Grade lines ────
    # The LLM sometimes rewrites these despite instructions. Force the true
    # system-computed values back in so UI and narrative always match.
    _true_fmt_score = format_data.get("format_score", 75) if format_data else 75
    _true_fmt_grade = format_data.get("letter_grade", "N/A") if format_data else "N/A"
    _true_fmt_label = format_data.get("label", "") if format_data else ""

    ats_result = re.sub(
        r'\*\*Format Score:\*\*.*',
        f'**Format Score:** {_true_fmt_score} / 100',
        ats_result
    )
    ats_result = re.sub(
        r'\*\*Format Grade:\*\*.*',
        f'**Format Grade:** {_true_fmt_grade} — {_true_fmt_label}',
        ats_result
    )
    # ─────────────────────────────────────────────────────────────────────

    def extract_section(pattern, text, default="N/A"):
        match = re.search(pattern, text, re.DOTALL)
        return match.group(1).strip() if match else default

    def extract_score(pattern, text, default=0):
        match = re.search(pattern, text)
        return int(match.group(1)) if match else default

    # Extract key sections
    _raw_name = extract_section(r"### 🏷️ Candidate Name(.*?)###", ats_result, "")
    candidate_name = re.sub(r"[*_`#\[\]<>]", "", _raw_name).strip()
    candidate_name = " ".join(candidate_name.split())
    _placeholder_values = {
        "not found", "n/a", "unknown", "none", "",
        "extract full name from resume header or contact section",
        "copy the candidate's full name exactly as it appears in the resume",
        "copy the candidates full name exactly as it appears in the resume",
        "name not found", "candidate name not found",
    }
    if candidate_name.lower() in _placeholder_values:
        candidate_name = "Not Found"
    edu_analysis = extract_section(r"### 🏫 Education Analysis(.*?)###", ats_result)
    exp_analysis = extract_section(r"### 💼 Experience Analysis(.*?)###", ats_result)
    skills_analysis = extract_section(r"### 🛠 Skills Analysis(.*?)###", ats_result)
    lang_analysis = extract_section(r"### 🗣 Language Quality Analysis(.*?)###", ats_result)
    keyword_analysis = extract_section(r"### 🔑 Keyword Analysis(.*?)###", ats_result)
    format_analysis = extract_section(r"### 📐 Format & ATS Compatibility Analysis(.*?)###", ats_result)
    final_thoughts = extract_section(r"### ✅ Final Assessment(.*)", ats_result)

    # Extract scores with improved patterns (LLM now scores directly using sidebar weights)
    edu_score     = extract_score(r"\*\*Score:\*\*\s*(\d+)", edu_analysis)
    exp_score     = extract_score(r"\*\*Score:\*\*\s*(\d+)", exp_analysis)
    skills_score  = extract_score(r"\*\*Score:\*\*\s*(\d+)", skills_analysis)
    keyword_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", keyword_analysis)
    # ⚡ Parse grammar score + feedback from ATS result (no separate LLM call needed)
    _grammar_score_match    = re.search(r"\*\*Score:\*\*\s*<evaluate.*?(\d+)>|Score.*?(\d+)\s*/\s*" + str(lang_weight), lang_analysis)
    _grammar_score_match2   = re.search(r"\*\*Score:\*\*\s*(\d+)", lang_analysis)
    _grammar_feedback_match = re.search(r"\*\*Grammar & Professional Tone:\*\*\s*(.+)", lang_analysis)
    _grammar_sugg_raw       = re.findall(r"^- (.+)", lang_analysis, re.MULTILINE)

    if _grammar_score_match2:
        grammar_score = int(_grammar_score_match2.group(1))
    # else keep the safe default already set above

    if _grammar_feedback_match:
        grammar_feedback = _grammar_feedback_match.group(1).strip()

    if _grammar_sugg_raw:
        grammar_suggestions = _grammar_sugg_raw

    lang_score = grammar_score  # use value parsed from ATS result

    # ── Clamp every score: min floor + hard upper cap to its own weight ──
    # Upper cap prevents LLM hallucinating over-max scores (e.g. 25/20)
    # which would silently push content_score above 100.
    edu_score     = max(int(edu_weight * 0.15),     min(edu_score,     edu_weight))
    exp_score     = max(int(exp_weight * 0.15),     min(exp_score,     exp_weight))
    skills_score  = max(int(skills_weight * 0.15),  min(skills_score,  skills_weight))
    keyword_score = max(int(keyword_weight * 0.10), min(keyword_score, keyword_weight))
    lang_score    = max(0,                          min(lang_score,    lang_weight))

    # Extract missing items with better parsing - now called "opportunities"
    missing_keywords_section = extract_section(r"\*\*Keyword Enhancement Opportunities:\*\*(.*?)(?:\*\*|###|\Z)", keyword_analysis)
    missing_skills_section = extract_section(r"\*\*Skills Gaps \(Development Opportunities\):\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)

    # Fallback to old patterns if new ones don't match
    if not missing_keywords_section.strip():
        missing_keywords_section = extract_section(r"\*\*Missing Critical Keywords:\*\*(.*?)(?:\*\*|###|\Z)", keyword_analysis)
    if not missing_skills_section.strip():
        missing_skills_section = extract_section(r"\*\*Skills Gaps \(Opportunities for Growth\):\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)
    if not missing_skills_section.strip():
        missing_skills_section = extract_section(r"\*\*Missing Critical Skills:\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)
    
    # Improved extraction - handle multiple formats and get all items
    def extract_list_items(text):
        if not text.strip():
            return "None identified"
        
        # Find all bullet points with various formats
        items = []
        lines = text.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Remove various bullet point formats
            cleaned_line = re.sub(r'^[-•*]\s*', '', line)  # Remove -, •, * bullets
            cleaned_line = re.sub(r'^\d+\.\s*', '', cleaned_line)  # Remove numbered lists
            cleaned_line = cleaned_line.strip()
            
            if cleaned_line and len(cleaned_line) > 2:  # Avoid empty or very short items
                items.append(cleaned_line)
        
        return ', '.join(items) if items else "None identified"
    
    missing_keywords = extract_list_items(missing_keywords_section)
    missing_skills = extract_list_items(missing_skills_section)

    # ── Score assembly — fully deterministic integer arithmetic ──────────
    # Step 1: sum the five LLM-scored components (clamped to their individual weights)
    content_score = edu_score + exp_score + skills_score + lang_score + keyword_score

    # Normalise LLM components to 90-pt scale (format takes the remaining 10 pts)
    # This keeps total = 100 while giving format meaningful, visible weight.
    weight_total = edu_weight + exp_weight + skills_weight + lang_weight + keyword_weight
    if weight_total > 0:
        content_score = round(content_score / weight_total * 90)
    content_score = max(0, min(90, content_score))

    # Step 2: format score (0–100) contributes a fixed 10-pt component
    # Scaled proportionally: 100 format → 10 pts, 0 format → 0 pts
    fmt_score_raw = format_data.get("format_score", 75) if format_data else 75
    fmt_score_raw = max(0, min(100, int(fmt_score_raw)))
    FORMAT_WEIGHT = 10
    format_component = round(fmt_score_raw / 100 * FORMAT_WEIGHT)

    # Step 3: combine content + format → pre-penalty total (0–100)
    pre_penalty_score = content_score + format_component
    pre_penalty_score = max(0, min(100, pre_penalty_score))

    # Step 4: subtract domain mismatch penalty ONCE — straight subtraction
    total_score = pre_penalty_score - domain_penalty

    # Step 5: clamp final result 15–100
    total_score = max(15, min(100, total_score))

    # ✅ Industry-standard score labels with clear hiring signal
    formatted_score = (
        "Exceptional Match — Top 10% Candidate"    if total_score >= 85 else
        "Strong Match — Recommend for Interview"    if total_score >= 70 else
        "Good Potential — Competitive Candidate"    if total_score >= 55 else
        "Fair Match — Needs Resume Optimization"    if total_score >= 40 else
        "Developing — Significant Skill Gaps"       if total_score >= 25 else
        "Poor Match — Major Role Misalignment"
    )

    # ✅ Format suggestions nicely
    suggestions_html = ""
    if grammar_suggestions:
        suggestions_html = "<ul>" + "".join([f"<li>{s}</li>" for s in grammar_suggestions]) + "</ul>"

    updated_lang_analysis = f"""
{lang_analysis}
<br><b>LLM Feedback Summary:</b> {grammar_feedback}
<br><b>Improvement Suggestions:</b> {suggestions_html}
"""

    # Enhanced final thoughts with domain analysis and industry benchmarks
    final_thoughts += f"""

**Technical Evaluation Details:**
- Content Score (LLM components, 90-pt scale): {content_score}/90
- Format Component (10-pt scale): {format_component}/10 (Format Score: {fmt_score_raw}/100)
- Pre-Penalty Score: {pre_penalty_score}/100
- Domain Penalty Applied: -{domain_penalty} pts (out of max -{MAX_DOMAIN_PENALTY} pts)
- Final ATS Score: {total_score}/100
- Domain Similarity: {similarity_score:.2f}/1.0 ({int(similarity_score * 100)}% alignment)
- Resume Domain Detected: {resume_domain}
- Target Job Domain: {job_domain}
- Language Pre-Score: {grammar_score}/{lang_weight}

**Score Interpretation (Industry Benchmarks):**
- 85–100: Top 10% candidates — Strong interview recommendation
- 70–84: Above average — Likely to advance past ATS screening
- 55–69: Competitive — May advance with strong cover letter
- 40–54: Below average — Needs resume optimization before applying
- 25–39: Significant gaps — Upskilling recommended
- 0–24: Major misalignment — Not suitable for this specific role

**ATS Scoring Notes:**
- Scoring model: LLM components (90 pts) + Format (10 pts) − Domain penalty
- Format score is a real 10-pt component (not a delta) — poor formatting meaningfully lowers the score
- Domain penalty subtracted once as a flat deduction (max {MAX_DOMAIN_PENALTY} pts)
- Format checker v2: uses PDF block-coordinate multi-column detection, tiered deductions, bonus credits
- Transferable skills, projects, and open-source contributions were credited
- Career stage (entry/mid/senior) considered in experience scoring
"""

    return ats_result, {
        "Candidate Name": candidate_name,
        "Education Score": edu_score,
        "Experience Score": exp_score,
        "Skills Score": skills_score,
        "Language Score": lang_score,
        "Keyword Score": keyword_score,
        "Format Score": fmt_score_raw,
        "Format Grade": format_data.get("letter_grade", "N/A") if format_data else "N/A",
        "Format Label": format_data.get("label", "") if format_data else "",
        "Format Issues": format_data.get("issues", []) if format_data else [],
        "Format Passes": format_data.get("passes", []) if format_data else [],
        "ATS Match %": total_score,
        "Formatted Score": formatted_score,
        "Education Analysis": edu_analysis,
        "Experience Analysis": exp_analysis,
        "Skills Analysis": skills_analysis,
        "Language Analysis": updated_lang_analysis,
        "Keyword Analysis": keyword_analysis,
        "Format Analysis": format_analysis,
        "Final Thoughts": final_thoughts,
        "Missing Keywords": missing_keywords,
        "Missing Skills": missing_skills,
        "Resume Domain": resume_domain,
        "Job Domain": job_domain,
        "Domain Penalty": domain_penalty,
        "Domain Similarity Score": similarity_score
    }

# Setup Vector DB
def setup_vectorstore(documents):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    if DEVICE == "cuda":
        embeddings.model = embeddings.model.to(torch.device("cuda"))
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    doc_chunks = text_splitter.split_text("\n".join(documents))
    return FAISS.from_texts(doc_chunks, embeddings)

# Create Conversational Chain
def create_chain(vectorstore):
    # ✅ Use get_healthy_keys() so dead/quota keys are skipped (reads key_failures
    #    and key_usage from Supabase — same tables that call_llm() maintains).
    all_keys    = load_groq_api_keys()
    healthy     = get_healthy_keys(all_keys)
    if not healthy:
        raise ValueError("❌ No healthy Groq API keys available for chat chain.")
    # healthy list is already shuffled by get_healthy_keys — just take the first
    groq_api_key = healthy[0]
    increment_key_usage(groq_api_key)   # keep usage count in sync with call_llm

    # ✅ Create the ChatGroq object
    llm = ChatGroq(model="llama-3.3-70b-versatile", temperature=0, groq_api_key=groq_api_key)

    # ✅ Build the chain — report failures back so llm_manager skips this key next time
    try:
        chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=vectorstore.as_retriever(),
            return_source_documents=True
        )
        return chain
    except Exception as e:
        reason = "quota" if any(w in str(e).lower() for w in ["quota", "rate limit", "429"]) else "error"
        mark_key_failure(groq_api_key, reason)
        raise

# Chat history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# ---------------- Sidebar Layout with Inline Images ----------------
st.sidebar.markdown("""
<p style='
    font-size: 0.72rem;
    font-weight: 700;
    letter-spacing: 0.10em;
    text-transform: uppercase;
    color: #4a5568;
    border-bottom: 1px solid rgba(255,255,255,0.06);
    padding-bottom: 8px;
    margin-bottom: 12px;
    font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
'>🏷️ Job Configuration</p>
""", unsafe_allow_html=True)

# ---------------- Job Information Dropdown ----------------
with st.sidebar.expander("![Job](https://img.icons8.com/ios-filled/20/briefcase.png) Enter Job Details", expanded=False):
    job_title = st.text_input(
        "![Job](https://img.icons8.com/ios-filled/20/briefcase.png) Job Title"
    )

    user_location = st.text_input(
        "![Location](https://img.icons8.com/ios-filled/20/marker.png) Preferred Job Location (City, Country)"
    )

    job_description = st.text_area(
        "![Description](https://img.icons8.com/ios-filled/20/document.png) Paste Job Description",
        height=200
    )

    if job_description.strip() == "":
        st.warning("Please enter a job description to evaluate the resumes.")

# ---------------- Advanced Weights Dropdown ----------------
with st.sidebar.expander("![Settings](https://img.icons8.com/ios-filled/20/settings.png) Customize ATS Scoring Weights", expanded=False):
    st.markdown(
        "<div style='font-size:0.72rem;color:#64748b;margin-bottom:8px;font-family:-apple-system,sans-serif;'>"
        "Format quality is scored automatically (10 pts fixed). "
        "Adjust the remaining <b>90 pts</b> below.</div>",
        unsafe_allow_html=True
    )
    edu_weight = st.slider("![Education](https://img.icons8.com/ios-filled/20/graduation-cap.png) Education Weight", 0, 50, 20)
    exp_weight = st.slider("![Experience](https://img.icons8.com/ios-filled/20/portfolio.png) Experience Weight", 0, 50, 35)
    skills_weight = st.slider("![Skills](https://img.icons8.com/ios-filled/20/gear.png) Skills Match Weight", 0, 50, 20)
    lang_weight = st.slider("![Language](https://img.icons8.com/ios-filled/20/language.png) Language Quality Weight", 0, 10, 5)
    keyword_weight = st.slider("![Keyword](https://img.icons8.com/ios-filled/20/key.png) Keyword Match Weight", 0, 20, 10)

    total_weight = edu_weight + exp_weight + skills_weight + lang_weight + keyword_weight

    # ---------------- Inline SVG Validation ----------------
    if total_weight != 90:
        st.markdown(
            f"""
            <div style="display:flex;align-items:center;gap:8px;
                        border:1px solid rgba(251,113,133,0.3);
                        background:linear-gradient(135deg,rgba(251,113,133,0.12) 0%,rgba(251,113,133,0.05) 100%);
                        padding:10px 14px;
                        border-radius:10px;
                        backdrop-filter:blur(12px);">
                <svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" fill="#fb7185" viewBox="0 0 24 24">
                    <path d="M12 2C6.48 2 2 6.48 2 12s4.48 10 10 10
                             10-4.48 10-10S17.52 2 12 2zm0 15
                             c-.83 0-1.5.67-1.5 1.5S11.17 20
                             12 20s1.5-.67 1.5-1.5S12.83 17
                             12 17zm1-4V7h-2v6h2z"/>
                </svg>
                <span style="color:#fca5a5;font-weight:600;font-size:0.8rem;font-family:-apple-system,sans-serif;">
                    Total = {total_weight}. Adjust to exactly 90 (Format = 10 pts fixed).
                </span>
            </div>
            """,
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            f"""
            <div style="display:flex;align-items:center;gap:8px;
                        border:1px solid rgba(52,211,153,0.28);
                        background:linear-gradient(135deg,rgba(52,211,153,0.12) 0%,rgba(52,211,153,0.05) 100%);
                        padding:10px 14px;
                        border-radius:10px;
                        backdrop-filter:blur(12px);">
                <svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" fill="#34d399" viewBox="0 0 24 24">
                    <path d="M9 16.2l-3.5-3.5-1.4 1.4L9
                             19 20.3 7.7l-1.4-1.4z"/>
                </svg>
                <span style="color:#6ee7b7;font-weight:600;font-size:0.8rem;font-family:-apple-system,sans-serif;">
                    Weights balanced · Content = 90 pts · Format = 10 pts · Total = 100
                </span>
            </div>
            """,
            unsafe_allow_html=True
        )

with tab1:
    # Slide message styles already defined in global CSS — no extra block needed

    uploaded_files = st.file_uploader(
        "📄 Upload PDF Resumes",
        type=["pdf"],
        accept_multiple_files=True,
        help="Upload one or more resumes in PDF format (max 200MB each)."
    )

    if uploaded_files:
        for uploaded_file in uploaded_files:
            with st.container():
                st.subheader(f"📄 Original Resume Preview: {uploaded_file.name}")

                try:
                    # ✅ Show PDF preview safely
                    pdf_viewer(
                        uploaded_file.read(),
                        key=f"pdf_viewer_{uploaded_file.name}"
                    )

                    # Reset pointer so file can be read again later
                    uploaded_file.seek(0)

                    # ✅ Extract text safely
                    resume_text = safe_extract_text(uploaded_file)

                    if resume_text:
                        st.markdown(f"""
                        <div class='slide-message success-msg'>
                            <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
                              stroke-width="2" viewBox="0 0 24 24"><path d="M5 13l4 4L19 7"/></svg>
                            ✅ Successfully processed <b>{uploaded_file.name}</b>
                        </div>
                        """, unsafe_allow_html=True)
                        # 🔹 Continue with ATS scoring, bias detection, etc. here
                    else:
                        st.markdown(f"""
                        <div class='slide-message warn-msg'>
                            ⚠️ <b>{uploaded_file.name}</b> does not contain valid resume text.
                        </div>
                        """, unsafe_allow_html=True)

                except Exception as e:
                    st.markdown(f"""
                    <div class='slide-message error-msg'>
                        ❌ Could not display or process <b>{uploaded_file.name}</b>: {e}
                    </div>
                    """, unsafe_allow_html=True)

# ✅ Initialize state
# Initialize session state
if "resume_data" not in st.session_state:
    st.session_state.resume_data = []

if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

resume_data = st.session_state.resume_data

# ✏️ Resume Evaluation Logic
if uploaded_files and job_description:
    all_text = []

    for uploaded_file in uploaded_files:
        if uploaded_file.name in st.session_state.processed_files:
            continue

        # ✅ Improved optimized scanner animation with better performance
        scanner_placeholder = st.empty()

        # ✅ IMPROVED: More efficient CSS animations with GPU acceleration
        OPTIMIZED_SCANNER_HTML = f"""
        <style>
        .scanner-overlay {{
            position: fixed;
            top: 0; left: 0;
            width: 100vw; height: 100vh;
            background: linear-gradient(135deg, #0b0c10 0%, #1a1c29 100%);
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            z-index: 9999;
            will-change: transform, opacity;
        }}
        
        .scanner-doc {{
            width: 280px;
            height: 340px;
            background: linear-gradient(145deg, #f8f9fa, #e9ecef);
            border-radius: 16px;
            position: relative;
            overflow: hidden;
            box-shadow: 0 20px 40px rgba(0, 191, 255, 0.3);
            transform: translateZ(0);
            will-change: transform;
            animation: docFloat 3s ease-in-out infinite alternate;
        }}
        
        @keyframes docFloat {{
            0% {{ transform: translateY(0px) scale(1); }}
            100% {{ transform: translateY(-8px) scale(1.02); }}
        }}
        
        .doc-header {{
            padding: 20px;
            text-align: center;
            border-bottom: 2px solid #e9ecef;
        }}
        
        .doc-avatar {{
            width: 50px;
            height: 50px;
            background: linear-gradient(135deg, #667eea, #764ba2);
            border-radius: 50%;
            margin: 0 auto 10px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 20px;
            color: white;
        }}
        
        .doc-title {{
            font-size: 16px;
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 5px;
            font-family: 'Segoe UI', sans-serif;
        }}
        
        .doc-content {{
            padding: 15px;
            font-size: 12px;
            color: #6c757d;
            line-height: 1.4;
        }}
        
        .scan-line {{
            position: absolute;
            top: 0; left: 0;
            width: 100%; height: 4px;
            background: linear-gradient(90deg, transparent, rgba(0,191,255,0.8), transparent);
            animation: scanMove 2.5s ease-in-out infinite;
            box-shadow: 0 0 20px rgba(0,191,255,0.6);
            transform: translateZ(0);
            will-change: transform;
        }}
        
        @keyframes scanMove {{
            0% {{ top: 0; opacity: 1; }}
            50% {{ opacity: 0.8; }}
            100% {{ top: 340px; opacity: 1; }}
        }}
        
        .scanner-text {{
            margin-top: 30px;
            font-family: 'Orbitron', 'Segoe UI', sans-serif;
            font-weight: 600;
            font-size: 18px;
            color: #00bfff;
            text-shadow: 0 0 10px rgba(0,191,255,0.5);
            animation: textPulse 2s ease-in-out infinite;
        }}
        
        @keyframes textPulse {{
            0%, 100% {{ opacity: 1; transform: scale(1); }}
            50% {{ opacity: 0.8; transform: scale(1.05); }}
        }}
        
        .progress-bar {{
            width: 200px;
            height: 4px;
            background: rgba(255,255,255,0.2);
            border-radius: 2px;
            margin-top: 20px;
            overflow: hidden;
        }}
        
        .progress-fill {{
            height: 100%;
            background: linear-gradient(90deg, #00bfff, #1e90ff);
            border-radius: 2px;
            animation: progressFill 3s ease-in-out infinite;
            transform: translateX(-100%);
        }}
        
        @keyframes progressFill {{
            0% {{ transform: translateX(-100%); }}
            100% {{ transform: translateX(0); }}
        }}
        
        /* Mobile optimizations */
        @media (max-width: 768px) {{
            .scanner-doc {{ width: 240px; height: 300px; }}
            .scanner-text {{ font-size: 16px; }}
        }}
        </style>
        
        <div class="scanner-overlay">
            <div class="scanner-doc">
                <div class="scan-line"></div>
                <div class="doc-header">
                    <div class="doc-avatar">👤</div>
                    <div class="doc-title">{job_title}</div>
                </div>
                <div class="doc-content">
                    • Analyzing candidate profile...<br>
                    • Extracting key skills...<br>
                    • Matching with job requirements...<br>
                    • Calculating ATS compatibility...<br>
                    • Checking for bias patterns...
                </div>
            </div>
            <div class="scanner-text">Scanning Resume...</div>
            <div class="progress-bar">
                <div class="progress-fill"></div>
            </div>
        </div>
        """
        
        scanner_placeholder.markdown(OPTIMIZED_SCANNER_HTML, unsafe_allow_html=True)

        # ✅ Save uploaded file
        file_path = os.path.join(working_dir, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # ✅ Reduced delay for better UX
        time.sleep(4)

        # ✅ Extract text from PDF
        text = extract_text_from_pdf(file_path)
        if not text:
            st.warning(f"⚠️ Could not extract text from {uploaded_file.name}. Skipping.")
            scanner_placeholder.empty()
            continue

        all_text.append(" ".join(text))
        full_text = " ".join(text)

        # ✅ Bias detection
        bias_score, masc_count, fem_count, detected_masc, detected_fem = detect_bias(full_text)

        # ⚡ MERGED: single LLM call returns BOTH plain-text rewrite AND JSON.
        # json_str is the structured JSON for DOCX generation — no second call needed.
        with st.spinner("✍️ Rewriting resume & generating optimised structure..."):
            try:
                highlighted_text, rewritten_text, _, _, _, _, json_str = rewrite_and_highlight(
                    full_text, replacement_mapping, user_location
                )
            except Exception:
                highlighted_text = full_text
                rewritten_text   = full_text
                json_str         = ""

        # ✅ Resume Optimization Module — reuse JSON already produced above (0 extra LLM calls)
        try:
            optimized_resume_data = extract_resume_json(json_str)
        except Exception:
            optimized_resume_data = extract_resume_json("")  # falls back to empty skeleton

        # ✅ Format check (industry standard — no LLM call)
        try:
            doc_check = fitz.open(file_path)
            num_pages = doc_check.page_count
            doc_check.close()
        except Exception:
            num_pages = 1
        format_data = check_resume_format(full_text, num_pages, pdf_path=file_path)

        # ✅ LLM-based ATS Evaluation (includes domain detection + grammar scoring internally)
        with st.spinner("🔍 Running ATS evaluation..."):
            ats_result, ats_scores = ats_percentage_score(
                resume_text=full_text,
                job_description=job_description,
                logic_profile_score=None,
                edu_weight=edu_weight,
                exp_weight=exp_weight,
                skills_weight=skills_weight,
                lang_weight=lang_weight,
                keyword_weight=keyword_weight,
                format_data=format_data,
            )

        # ✅ Extract structured ATS values
        candidate_name = ats_scores.get("Candidate Name", "Not Found")

        # ── Filename-based name fallback (reliable ground truth) ─────────────
        def _name_from_filename(fname: str) -> str:
            stop_words = {
                "resume", "cv", "curriculum", "vitae", "updated", "final",
                "new", "latest", "copy", "draft", "version", "doc",
                "v1", "v2", "v3", "v4", "v5",
                "2022", "2023", "2024", "2025", "2026",
            }
            base = os.path.splitext(fname)[0]
            base = re.sub(r"[\(\)\[\]_\-\.]", " ", base)
            base = re.sub(r"\s+", " ", base).strip()
            parts = []
            for word in base.split():
                if word.lower() in stop_words or word.isdigit():
                    break
                if re.match(r"^[A-Za-z]+$", word):
                    parts.append(word.title())
            return " ".join(parts) if len(parts) >= 1 else ""

        _filename_name = _name_from_filename(uploaded_file.name)
        _bad_name_values = {"not found", "n/a", "unknown", "none", ""}

        if candidate_name.lower() in _bad_name_values and _filename_name:
            candidate_name = _filename_name
        elif _filename_name and candidate_name.lower() not in _bad_name_values:
            llm_chars  = set(candidate_name.lower().replace(" ", ""))
            file_chars = set(_filename_name.lower().replace(" ", ""))
            overlap = len(llm_chars & file_chars) / max(len(file_chars), 1)
            if overlap < 0.4:
                candidate_name = _filename_name
        # ─────────────────────────────────────────────────────────────────────

        ats_score = ats_scores.get("ATS Match %", 0)
        edu_score = ats_scores.get("Education Score", 0)
        exp_score = ats_scores.get("Experience Score", 0)
        skills_score = ats_scores.get("Skills Score", 0)
        lang_score = ats_scores.get("Language Score", 0)
        keyword_score = ats_scores.get("Keyword Score", 0)
        fmt_score = ats_scores.get("Format Score", format_data.get("format_score", 0))
        formatted_score = ats_scores.get("Formatted Score", "N/A")
        fit_summary = ats_scores.get("Final Thoughts", "N/A")
        language_analysis_full = ats_scores.get("Language Analysis", "N/A")

        missing_keywords_raw = ats_scores.get("Missing Keywords", "N/A")
        missing_skills_raw = ats_scores.get("Missing Skills", "N/A")
        missing_keywords = [kw.strip() for kw in missing_keywords_raw.split(",") if kw.strip()] if missing_keywords_raw != "N/A" else []
        missing_skills = [sk.strip() for sk in missing_skills_raw.split(",") if sk.strip()] if missing_skills_raw != "N/A" else []

        bias_flag = "High Bias" if bias_score > 0.6 else "Fair"
        ats_flag  = "Low ATS"   if ats_score < 50   else "Good ATS"

        # Reuse domain already detected inside ats_percentage_score — no extra LLM call
        domain = ats_scores.get("Resume Domain", "Unknown")

        # ✅ Store everything in session state
        st.session_state.resume_data.append({
            "Resume Name": uploaded_file.name,
            "Candidate Name": candidate_name,
            "ATS Report": ats_result,
            "ATS Match %": ats_score,
            "Formatted Score": formatted_score,
            "Education Score": edu_score,
            "Experience Score": exp_score,
            "Skills Score": skills_score,
            "Language Score": lang_score,
            "Keyword Score": keyword_score,
            "Format Score": ats_scores.get("Format Score", 0),
            "Format Grade": ats_scores.get("Format Grade", "N/A"),
            "Format Label": ats_scores.get("Format Label", ""),
            "Format Issues": ats_scores.get("Format Issues", []),
            "Format Passes": ats_scores.get("Format Passes", []),
            "Education Analysis": ats_scores.get("Education Analysis", ""),
            "Experience Analysis": ats_scores.get("Experience Analysis", ""),
            "Skills Analysis": ats_scores.get("Skills Analysis", ""),
            "Language Analysis": language_analysis_full,
            "Keyword Analysis": ats_scores.get("Keyword Analysis", ""),
            "Format Analysis": ats_scores.get("Format Analysis", ""),
            "Final Thoughts": fit_summary,
            "Missing Keywords": missing_keywords,
            "Missing Skills": missing_skills,
            "Bias Score (0 = Fair, 1 = Biased)": bias_score,
            "Bias Status": bias_flag,
            "Masculine Words": masc_count,
            "Feminine Words": fem_count,
            "Detected Masculine Words": detected_masc,
            "Detected Feminine Words": detected_fem,
            "Text Preview": full_text[:300] + "...",
            "Highlighted Text": highlighted_text,
            "Rewritten Text": rewritten_text,
            "Optimized Resume Data": optimized_resume_data,
            "Domain": domain,
            "Domain Penalty": ats_scores.get("Domain Penalty", 0),
            "Domain Similarity Score": ats_scores.get("Domain Similarity Score", 1.0),
            "Resume Domain": ats_scores.get("Resume Domain", domain),
            "Job Domain": ats_scores.get("Job Domain", "Unknown"),
        })

        insert_candidate(
            (
                uploaded_file.name,
                candidate_name,
                ats_score,
                edu_score,
                exp_score,
                skills_score,
                lang_score,
                keyword_score,
                bias_score,
                fmt_score,   # ← format_score now saved to DB
            ),
            job_title=job_title,
            job_description=job_description
        )

        st.session_state.processed_files.add(uploaded_file.name)

        # ✅ IMPROVED: Smoother success animation with better transitions
        SUCCESS_HTML = """
        <style>
        .success-overlay {
            position: fixed;
            top: 0; left: 0;
            width: 100vw; height: 100vh;
            background: linear-gradient(135deg, #0b0c10 0%, #1a1c29 100%);
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            z-index: 9999;
            animation: fadeIn 0.5s ease-out;
        }
        
        @keyframes fadeIn {
            0% { opacity: 0; }
            100% { opacity: 1; }
        }
        
        .success-circle {
            width: 140px;
            height: 140px;
            border: 3px solid #00bfff;
            border-radius: 50%;
            position: relative;
            display: flex;
            align-items: center;
            justify-content: center;
            background: radial-gradient(circle, rgba(0,191,255,0.1) 0%, rgba(0,191,255,0.05) 50%, transparent 100%);
            animation: successPulse 2s ease-in-out infinite;
        }
        
        @keyframes successPulse {
            0%, 100% { 
                transform: scale(1);
                box-shadow: 0 0 20px rgba(0,191,255,0.3);
            }
            50% { 
                transform: scale(1.05);
                box-shadow: 0 0 30px rgba(0,191,255,0.6);
            }
        }
        
        .success-checkmark {
            font-size: 48px;
            color: #00ff7f;
            animation: checkmarkPop 0.8s ease-out;
        }
        
        @keyframes checkmarkPop {
            0% { transform: scale(0) rotate(-45deg); opacity: 0; }
            50% { transform: scale(1.2) rotate(-10deg); opacity: 0.8; }
            100% { transform: scale(1) rotate(0deg); opacity: 1; }
        }
        
        .success-text {
            margin-top: 25px;
            font-family: 'Orbitron', 'Segoe UI', sans-serif;
            font-size: 20px;
            font-weight: 600;
            color: #00bfff;
            text-shadow: 0 0 10px rgba(0,191,255,0.5);
            animation: textSlideUp 0.8s ease-out 0.3s both;
        }
        
        @keyframes textSlideUp {
            0% { transform: translateY(20px); opacity: 0; }
            100% { transform: translateY(0); opacity: 1; }
        }
        
        .success-subtitle {
            margin-top: 10px;
            font-size: 14px;
            color: #8e9aaf;
            animation: textSlideUp 0.8s ease-out 0.5s both;
        }
        </style>
        
        <div class="success-overlay">
            <div class="success-circle">
                <div class="success-checkmark">✓</div>
            </div>
            <div class="success-text">Scan Complete!</div>
            <div class="success-subtitle">Resume analysis ready</div>
        </div>
        """
        
        # Clear scanner and show success animation
        scanner_placeholder.empty()
        success_placeholder = st.empty()
        success_placeholder.markdown(SUCCESS_HTML, unsafe_allow_html=True)

        # ⏳ Shorter delay for better UX, then clear and rerun
        time.sleep(3)
        success_placeholder.empty()
        st.rerun()

    # ✅ Optional vectorstore setup
    if all_text:
        st.session_state.vectorstore = setup_vectorstore(all_text)
        st.session_state.chain = create_chain(st.session_state.vectorstore)

# 🔄 Developer Reset Button
with tab1:
    if st.button("🔄 Refresh view"):
        st.session_state.processed_files.clear()
        st.session_state.resume_data.clear()

        # Temporary placeholder for sliding success message
        msg_placeholder = st.empty()
        msg_placeholder.markdown("""
        <div class='slide-message success-msg'>
            <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
              stroke-width="2" viewBox="0 0 24 24"><path d="M5 13l4 4L19 7"/></svg>
            ✅ Cleared uploaded resume history. You can re-upload now.
        </div>
        """, unsafe_allow_html=True)

        # Wait 3 seconds then clear message
        time.sleep(3)
        msg_placeholder.empty()

def generate_resume_report_html(resume):
    candidate_name = resume.get('Candidate Name', 'Not Found')
    resume_name = resume.get('Resume Name', 'Unknown')
    rewritten_text = resume.get('Rewritten Text', '').replace("\n", "<br/>")

    masculine_words_list = resume.get("Detected Masculine Words", [])
    masculine_words = "".join(
        f"<b>{item.get('word','')}</b>: {item.get('sentence','')}<br/>"
        for item in masculine_words_list
    ) if masculine_words_list else "<i>None detected.</i>"

    feminine_words_list = resume.get("Detected Feminine Words", [])
    feminine_words = "".join(
        f"<b>{item.get('word','')}</b>: {item.get('sentence','')}<br/>"
        for item in feminine_words_list
    ) if feminine_words_list else "<i>None detected.</i>"

    ats_report_html = resume.get("ATS Report", "").replace("\n", "<br/>")

    def style_analysis(analysis, fallback="N/A"):
        if not analysis or analysis == "N/A":
            return f"<p><i>{fallback}</i></p>"

        if "**Score:**" in analysis:
            parts = analysis.split("**Score:**")
            rest = parts[1].split("**", 1)
            score_text = rest[0].strip()
            remaining = rest[1].strip() if len(rest) > 1 else ""
            return f"<p><b>Score:</b> {score_text}</p><p>{remaining}</p>"
        else:
            return f"<p>{analysis}</p>"

    edu_analysis = style_analysis(resume.get("Education Analysis", "").replace("\n", "<br/>"))
    exp_analysis = style_analysis(resume.get("Experience Analysis", "").replace("\n", "<br/>"))
    skills_analysis = style_analysis(resume.get("Skills Analysis", "").replace("\n", "<br/>"))
    keyword_analysis = style_analysis(resume.get("Keyword Analysis", "").replace("\n", "<br/>"))
    final_thoughts = resume.get("Final Thoughts", "N/A").replace("\n", "<br/>")

    lang_analysis_raw = resume.get("Language Analysis", "").replace("\n", "<br/>")
    lang_analysis = f"<div>{lang_analysis_raw}</div>" if lang_analysis_raw else "<p><i>No language analysis available.</i></p>"

    ats_match = resume.get('ATS Match %', 'N/A')
    edu_score = resume.get('Education Score', 'N/A')
    exp_score = resume.get('Experience Score', 'N/A')
    skills_score = resume.get('Skills Score', 'N/A')
    lang_score = resume.get('Language Score', 'N/A')
    keyword_score = resume.get('Keyword Score', 'N/A')
    format_score = resume.get('Format Score', 'N/A')
    format_grade = resume.get('Format Grade', 'N/A')
    format_label = resume.get('Format Label', '')
    masculine_count = len(masculine_words_list)
    feminine_count = len(feminine_words_list)
    bias_score = resume.get('Bias Score (0 = Fair, 1 = Biased)', 'N/A')

    return f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Helvetica, sans-serif;
                font-size: 12pt;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.4em;
            }}
            li {{
                margin-bottom: 5px;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 12px;
                border-left: 4px solid #666;
            }}
            .box {{
                padding: 10px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;
            }}
        </style>
    </head>
    <body>

    <h1>Resume Analysis Report</h1>

    <h2>Candidate: {candidate_name}</h2>
    <p><b>Resume File:</b> {resume_name}</p>

    <h2>ATS Evaluation</h2>
    <table>
        <tr><td><b>Overall ATS Match</b></td><td>{ats_match}%</td></tr>
        <tr><td><b>Education Score</b></td><td>{edu_score}</td></tr>
        <tr><td><b>Experience Score</b></td><td>{exp_score}</td></tr>
        <tr><td><b>Skills Score</b></td><td>{skills_score}</td></tr>
        <tr><td><b>Language Score</b></td><td>{lang_score}</td></tr>
        <tr><td><b>Keyword Score</b></td><td>{keyword_score}</td></tr>
        <tr><td><b>Format Score</b></td><td>{format_score}/100 — {format_grade} ({format_label})</td></tr>
    </table>

    <div class="section-title">ATS Report</div>
    <div class="box">{ats_report_html}</div>

    <div class="section-title">Education Analysis</div>
    <div class="box">{edu_analysis}</div>

    <div class="section-title">Experience Analysis</div>
    <div class="box">{exp_analysis}</div>

    <div class="section-title">Skills Analysis</div>
    <div class="box">{skills_analysis}</div>

    <div class="section-title">Language Analysis</div>
    <div class="box">{lang_analysis}</div>

    <div class="section-title">Keyword Analysis</div>
    <div class="box">{keyword_analysis}</div>

    <div class="section-title">Final Thoughts</div>
    <div class="box">{final_thoughts}</div>

    <h2>Gender Bias Analysis</h2>
    <table>
        <tr><td><b>Masculine Words</b></td><td>{masculine_count}</td></tr>
        <tr><td><b>Feminine Words</b></td><td>{feminine_count}</td></tr>
        <tr><td><b>Bias Score (0 = Fair, 1 = Biased)</b></td><td>{bias_score}</td></tr>
    </table>

    <div class="section-title">Masculine Words Detected</div>
    <div class="box">{masculine_words}</div>

    <div class="section-title">Feminine Words Detected</div>
    <div class="box">{feminine_words}</div>

    <h2>Rewritten Bias-Free Resume</h2>
    <div class="box">{rewritten_text}</div>

    </body>
    </html>
    """

# === TAB 1: Dashboard ===
with tab1:
    resume_data = st.session_state.get("resume_data", [])

    if resume_data:
        # ✅ Calculate total counts safely
        total_masc = sum(len(r.get("Detected Masculine Words", [])) for r in resume_data)
        total_fem = sum(len(r.get("Detected Feminine Words", [])) for r in resume_data)
        avg_bias = round(np.mean([r.get("Bias Score (0 = Fair, 1 = Biased)", 0) for r in resume_data]), 2)
        total_resumes = len(resume_data)

        st.markdown("<p class='section-label'>Session Summary</p>", unsafe_allow_html=True)
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Resumes Uploaded", total_resumes)
        with col2:
            st.metric("Avg. Bias Score", avg_bias)
        with col3:
            st.metric("Total Masculine Words", total_masc)
        with col4:
            st.metric("Total Feminine Words", total_fem)

        st.markdown("<p class='section-label'>Resumes Overview</p>", unsafe_allow_html=True)
        df = pd.DataFrame(resume_data)

        # ✅ Add calculated count columns safely
        df["Masculine Words Count"] = df["Detected Masculine Words"].apply(lambda x: len(x) if isinstance(x, list) else 0)
        df["Feminine Words Count"] = df["Detected Feminine Words"].apply(lambda x: len(x) if isinstance(x, list) else 0)

        overview_cols = [
            "Resume Name", "Candidate Name", "ATS Match %", "Education Score",
            "Experience Score", "Skills Score", "Language Score", "Keyword Score",
            "Format Score",
            "Bias Score (0 = Fair, 1 = Biased)", "Masculine Words Count", "Feminine Words Count"
        ]

        st.dataframe(df[overview_cols], use_container_width=True)

        st.markdown("<p class='section-label'>Visual Analysis</p>", unsafe_allow_html=True)
        chart_tab1, chart_tab2 = st.tabs(["Bias Score Chart", "Gender-Coded Words"])
        with chart_tab1:
            st.subheader("Bias Score Comparison Across Resumes")
            bias_chart_df = df[["Resume Name", "Bias Score (0 = Fair, 1 = Biased)"]].copy()
            bias_chart_df.columns = ["Resume", "Bias Score"]
            bias_altair = alt.Chart(bias_chart_df).mark_bar(
                cornerRadiusTopLeft=4,
                cornerRadiusTopRight=4,
                color="#4f8cff"
            ).encode(
                x=alt.X("Resume:N", sort=None, axis=alt.Axis(labelAngle=-35, labelFontSize=11, titleFontSize=12)),
                y=alt.Y("Bias Score:Q", scale=alt.Scale(domain=[0, 1]), axis=alt.Axis(titleFontSize=12)),
                tooltip=["Resume", alt.Tooltip("Bias Score:Q", format=".2f")]
            ).properties(height=260).configure_view(strokeWidth=0).configure_axis(
                grid=False, domainColor="#2d3748"
            )
            st.altair_chart(bias_altair, use_container_width=True)
        with chart_tab2:
            st.subheader("Masculine vs Feminine Word Usage")
            gender_df = pd.DataFrame({
                "Resume": list(df["Resume Name"]) * 2,
                "Type": ["Masculine"] * len(df) + ["Feminine"] * len(df),
                "Count": list(df["Masculine Words Count"]) + list(df["Feminine Words Count"])
            })
            color_scale = alt.Scale(domain=["Masculine", "Feminine"], range=["#4f8cff", "#fb7185"])
            gender_altair = alt.Chart(gender_df).mark_bar(cornerRadiusTopLeft=3, cornerRadiusTopRight=3).encode(
                x=alt.X("Resume:N", sort=None, axis=alt.Axis(labelAngle=-35, labelFontSize=11, titleFontSize=12)),
                y=alt.Y("Count:Q", axis=alt.Axis(titleFontSize=12)),
                color=alt.Color("Type:N", scale=color_scale, legend=alt.Legend(orient="top", titleFontSize=11)),
                xOffset="Type:N",
                tooltip=["Resume", "Type", "Count"]
            ).properties(height=260).configure_view(strokeWidth=0).configure_axis(
                grid=False, domainColor="#2d3748"
            )
            st.altair_chart(gender_altair, use_container_width=True)

        st.markdown("<p class='section-label'>Detailed Resume Reports</p>", unsafe_allow_html=True)
        for resume in resume_data:
            candidate_name = resume.get("Candidate Name", "Not Found")
            resume_name = resume.get("Resume Name", "Unknown")
            missing_keywords = resume.get("Missing Keywords", [])
            missing_skills = resume.get("Missing Skills", [])

            with st.expander(f"{resume_name} | {candidate_name}"):
                st.markdown(f"""
                <div style="
                    background: linear-gradient(135deg, rgba(56,189,248,0.10) 0%, rgba(79,163,227,0.05) 100%);
                    border: 1px solid rgba(56,189,248,0.18);
                    border-radius: 14px;
                    padding: 18px 22px;
                    margin-bottom: 20px;
                ">
                    <div style="
                        font-family: -apple-system, BlinkMacSystemFont, 'SF Pro Display', sans-serif;
                        font-size: 1rem;
                        font-weight: 700;
                        color: #f0f4f8;
                        letter-spacing: -0.01em;
                    ">ATS Evaluation — <span style='color:#38bdf8;'>{candidate_name}</span></div>
                    <div style="font-size:0.75rem; color:#64748b; margin-top:4px; font-family: -apple-system, sans-serif; text-transform:uppercase; letter-spacing:0.05em;">Resume Intelligence Report</div>
                </div>
                """, unsafe_allow_html=True)

                # ── SVG icon helper ──────────────────────────────────────────────
                SVG_ICONS = {
                    "overall": '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg>',
                    "grade":   '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M22 11.08V12a10 10 0 1 1-5.93-9.14"/><polyline points="22 4 12 14.01 9 11.01"/></svg>',
                    "edu":     '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M22 10v6M2 10l10-5 10 5-10 5z"/><path d="M6 12v5c3 3 9 3 12 0v-5"/></svg>',
                    "exp":     '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="2" y="7" width="20" height="14" rx="2"/><path d="M16 7V5a2 2 0 0 0-2-2h-4a2 2 0 0 0-2 2v2"/></svg>',
                    "skills":  '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="3"/><path d="M12 1v4M12 19v4M4.22 4.22l2.83 2.83M16.95 16.95l2.83 2.83M1 12h4M19 12h4M4.22 19.78l2.83-2.83M16.95 7.05l2.83-2.83"/></svg>',
                    "lang":    '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M21 15a2 2 0 0 1-2 2H7l-4 4V5a2 2 0 0 1 2-2h14a2 2 0 0 1 2 2z"/></svg>',
                    "keyword": '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="11" cy="11" r="8"/><line x1="21" y1="21" x2="16.65" y2="16.65"/></svg>',
                    "format":  '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="9 11 12 14 22 4"/><path d="M21 12v7a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11"/></svg>',
                    "pass":    '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#34d399" stroke-width="2.5"><polyline points="20 6 9 17 4 12"/></svg>',
                    "fail":    '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#f87171" stroke-width="2.5"><line x1="18" y1="6" x2="6" y2="18"/><line x1="6" y1="6" x2="18" y2="18"/></svg>',
                }

                def svg_ats_card(svg_key, label, value, tooltip=None):
                    tooltip_attr = f'title="{tooltip}"' if tooltip else ""
                    return f"""
                    <div style="
                        background: rgba(15,23,42,0.85);
                        border: 1px solid rgba(56,189,248,0.25);
                        border-radius: 12px;
                        padding: 14px 16px;
                        margin-bottom: 8px;
                        height: 86px;
                        display: flex;
                        flex-direction: column;
                        justify-content: center;
                        overflow: hidden;
                        box-sizing: border-box;
                    ">
                        <div style="display:flex;align-items:center;gap:6px;font-size:0.72rem; color:#94a3b8;">
                            <span style="color:#38bdf8;flex-shrink:0;">{SVG_ICONS.get(svg_key,"")}</span>
                            <span style="white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{label}</span>
                        </div>
                        <div {tooltip_attr} style="font-size:1.35rem; font-weight:700; color:#f0f4f8; margin-top:6px; white-space:nowrap; overflow:hidden; text-overflow:ellipsis;">{value}</div>
                    </div>"""

                # ── Overall Score Gauge (SVG) ────────────────────────────────────
                overall_pct = resume.get("ATS Match %", 0)
                fmt_score   = resume.get("Format Score", 0)
                fmt_grade   = resume.get("Format Grade", "N/A")
                fmt_label   = resume.get("Format Label", "")

                # Gauge colour
                if overall_pct >= 75:
                    gauge_color = "#22c55e"
                elif overall_pct >= 55:
                    gauge_color = "#f59e0b"
                else:
                    gauge_color = "#ef4444"

                # SVG arc gauge
                radius = 70
                cx, cy = 90, 90
                circumference = 3.14159 * radius  # half-circle arc = π*r
                arc_offset = circumference * (1 - overall_pct / 100)

                st.markdown(f"""
                <div style="display:flex;align-items:center;gap:32px;padding:20px 24px;
                            background:rgba(15,23,42,0.9);border:1px solid rgba(56,189,248,0.2);
                            border-radius:16px;margin-bottom:20px;flex-wrap:wrap;">
                    <!-- Gauge -->
                    <div style="flex-shrink:0;text-align:center;">
                        <svg width="180" height="100" viewBox="0 0 180 100">
                            <!-- Track -->
                            <path d="M 20 90 A 70 70 0 0 1 160 90" fill="none" stroke="rgba(255,255,255,0.08)" stroke-width="12" stroke-linecap="round"/>
                            <!-- Value arc -->
                            <path d="M 20 90 A 70 70 0 0 1 160 90" fill="none"
                                stroke="{gauge_color}" stroke-width="12" stroke-linecap="round"
                                stroke-dasharray="{circumference}" stroke-dashoffset="{arc_offset}"
                                style="transition:stroke-dashoffset 0.8s ease;"/>
                            <!-- Score text -->
                            <text x="90" y="80" text-anchor="middle" font-size="28" font-weight="700"
                                fill="{gauge_color}" font-family="-apple-system,sans-serif">{overall_pct}</text>
                            <text x="90" y="98" text-anchor="middle" font-size="11" fill="#64748b"
                                font-family="-apple-system,sans-serif">/ 100</text>
                        </svg>
                        <div style="font-size:0.75rem;color:#64748b;margin-top:2px;font-family:-apple-system,sans-serif;letter-spacing:0.04em;text-transform:uppercase;">Overall ATS Score</div>
                    </div>
                    <!-- Label & Format quick-look -->
                    <div style="flex:1;min-width:200px;">
                        <div style="font-size:1.1rem;font-weight:700;color:#f0f4f8;font-family:-apple-system,sans-serif;">{resume.get("Formatted Score","N/A")}</div>
                        <div style="margin-top:12px;display:flex;align-items:center;gap:10px;">
                            <span style="color:#38bdf8;">{SVG_ICONS["format"]}</span>
                            <span style="font-size:0.82rem;color:#94a3b8;">Format Score:</span>
                            <span style="font-size:0.95rem;font-weight:700;color:#f0f4f8;">{fmt_score}/100</span>
                            <span style="background:rgba(56,189,248,0.12);border:1px solid rgba(56,189,248,0.25);
                                        border-radius:6px;padding:2px 8px;font-size:0.75rem;font-weight:700;color:#38bdf8;">{fmt_grade}</span>
                        </div>
                        <div style="margin-top:6px;font-size:0.78rem;color:#64748b;">{fmt_label}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # ── Score cards row 1 ──────────────────────────────────────────
                formatted_val = resume.get("Formatted Score", "N/A")
                score_col1, score_col2, score_col3 = st.columns(3)
                with score_col1:
                    st.markdown(svg_ats_card("overall", "Overall ATS Match", f"{resume.get('ATS Match %', 'N/A')}%"), unsafe_allow_html=True)
                with score_col2:
                    st.markdown(svg_ats_card("grade", "Hire Signal", formatted_val, tooltip=formatted_val), unsafe_allow_html=True)
                with score_col3:
                    st.markdown(svg_ats_card("lang", "Language Quality", f"{resume.get('Language Score', 'N/A')} / {lang_weight}"), unsafe_allow_html=True)

                # ── Score cards row 2 ──────────────────────────────────────────
                col_a, col_b, col_c, col_d = st.columns(4)
                with col_a:
                    st.markdown(svg_ats_card("edu", "Education", f"{resume.get('Education Score', 'N/A')} / {edu_weight}"), unsafe_allow_html=True)
                with col_b:
                    st.markdown(svg_ats_card("exp", "Experience", f"{resume.get('Experience Score', 'N/A')} / {exp_weight}"), unsafe_allow_html=True)
                with col_c:
                    st.markdown(svg_ats_card("skills", "Skills", f"{resume.get('Skills Score', 'N/A')} / {skills_weight}"), unsafe_allow_html=True)
                with col_d:
                    st.markdown(svg_ats_card("keyword", "Keywords", f"{resume.get('Keyword Score', 'N/A')} / {keyword_weight}"), unsafe_allow_html=True)

                # ── Score cards row 3: bias + domain status ────────────────────
                SVG_BIAS  = '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>'
                SVG_DOM   = '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="10"/><line x1="2" y1="12" x2="22" y2="12"/><path d="M12 2a15.3 15.3 0 0 1 4 10 15.3 15.3 0 0 1-4 10 15.3 15.3 0 0 1-4-10 15.3 15.3 0 0 1 4-10z"/></svg>'

                bias_raw   = resume.get("Bias Score (0 = Fair, 1 = Biased)", 0)
                bias_pct   = round(bias_raw * 100)
                bias_label = "High Bias" if bias_raw > 0.6 else ("Moderate" if bias_raw > 0.3 else "Fair")
                bias_color = "#ef4444" if bias_raw > 0.6 else ("#f59e0b" if bias_raw > 0.3 else "#22c55e")

                dom_penalty = resume.get("Domain Penalty", 0)
                dom_penalty = dom_penalty if isinstance(dom_penalty, (int, float)) else 0
                dom_sim     = resume.get("Domain Similarity Score", 1.0)
                dom_sim     = dom_sim if isinstance(dom_sim, (int, float)) else 1.0
                dom_pct     = round(dom_sim * 100)
                dom_label   = resume.get("Resume Domain", resume.get("Domain", "Unknown"))

                r3c1, r3c2, r3c3, r3c4 = st.columns(4)
                with r3c1:
                    st.markdown(f"""
                    <div style="background:rgba(15,23,42,0.85);border:1px solid rgba(56,189,248,0.25);
                                border-radius:12px;padding:14px 16px;margin-bottom:8px;height:86px;
                                display:flex;flex-direction:column;justify-content:center;overflow:hidden;">
                        <div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;color:#94a3b8;">
                            <span style="color:{bias_color};flex-shrink:0;">{SVG_BIAS}</span>
                            <span>Bias Status</span>
                        </div>
                        <div style="font-size:1.1rem;font-weight:700;color:{bias_color};margin-top:6px;">
                            {bias_label} <span style="font-size:0.8rem;color:#64748b;">({bias_pct}%)</span>
                        </div>
                    </div>""", unsafe_allow_html=True)
                with r3c2:
                    st.markdown(f"""
                    <div style="background:rgba(15,23,42,0.85);border:1px solid rgba(56,189,248,0.25);
                                border-radius:12px;padding:14px 16px;margin-bottom:8px;height:86px;
                                display:flex;flex-direction:column;justify-content:center;overflow:hidden;">
                        <div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;color:#94a3b8;">
                            <span style="color:#38bdf8;flex-shrink:0;">{SVG_DOM}</span>
                            <span>Domain Match</span>
                        </div>
                        <div style="font-size:1.1rem;font-weight:700;color:#f0f4f8;margin-top:6px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">
                            {dom_pct}% <span style="font-size:0.75rem;color:#64748b;">(-{dom_penalty} pts)</span>
                        </div>
                    </div>""", unsafe_allow_html=True)
                with r3c3:
                    st.markdown(svg_ats_card("format", "Format Score", f"{resume.get('Format Score', 'N/A')}/100 · {resume.get('Format Grade','N/A')}"), unsafe_allow_html=True)
                with r3c4:
                    masc_c = len(resume.get("Detected Masculine Words", []))
                    fem_c  = len(resume.get("Detected Feminine Words", []))
                    SVG_WORDS = '<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M21 15a2 2 0 0 1-2 2H7l-4 4V5a2 2 0 0 1 2-2h14a2 2 0 0 1 2 2z"/></svg>'
                    st.markdown(f"""
                    <div style="background:rgba(15,23,42,0.85);border:1px solid rgba(56,189,248,0.25);
                                border-radius:12px;padding:14px 16px;margin-bottom:8px;height:86px;
                                display:flex;flex-direction:column;justify-content:center;overflow:hidden;">
                        <div style="display:flex;align-items:center;gap:6px;font-size:0.72rem;color:#94a3b8;">
                            <span style="color:#38bdf8;flex-shrink:0;">{SVG_WORDS}</span>
                            <span>Gender Words</span>
                        </div>
                        <div style="font-size:0.95rem;font-weight:700;color:#f0f4f8;margin-top:6px;">
                            <span style="color:#60a5fa;">{masc_c} M</span>
                            <span style="color:#64748b;margin:0 4px;">/</span>
                            <span style="color:#f87171;">{fem_c} F</span>
                        </div>
                    </div>""", unsafe_allow_html=True)

                # ── Format Checker Panel ───────────────────────────────────────
                fmt_issues = resume.get("Format Issues", [])
                fmt_passes = resume.get("Format Passes", [])
                st.markdown("""
                <div style="margin:16px 0 6px;font-size:0.72rem;font-weight:700;color:#64748b;
                            letter-spacing:0.08em;text-transform:uppercase;font-family:-apple-system,sans-serif;">
                    Format &amp; ATS Compatibility Check
                </div>""", unsafe_allow_html=True)
                
                issues_html = "".join(
                    f"<div style='display:flex;align-items:flex-start;gap:6px;margin-bottom:5px;font-size:0.8rem;color:#fca5a5;'>{SVG_ICONS['fail']}<span>{iss}</span></div>"
                    for iss in fmt_issues
                ) if fmt_issues else "<div style='font-size:0.8rem;color:#94a3b8;'>No critical issues detected.</div>"
                passes_html = "".join(
                    f"<div style='display:flex;align-items:flex-start;gap:6px;margin-bottom:5px;font-size:0.8rem;color:#6ee7b7;'>{SVG_ICONS['pass']}<span>{p}</span></div>"
                    for p in fmt_passes
                ) if fmt_passes else ""

                fmt_col1, fmt_col2 = st.columns(2)
                with fmt_col1:
                    st.markdown(f"""
                    <div style="background:rgba(239,68,68,0.06);border:1px solid rgba(239,68,68,0.2);
                                border-radius:10px;padding:12px 14px;">
                        <div style="font-size:0.72rem;font-weight:700;color:#f87171;text-transform:uppercase;
                                    letter-spacing:0.06em;margin-bottom:8px;display:flex;align-items:center;gap:6px;">
                            {SVG_ICONS['fail']} Issues ({len(fmt_issues)})
                        </div>
                        {issues_html}
                    </div>""", unsafe_allow_html=True)
                with fmt_col2:
                    st.markdown(f"""
                    <div style="background:rgba(52,211,153,0.06);border:1px solid rgba(52,211,153,0.2);
                                border-radius:10px;padding:12px 14px;">
                        <div style="font-size:0.72rem;font-weight:700;color:#34d399;text-transform:uppercase;
                                    letter-spacing:0.06em;margin-bottom:8px;display:flex;align-items:center;gap:6px;">
                            {SVG_ICONS['pass']} Passed ({len(fmt_passes)})
                        </div>
                        {passes_html}
                    </div>""", unsafe_allow_html=True)

                # Fit summary
                st.markdown("""
                <div style="margin:18px 0 6px;font-size:0.72rem;font-weight:700;color:#64748b;
                            letter-spacing:0.08em;text-transform:uppercase;font-family:-apple-system,sans-serif;">
                    Fit Summary
                </div>""", unsafe_allow_html=True)
                st.write(resume.get('Final Thoughts', 'N/A'))

                # ATS Report
                if resume.get("ATS Report"):
                    st.markdown("<p class='section-label'>ATS Evaluation Report</p>", unsafe_allow_html=True)
                    st.markdown(resume["ATS Report"], unsafe_allow_html=True)

                # ATS Chart
                st.markdown("<p class='section-label'>ATS Score Breakdown</p>", unsafe_allow_html=True)
                # Normalize each component score to 0–100 scale for fair visual comparison
                def _pct(score, weight):
                    return round(score / weight * 100) if weight > 0 else 0
                ats_df = pd.DataFrame({
                    'Component': ['Education', 'Experience', 'Skills', 'Language', 'Keywords', 'Format'],
                    'Score': [
                        _pct(resume.get("Education Score", 0), edu_weight),
                        _pct(resume.get("Experience Score", 0), exp_weight),
                        _pct(resume.get("Skills Score", 0), skills_weight),
                        _pct(resume.get("Language Score", 0), lang_weight) if lang_weight > 0 else 0,
                        _pct(resume.get("Keyword Score", 0), keyword_weight),
                        resume.get("Format Score", 0),  # Already on 0–100 scale
                    ]
                })
                ats_chart = alt.Chart(ats_df).mark_bar().encode(
                    x=alt.X('Component', sort=None),
                    y=alt.Y('Score', scale=alt.Scale(domain=[0, 100]), title='Score (% of weight)'),
                    color='Component',
                    tooltip=['Component', 'Score']
                ).properties(
                    title="ATS Evaluation Breakdown (All scores normalized to 0–100%)",
                    width=600,
                    height=300
                )
                st.altair_chart(ats_chart, use_container_width=True)

                st.markdown("<p class='section-label'>Detailed ATS Section Analyses</p>", unsafe_allow_html=True)
                for section_title, key in [
                    ("Education Analysis", "Education Analysis"),
                    ("Experience Analysis", "Experience Analysis"),
                    ("Skills Analysis", "Skills Analysis"),
                    ("Language Quality", "Language Analysis"),
                    ("Keyword Analysis", "Keyword Analysis"),
                    ("Format & ATS Compatibility", "Format Analysis"),
                    ("Final Assessment", "Final Thoughts")
                ]:
                    analysis_content = resume.get(key, "N/A")
                    if "**Score:**" in analysis_content:
                        parts = analysis_content.split("**Score:**")
                        rest = parts[1].split("**", 1)
                        score_text = rest[0].strip()
                        remaining = rest[1].strip() if len(rest) > 1 else ""
                        score_html = f"<span class='score-badge'>Score: {score_text}</span>"
                        body_html = f"{score_html}<div style='margin-top:8px;'>{remaining}</div>"
                    else:
                        body_html = f"<div>{analysis_content}</div>"

                    st.markdown(f"""
<div class="ats-section-header">{section_title}</div>
<div class="ats-section-body">{body_html}</div>
""", unsafe_allow_html=True)

                st.divider()

                detail_tab1, detail_tab2 = st.tabs(["Bias Analysis", "Rewritten Resume"])

                with detail_tab1:
                    st.markdown("""
                    <div style="display:flex;align-items:center;gap:8px;margin:12px 0 6px;">
                        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2"><circle cx="11" cy="11" r="8"/><line x1="21" y1="21" x2="16.65" y2="16.65"/></svg>
                        <span class='section-label' style="margin:0;">Bias-Highlighted Original Text</span>
                    </div>""", unsafe_allow_html=True)
                    st.markdown(resume["Highlighted Text"], unsafe_allow_html=True)

                    st.markdown("""
                    <div style="display:flex;align-items:center;gap:8px;margin:14px 0 6px;">
                        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2"><line x1="8" y1="6" x2="21" y2="6"/><line x1="8" y1="12" x2="21" y2="12"/><line x1="8" y1="18" x2="21" y2="18"/><line x1="3" y1="6" x2="3.01" y2="6"/><line x1="3" y1="12" x2="3.01" y2="12"/><line x1="3" y1="18" x2="3.01" y2="18"/></svg>
                        <span class='section-label' style="margin:0;">Gender-Coded Word Counts</span>
                    </div>""", unsafe_allow_html=True)
                    bias_col1, bias_col2 = st.columns(2)

                    with bias_col1:
                        st.metric("Masculine Words", len(resume["Detected Masculine Words"]))
                        if resume["Detected Masculine Words"]:
                            st.markdown("<p class='section-label'>Masculine Words with Context</p>", unsafe_allow_html=True)
                            for item in resume["Detected Masculine Words"]:
                                word = item['word']
                                sentence = item['sentence']
                                st.markdown(f"""<div style='margin-bottom:6px;font-size:0.85rem;'>
                                    <span style='color:#60a5fa;font-weight:600;'>{word}</span>: {sentence}</div>""",
                                    unsafe_allow_html=True)
                        else:
                            st.info("No masculine words detected.")

                    with bias_col2:
                        st.metric("Feminine Words", len(resume["Detected Feminine Words"]))
                        if resume["Detected Feminine Words"]:
                            st.markdown("<p class='section-label'>Feminine Words with Context</p>", unsafe_allow_html=True)
                            for item in resume["Detected Feminine Words"]:
                                word = item['word']
                                sentence = item['sentence']
                                st.markdown(f"""<div style='margin-bottom:6px;font-size:0.85rem;'>
                                    <span style='color:#f87171;font-weight:600;'>{word}</span>: {sentence}</div>""",
                                    unsafe_allow_html=True)
                        else:
                            st.info("No feminine words detected.")

                with detail_tab2:
                    st.markdown("""
                    <div style="display:flex;align-items:center;gap:8px;margin:12px 0 6px;">
                        <svg width="15" height="15" viewBox="0 0 24 24" fill="none" stroke="#34d399" stroke-width="2"><polyline points="9 11 12 14 22 4"/><path d="M21 12v7a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11"/></svg>
                        <span class='section-label' style="margin:0;">Bias-Free Rewritten Resume</span>
                    </div>""", unsafe_allow_html=True)

                    # ── Job Title Suggestions (Analysis Module — displayed here, NOT in DOCX) ──
                    rewritten_raw = resume.get("Rewritten Text", "")
                    if "### 🎯 Suggested Job Titles" in rewritten_raw:
                        split_parts = rewritten_raw.split("### 🎯 Suggested Job Titles")
                        resume_text_display = split_parts[0].strip()
                        job_suggestions_display = "### 🎯 Suggested Job Titles" + split_parts[1]
                    else:
                        resume_text_display = rewritten_raw
                        job_suggestions_display = ""

                    st.write(resume_text_display)

                    if job_suggestions_display:
                        st.markdown("""
                        <div style="margin:18px 0 8px;font-size:0.72rem;font-weight:700;color:#64748b;
                                    letter-spacing:0.08em;text-transform:uppercase;font-family:-apple-system,sans-serif;">
                            Job Title Suggestions (for reference only — not included in resume files)
                        </div>""", unsafe_allow_html=True)
                        st.markdown(job_suggestions_display)

                    # ── 3-Template DOCX Download Buttons (Optimization Module — JSON data only) ──
                    st.markdown("""
                    <div style="margin:20px 0 10px;font-size:0.72rem;font-weight:700;color:#64748b;
                                letter-spacing:0.08em;text-transform:uppercase;font-family:-apple-system,sans-serif;">
                        Download Optimized Resume — Choose Template
                    </div>""", unsafe_allow_html=True)

                    optimized_data = resume.get("Optimized Resume Data", {})
                    base_name = resume['Resume Name'].split('.')[0]

                    dl_col1, dl_col2, dl_col3 = st.columns(3)

                    with dl_col1:
                        try:
                            modern_buf = generate_modern_docx(optimized_data)
                            st.download_button(
                                label="⬇ Modern (ATS)",
                                data=modern_buf,
                                file_name=f"{base_name}_modern_ats.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key=f"dl_modern_{resume['Resume Name']}",
                                help="Navy headings · Calibri · Labeled Skills block · ATS section order · Workday/Greenhouse optimized"
                            )
                        except Exception as e:
                            st.error(f"Modern template error: {e}")

                    with dl_col2:
                        try:
                            minimal_buf = generate_minimal_docx(optimized_data)
                            st.download_button(
                                label="⬇ Minimal (ATS)",
                                data=minimal_buf,
                                file_name=f"{base_name}_minimal_ats.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key=f"dl_minimal_{resume['Resume Name']}",
                                help="Pure black/white Arial · Maximum parse accuracy · Taleo/iCIMS/SmartRecruiters compatible"
                            )
                        except Exception as e:
                            st.error(f"Minimal template error: {e}")

                    with dl_col3:
                        try:
                            creative_buf = generate_creative_docx(optimized_data)
                            st.download_button(
                                label="⬇ Executive (ATS)",
                                data=creative_buf,
                                file_name=f"{base_name}_executive_ats.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key=f"dl_creative_{resume['Resume Name']}",
                                help="Teal/navy accents · ATS-safe symbols · Consistent Calibri body · Standard section labels"
                            )
                        except Exception as e:
                            st.error(f"Executive template error: {e}")

                    html_report = generate_resume_report_html(resume)
                    pdf_file = html_to_pdf_bytes(html_report)
                    st.download_button(
                        label="Download Full Analysis Report (.pdf)",
                        data=pdf_file,
                        file_name=f"{base_name}_report.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key=f"download_pdf_{resume['Resume Name']}"
                    )

    else:           
        st.warning("⚠️ Please upload resumes to view dashboard analytics.")
